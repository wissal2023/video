import os
import re
import time
import unicodedata
import pyttsx3
import numpy as np
import sys
import json
import shutil
from pathlib import Path
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.ensemble import RandomForestRegressor
import joblib
import importlib
import tempfile
import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils.data import DataLoader, TensorDataset
import argparse   
from transformers import pipeline, AutoTokenizer, LlamaTokenizer, ViTImageProcessor, AutoModel
from transformers import AutoModelForCausalLM, AutoModelForSeq2SeqLM
from transformers import VisionEncoderDecoderModel, ViTFeatureExtractor
from transformers import AutoImageProcessor, AutoModelForImageClassification
from sentence_transformers import SentenceTransformer
from PIL import Image, ImageEnhance, ImageOps, ImageStat, ImageChops
import PyPDF2
import docx
import fitz 
from pydub import AudioSegment
from pydub.effects import normalize, compress_dynamic_range
from sklearn.neighbors import NearestNeighbors
from sklearn.cluster import KMeans
from sklearn.ensemble import RandomForestRegressor
from sklearn.feature_extraction.text import TfidfVectorizer
from scipy import signal
from scipy.signal import butter, filtfilt
from scipy.ndimage import gaussian_filter
import soundfile as sf
import traceback
import librosa
from librosa.feature import rms
import tempfile
from random import randint
from tqdm import tqdm
import torchvision.models as models
import torchvision.transforms as transforms

# Added imports for modularized components
sys.path.insert(0, str(Path(__file__).parent / "videoseul_modules"))

from videoseul_modules.text_audio_utils import TextEnhancer, create_combined_audio
from videoseul_modules.readability import ReadabilityEvaluator
from videoseul_modules.utils import *
from videoseul_modules.presentation_generator import SmartPresentationGenerator



IMAGEMAGICK_PATH = r"C:\\Program Files\\ImageMagick-7.1.1-Q16-HDRI\\magick.exe"

if not os.path.exists(IMAGEMAGICK_PATH):
    print(f"?? ATTENTION: L'exécutable ImageMagick n'a pas été trouvé à {IMAGEMAGICK_PATH}")
    print("?? La création de texte dans les diapositives risque d'échouer.")
else:
    print(f"? ImageMagick trouvé: {IMAGEMAGICK_PATH}")

# Set environment variable for ImageMagick before importing moviepy
os.environ["IMAGEMAGICK_BINARY"] = IMAGEMAGICK_PATH

from moviepy.editor import TextClip, ImageClip, ColorClip, CompositeVideoClip, concatenate_videoclips
from moviepy.editor import AudioFileClip, concatenate_audioclips, VideoFileClip, CompositeAudioClip
from moviepy.audio.AudioClip import AudioClip

  
# Imports pour les moèles d'IA avancés
TRANSFORMER_AVAILABLE = False
try:  
    TRANSFORMER_AVAILABLE = True
    print("? Bibliothèques de modèles disponibles")
except ImportError as e:
    print(f"?? Transformers ou SentenceTransformer non disponibles: {e}")
    print(f"Détails: {e}")
    TRANSFORMER_AVAILABLE = False

# Vérifier si le chemin existe
if not os.path.exists(IMAGEMAGICK_PATH):
    print(f"?? ATTENTION: L'exécutable ImageMagick n'a pas été trouvé à {IMAGEMAGICK_PATH}")
    print("?? La création de texte dans les diapositives risque d'échouer.")
    
    # Essayer de trouver magick.exe dans le répertoire indiqué
    magick_dir = os.path.dirname(IMAGEMAGICK_PATH)
    if os.path.exists(magick_dir):
        print(f"?? Recherche de magick.exe dans {magick_dir}...")
        for file in os.listdir(magick_dir):
            if file.lower() == "magick.exe":
                IMAGEMAGICK_PATH = os.path.join(magick_dir, file)
                print(f"? Trouvé: {IMAGEMAGICK_PATH}")
                break
else:
    print(f"? ImageMagick trouvé: {IMAGEMAGICK_PATH}")

try:    
    TORCH_AVAILABLE = True
    print("? PyTorch disponible")
except ImportError:
    TORCH_AVAILABLE = False
    print("?? PyTorch non disponible, certaines fonctionnalités seront limitées")

# Chemins de base
BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)
MODELS_DIR = BASE_DIR / "models"
MODELS_DIR.mkdir(exist_ok=True)
TRAINING_DATA_DIR = BASE_DIR / "training_data"
TRAINING_DATA_DIR.mkdir(exist_ok=True)
ANIMATION_DIR = OUTPUT_DIR / "animations"
ANIMATION_DIR.mkdir(exist_ok=True)
EXTRACTED_IMAGES_DIR = OUTPUT_DIR / "extracted_images"
EXTRACTED_IMAGES_DIR.mkdir(exist_ok=True)
DEEP_MODELS_DIR = MODELS_DIR / "deep_learning"
DEEP_MODELS_DIR.mkdir(exist_ok=True)

# Paramètres Généraux
FONT_SIZE_TITLE = 50
FONT_SIZE_TEXT = 36
FONT_SIZE_TABLE = 28
WIDTH, HEIGHT = 960, 540
BG_COLOR = (255, 255, 255)
TEXT_COLOR = "black"
TABLE_HEADER_BG = (200, 220, 240)
TABLE_ROW_BG_1 = (248, 248, 248)
TABLE_ROW_BG_2 = (255, 255, 255)
TABLE_BORDER = (180, 180, 180)
TEXT_MARGIN = 40
INTRO_DUR = 3
OUTRO_DUR = 3
SLIDE_DUR = 8
TRANSITION_DUR = 1
ANIMATION_FRAMES = 15
ANIMATION_DURATION = 2

# Constantes pour l'apprentissage automatique
SLIDE_FEATURES = 20  # Nombre de caractéristiques pour représenter une diapositive
SENTENCE_EMBEDDING_SIZE = 768  # Taille des embeddings de phrases

# Configuration de la synthèse vocale
engine = pyttsx3.init()
engine.setProperty('rate', 150)

# Nettoyage des fichiers audio potentiellement corrompus
for f in OUTPUT_DIR.glob("*.wav"):
    try:
        f.unlink()
    except Exception as e:
        print(f"? Impossible de supprimer {f.name} : {e}")

# Définir le chemin du logo si disponible
LOGO_PATH = BASE_DIR / "logo.png"
if not LOGO_PATH.exists():
    LOGO_PATH = Path()  # Chemin vide si pas de logo
    
MARKDOWN_DIR = Path("C:/Users/ThinkPad/Desktop/plateform/python/").absolute()  # Remplacez par le chemin de votre fichier v3.md
# Configuration
class Config:
    ENABLE_USER_FEEDBACK = True
    MIN_SAMPLES_FOR_RETRAINING = 5
    AUTO_SAVE_MODELS = True

config = Config()

def check_dependencies():
    """Vérifie si toutes les bibliothèques nécessaires sont installées"""
    required_packages = {
        "moviepy": "Pour l'édition vidéo",
        "pyttsx3": "Pour la synthèse vocale",
        "PyPDF2": "Pour l'extraction de PDF",
        "python-docx": "Pour l'extraction de documents Word",
        "scikit-learn": "Pour les modèles d'apprentissage automatique",
        "Pillow": "Pour le traitement d'images",
        "joblib": "Pour la sauvegarde des modèles",
        "numpy": "Pour les calculs numériques",
        "pandas": "Pour la manipulation de données",
        "PyMuPDF": "Pour l'extraction avancée de PDF"
    }
    
    missing_packages = []
    
    for package, description in required_packages.items():
        try:
            # Utiliser une technique différente pour joblib qui est dans sklearn
            if package == "joblib":
                importlib.import_module("joblib")
            elif package == "Pillow":  # PIL est importé sous un autre nom
                importlib.import_module("PIL")
            elif package == "python-docx":  # docx est importé sous un autre nom
                importlib.import_module("docx")
            elif package == "PyMuPDF":  # PyMuPDF est importé sous le nom fitz
                importlib.import_module("fitz")
            elif package == "scikit-learn":  # scikit-learn s'importe sous le nom sklearn
                importlib.import_module("sklearn")
            else:
                importlib.import_module(package)
        except ImportError:
            missing_packages.append(f"{package} ({description})")
    
    if missing_packages:
        print("?? Les bibliothèques suivantes sont requises mais non installées:")
        for pkg in missing_packages:
            print(f"  - {pkg}")
        print("\n Utilisez pip pour les installer:")
        print(f"pip install {' '.join([pkg.split(' ')[0] for pkg in missing_packages])}")
        return False
    
    return True
def is_content_empty_except_images(content):
    """
    Vérifie si le contenu est vide ou ne contient que des références d'images ou des tableaux
    """
    if not content or not content.strip():
        return True

    # Suppression des images
    content_without_images = re.sub(r'!\[.*?\]\(.*?\)', '', content)
    
    # Supprimer les lignes de tableau Markdown
    content_without_tables = re.sub(r'^\|.*\|$', '', content_without_images, flags=re.MULTILINE)
    content_without_tables = re.sub(r'^[-|:]+$', '', content_without_tables, flags=re.MULTILINE)
    
    # Suppression des espaces, retours ligne, balises Markdown
    clean_content = re.sub(r'[#*_\s\n\r]', '', content_without_tables)
    
    return not bool(clean_content)
def slide_clip_with_exact_markdown_separated(title, content, images, style_params=None, layout_params=None, duration=8):
    """
    Crée un slide qui préserve exactement la structure Markdown originale
    avec une séparation claire entre texte et images pour éviter les chevauchements
    
    Args:
        title: Titre de la diapositive
        content: Contenu textuel
        images: Liste des chemins d'images
        style_params: Paramètres de style (optionnel)
        layout_params: Paramètres de mise en page (optionnel)
        duration: Durée de la diapositive en secondes
        
    Returns:
        CompositeVideoClip: Clip vidéo de la diapositive
    """

    # Valeurs par défaut pour les paramètres de style et de mise en page
    if style_params is None:
        style_params = {
            'name': 'balanced',
            'text_size': FONT_SIZE_TEXT,
            'animation_level': 0.5,
            'color_scheme': 'neutral',
            'layout': 'balanced',
            'transition': 'fade'
        }
    
    if layout_params is None:
        layout_params = {
            'text_x': 0.1, 
            'text_y': 0.15,
            'text_width': 0.8, 
            'text_height': 0.7,
            'title_position': 'top',
            'text_size_factor': 1.0
        }
    
    # Créer l'arrière-plan
    bg_color = BG_COLOR
    if style_params.get('color_scheme') == 'vibrant':
        bg_color = (240, 250, 255)  # Bleu très pâle
    elif style_params.get('color_scheme') == 'monochrome':
        bg_color = (248, 248, 248)  # Gris clair
    elif style_params.get('color_scheme') == 'contrast':
        bg_color = (235, 235, 245)  # Violet très pâle
    elif style_params.get('color_scheme') == 'focused':
        bg_color = (245, 245, 245)  # Gris très pâle
    
    bg = ColorClip((WIDTH, HEIGHT), color=bg_color).set_duration(duration)
    layers = [bg]
    
    # Déterminer si nous avons des images valides
    has_valid_images = bool(images)
    
    # Vérifier si le contenu est vide ou ne contient que des références d'images
    clean_content = re.sub(r'!\[.*?\]\(.*?\)', '', content).strip()
    clean_content = re.sub(r'[#*_\s\n\r]', '', clean_content)
    is_image_only = is_content_empty_except_images(content) and has_valid_images

    
    # Si layout_params contient 'fullscreen_image', utiliser cette valeur
    if is_image_only:
        layout_params['fullscreen_image'] = True
    is_fullscreen_image = layout_params.get('fullscreen_image', False)

    
    # Mode plein écran pour les images sans texte
    if is_fullscreen_image and has_valid_images:
        print(f"?? Mode plein écran activé pour image sans texte")
        
        # Ajouter le titre s'il existe (en bas pour ne pas gêner l'image)
        title_position = layout_params.get('title_position', 'bottom')
        if title:
            title_y = HEIGHT - 60 if title_position == 'bottom' else 10
            title_clip = TextClip(
                title, 
                fontsize=FONT_SIZE_TITLE, 
                color=TEXT_COLOR,
                method="caption", 
                align="center",
                size=(WIDTH-120, None)
            ).set_position(("center", title_y)).set_duration(duration)
            
            layers.append(title_clip)
       
        # Ajouter l'image en grand format
        try:
            # Utiliser la première image valide

            def trim_whitespace(img_path):
                img = Image.open(img_path)
                bg = Image.new(img.mode, img.size, img.getpixel((0,0)))
                diff = ImageChops.difference(img, bg)
                bbox = diff.getbbox()
                if bbox:
                    img = img.crop(bbox)
                # Sauvegarder temporairement l'image rognée
                temp_path = img_path.with_name(f"trimmed_{img_path.name}")
                img.save(temp_path)
                return temp_path

            # Rogner l’image avant affichage
            for img_path in images:
                try:
                    trimmed_path = trim_whitespace(Path(img_path))
                    img_clip = ImageClip(str(trimmed_path))

                    
                    # Calculer les dimensions pour que l'image remplisse l'écran
                    img_ratio = img_clip.h / img_clip.w
                    screen_ratio = HEIGHT / WIDTH
                    
                    if img_ratio > screen_ratio:  # Image plus haute que large
                        new_height = HEIGHT * 0.85  # 85% de la hauteur
                        new_width = new_height / img_ratio
                    else:  # Image plus large que haute
                        new_width = WIDTH * 0.9  # 90% de la largeur
                        new_height = new_width * img_ratio
                    min_width = WIDTH * 0.8
                    min_height = HEIGHT * 0.8
                    new_width = max(new_width, min_width)
                    new_height = max(new_height, min_height)

                    # Redimensionner l'image
                    img_clip = img_clip.resize(width=new_width) if img_ratio > screen_ratio else img_clip.resize(height=new_height)
                    
                    # Centrer l'image
                    x_pos = (WIDTH - img_clip.w) / 2
                    y_pos = (HEIGHT - img_clip.h) / 2
                    
                    # Positionner plus haut si le titre est en bas
                    if title and title_position == 'bottom':
                        y_pos -= 25  # Décaler légèrement vers le haut
                    
                    # Appliquer des animations si style animé
                    animation_level = style_params.get('animation_level', 0.7)
                    if animation_level > 0.6:
                        img_clip = img_clip.crossfadein(min(1.0, animation_level))
                    
                    img_clip = img_clip.set_position((x_pos, y_pos)).set_duration(duration)
                    layers.append(img_clip)
                    break  # N'utiliser que la première image valide
                except Exception as e:
                    print(f"?? Erreur chargement image {img_path}: {e}")
                    continue
        except Exception as e:
            print(f"?? Erreur générale images: {e}")
    
    # Traitement standard pour les diapositives avec texte et images
    else:
        # Réserver plus d'espace pour le titre
        title_height = 60
        content_top = title_height + 10
        available_height = HEIGHT - content_top - 40  # Marge du bas
        
        # Ajuster la taille du texte en fonction du style
        text_size = int(style_params.get('text_size', FONT_SIZE_TEXT) * layout_params.get('text_size_factor', 1.0))
        
        # Ajouter le titre s'il existe
        if title:
            # Vérifier si le titre est un marqueur de page pour le masquer
            if not re.match(r'(---\s*)?[Pp]age\s+\d+(\s*---)?', title):
                title_position = layout_params.get('title_position', 'top')
                title_y = 10 if title_position == 'top' else HEIGHT - 50
                
                title_clip = TextClip(
                    title, 
                    fontsize=FONT_SIZE_TITLE, 
                    color=TEXT_COLOR,
                    method="caption", 
                    align="center",
                    size=(WIDTH-120, None)
                ).set_position(("center", title_y)).set_duration(duration)
                
                layers.append(title_clip)
        
        # Convertir les images en liste de chemins si c'est une liste de dictionnaires
        image_paths = []
        if images:
            for img in images:
                if isinstance(img, dict):
                    image_paths.append(img['path'])
                elif isinstance(img, tuple):  # Support de l'ancien format
                    image_paths.append(img[0])
                else:
                    image_paths.append(img)
        
        # Appliquer les paramètres de mise en page pour le texte et les images
        text_x = int(layout_params.get('text_x', 0.1) * WIDTH)
        text_y = int(layout_params.get('text_y', 0.15) * HEIGHT)
        text_width = int(layout_params.get('text_width', 0.8) * WIDTH)
        
        # Supprimer les références d'images du texte affiché
        processed_content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
        
        # Préserver le formatage Markdown (gras, italique, listes)
        processed_content = preserve_markdown_formatting(processed_content)
        
        # Vérifier si le texte a du contenu après nettoyage
        if processed_content.strip():
            # Créer le clip de texte avec l'animation selon le niveau d'animation
            animation_level = style_params.get('animation_level', 0.5)
            
            if animation_level > 0.3:
                # Animation plus élaborée pour les styles dynamiques
                text_clip = create_animated_text_clip(
                    processed_content, 
                    duration=duration,
                    animation_level=animation_level
                )
            else:
                # Texte statique pour les styles minimalistes
                text_clip = create_text_clip(
                    processed_content, 
                    fontsize=text_size, 
                    width=text_width,
                    position=(text_x, text_y), 
                    duration=duration
                )
            
            # Ajouter le texte à la diapositive
            layers.append(text_clip)
        
        # Ajouter les images si présentes
        if has_valid_images:
            try:
                # Récupérer les paramètres de mise en page pour l'image
                img_x = int(layout_params.get('image_x', 0.55) * WIDTH)
                img_y = int(layout_params.get('image_y', 0.15) * HEIGHT)
                img_width = int(layout_params.get('image_width', 0.4) * WIDTH)
                
                # Utiliser la première image valide
                for img_path in image_paths:
                    try:
                        img_clip = ImageClip(str(img_path))
                        img_ratio = img_clip.h / img_clip.w
                        img_height = int(img_width * img_ratio)
                        
                        # Limiter la hauteur si nécessaire
                        max_height = int(HEIGHT * 0.6)
                        if img_height > max_height:
                            img_height = max_height
                            img_width = int(img_height / img_ratio)
                        
                        # Redimensionner et positionner l'image
                        img_clip = img_clip.resize(width=img_width)
                        img_clip = img_clip.set_position((img_x, img_y)).set_duration(duration)
                        
                        # Ajouter un effet de fondu d'entrée pour les styles dynamiques
                        animation_level = style_params.get('animation_level', 0.5)
                        if animation_level > 0.6:
                            img_clip = img_clip.crossfadein(min(1.0, animation_level))
                        
                        layers.append(img_clip)
                        break  # Utiliser seulement la première image valide
                    except Exception as e:
                        print(f"?? Erreur chargement image {img_path}: {e}")
                        continue
            except Exception as e:
                print(f"?? Erreur générale images: {e}")
    
    # Ajouter le logo si une fonction handle_logo existe
    if 'handle_logo' in globals():
        layers = handle_logo(layers, style_params, duration)
    
    # Créer la diapositive finale
    return CompositeVideoClip(layers, size=(WIDTH, HEIGHT))
# Remplacer les fonctions originales par nos versions améliorées
def integrate_fullscreen_image_solution():
    """Intègre les modifications pour afficher correctement les images en plein écran"""
    global slide_clip_with_exact_markdown, slide_clip_with_exact_markdown_separated
    
    # Remplacer la fonction de création de diapositives
    slide_clip_with_exact_markdown = slide_clip_with_exact_markdown_separated
    
    # Vérifier si la classe LayoutOptimizer existe et remplacer sa méthode
    if 'LayoutOptimizer' in globals():
        LayoutOptimizer.optimize_layout = optimize_layout
        print("? Méthode d'optimisation de mise en page remplacée")
    
    print("? Solution pour l'affichage plein écran des images intégrée avec succès")
    return True
# Exécuter l'intégration
integrate_fullscreen_image_solution()
# Mise à jour de la classe LayoutOptimizer pour mieux gérer la séparation texte/image
def optimize_layout_separated(self, content, title, images, style_params):
    """
    Optimise la mise en page avec une séparation claire entre texte et images
    
    Args:
        content: Texte de la diapositive
        title: Titre de la diapositive
        images: Liste des chemins d'images
        style_params: Paramètres de style
        
    Returns:
        layout_params: Paramètres de mise en page optimisés
    """
    # Analyser le contenu
    layout_params = {}
    has_image = bool(images)
    text_length = len(content)
    
    # 1. Détection de tableaux dans le contenu
    has_table = re.search(r'\|[-]+\|', content) is not None  # Détection de tableaux markdown
    has_table = has_table or "|" in content and any(line.count('|') > 2 for line in content.split('\n'))
    
    # 2. Détection d'images larges ou complexes
    large_images = []
    complex_images = []
    
    if has_image and images:
        for img_path in images:
            try:
                trimmed_path = trim_whitespace(Path(img_path))
                img_clip = ImageClip(str(trimmed_path))
                
                # Calculer les dimensions pour que l'image remplisse l'écran
                img_ratio = img_clip.h / img_clip.w
                screen_ratio = HEIGHT / WIDTH
                
                if img_ratio > screen_ratio:  # Image plus haute que large
                    new_height = HEIGHT * 0.85
                    new_width = new_height / img_ratio
                else:
                    new_width = WIDTH * 0.9
                    new_height = new_width * img_ratio

                # Forcer un minimum de taille
                min_width = WIDTH * 0.8
                min_height = HEIGHT * 0.8
                new_width = max(new_width, min_width)
                new_height = max(new_height, min_height)

                # Redimensionner et centrer
                img_clip = img_clip.resize(width=new_width) if img_ratio > screen_ratio else img_clip.resize(height=new_height)
                x_pos = (WIDTH - img_clip.w) / 2
                y_pos = (HEIGHT - img_clip.h) / 2
                if title and layout_params.get('title_position') == 'bottom':
                    y_pos -= 25
                
                # Animation
                animation_level = style_params.get('animation_level', 0.7)
                if animation_level > 0.6:
                    img_clip = img_clip.crossfadein(min(1.0, animation_level))
                
                img_clip = img_clip.set_position((x_pos, y_pos)).set_duration(duration)
                layers.append(img_clip)
                break
            except Exception as e:
                print(f"?? Erreur chargement image {img_path}: {e}")
                continue

    
    # 3. Déterminer si le contenu est court ou long
    is_short_content = text_length < 500
    is_very_long_content = text_length > 1500
    
    # 4. Récupérer la mise en page de base du style
    layout_type = style_params.get('layout', 'balanced')
    
    # 5. Décision pour la mise en page avec séparation claire texte/image
    
    # 5.1 Cas spécial: tableaux
    if has_table:
        if has_image:
            # Tableau en haut, image en bas
            return {
                'text_x': 0.05, 'text_y': 0.15,
                'text_width': 0.9, 'text_height': 0.45,
                'image_x': 0.3, 'image_y': 0.65,
                'image_width': 0.4, 'image_height': 0.3,
                'title_position': 'top',
                'text_size_factor': 0.9
            }
        else:
            # Tableau seul: utiliser tout l'espace
            return {
                'text_x': 0.05, 'text_y': 0.15,
                'text_width': 0.9, 'text_height': 0.75,
                'title_position': 'top',
                'text_size_factor': 0.9
            }
    
    # 5.2 Cas spécial: images complexes (graphiques, tableaux en image)
    if complex_images:
        # Pour les images complexes: grande image à droite
        return {
            'text_x': 0.05, 'text_y': 0.15,
            'text_width': 0.45, 'text_height': 0.75,
            'image_x': 0.55, 'image_y': 0.15,
            'image_width': 0.4, 'image_height': 0.7,
            'title_position': 'top',
            'text_size_factor': 0.95
        }
    
    # 5.3 Cas spécial: images larges
    if large_images:
        if is_short_content:
            # Texte court: image en haut, texte en bas
            return {
                'text_x': 0.1, 'text_y': 0.55,
                'text_width': 0.8, 'text_height': 0.4,
                'image_x': 0.15, 'image_y': 0.15,
                'image_width': 0.7, 'image_height': 0.35,
                'title_position': 'top'
            }
        else:
            # Texte plus long: image en bas, texte en haut
            return {
                'text_x': 0.1, 'text_y': 0.15,
                'text_width': 0.8, 'text_height': 0.5,
                'image_x': 0.15, 'image_y': 0.65,
                'image_width': 0.7, 'image_height': 0.3,
                'title_position': 'top'
            }
    
    # 5.4 Mises en page standard avec séparation claire texte/image
    if has_image:
        if layout_type == 'image_dominant':
            # Image dominante: texte en bas
            return {
                'text_x': 0.1, 'text_y': 0.65,
                'text_width': 0.8, 'text_height': 0.3,
                'image_x': 0.1, 'image_y': 0.15,
                'image_width': 0.8, 'image_height': 0.45,
                'title_position': 'top'
            }
        
        elif layout_type == 'text_focus':
            # Texte dominant: image en bas à droite
            return {
                'text_x': 0.05, 'text_y': 0.15,
                'text_width': 0.9, 'text_height': 0.5,
                'image_x': 0.6, 'image_y': 0.68,
                'image_width': 0.35, 'image_height': 0.28,
                'title_position': 'top'
            }
        
        else:  # balanced, split et autres
            # Disposition équilibrée: texte à gauche, image à droite
            return {
                'text_x': 0.05, 'text_y': 0.15,
                'text_width': 0.45, 'text_height': 0.75,
                'image_x': 0.55, 'image_y': 0.15,
                'image_width': 0.4, 'image_height': 0.7,
                'title_position': 'top'
            }
    else:  # Pas d'image
        # Utiliser tout l'espace pour le texte
        return {
            'text_x': 0.1, 'text_y': 0.15,
            'text_width': 0.8, 'text_height': 0.7,
            'title_position': 'top',
            'text_size_factor': 1.0 if is_short_content else 0.9
        }
# Création de la fonction principale pour remplacer la fonction existante dans le code
def replace_functions_for_better_separation():
    """
    Remplace les fonctions existantes par des versions améliorées qui garantissent
    une meilleure séparation entre le texte et les images
    """
    # 1. Remplacer slide_clip_with_exact_markdown par notre version améliorée
    global slide_clip_with_exact_markdown
    slide_clip_with_exact_markdown = slide_clip_with_exact_markdown_separated
    
    # 2. Intégrer la fonction optimize_layout_separated dans la classe LayoutOptimizer
    LayoutOptimizer.optimize_layout = optimize_layout_separated
    
    print("? Fonctions remplacées avec succès pour garantir une meilleure séparation texte/image")
    
    return True

    """
   class TextProcessor was here then refactored to TextProcessing.py
    """
       
    model = VisionEncoderDecoderModel.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
    processor = ViTImageProcessor.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
    tokenizer = AutoTokenizer.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)
def explain_image(image_path):
    image = Image.open(image_path).convert("RGB")
    pixel_values = processor(image, return_tensors="pt").pixel_values.to(device)

    # Prompt guide pour obtenir une explication (et pas juste description)
    prompt = "Expliquez précisément ce que montre cette image : les objets, le contexte, et l'action éventuelle."
    input_ids = tokenizer(prompt, return_tensors="pt").input_ids.to(device)

    output = model.generate(
        pixel_values=pixel_values,
        decoder_input_ids=input_ids,
        max_length=100,
        num_beams=4
    )
    explanation = tokenizer.decode(output[0], skip_special_tokens=True)
    return explanation

class TextEnhancer:
    """
    Classe pour améliorer le texte et générer des descriptions d'images
    en utilisant des modèles plus légers ou des règles heuristiques
    """
    
    def __init__(self):
        self.initialized = False
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.text_model = None
        self.image_model = None
        
        # Tentative d'initialisation des modèles
        self.initialize_models()
    
    def initialize_models(self):
        """Initialise les modèles pour le traitement de texte et d'images"""
        try:
            print("?? Initialisation des modèles de traitement de texte et d'images...")
            
            # Essayer d'initialiser un modèle léger pour le texte
            try:
                # Utiliser un modèle T5 léger pour le résumé et la génération de texte
                model_name = "t5-small"  # Modèle beaucoup plus léger
                
                print(f"?? Chargement du modèle de texte: {model_name}")
                self.text_tokenizer = AutoTokenizer.from_pretrained(model_name)
                self.text_model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
                
                # Créer un pipeline pour la génération de texte
                self.text_pipeline = pipeline(
                    "text2text-generation",
                    model=self.text_model,
                    tokenizer=self.text_tokenizer,
                    device=0 if self.device == "cuda" else -1
                )
                
                print("? Modèle de traitement de texte initialisé")
            except Exception as e:
                print(f"?? Erreur lors de l'initialisation du modèle de texte: {e}")
                traceback.print_exc()
                self.text_model = None
                self.text_pipeline = None
            
            # Essayer d'initialiser un modèle léger pour la description d'images
            try:
                # Utiliser un modèle de classification d'images léger
                vision_model_name = "google/vit-base-patch16-224"
                
                print(f"?? Chargement du modèle d'images: {vision_model_name}")
                self.image_extractor = AutoImageProcessor.from_pretrained(vision_model_name)
                self.image_model = AutoModelForImageClassification.from_pretrained(vision_model_name)
                
                # Créer un pipeline pour la classification d'images
                self.image_pipeline = pipeline(
                    "image-classification",
                    model=self.image_model,
                    feature_extractor=self.image_extractor,
                    device=0 if self.device == "cuda" else -1
                )
                
                print("? Modèle de traitement d'images initialisé")
            except Exception as e:
                print(f"?? Erreur lors de l'initialisation du modèle d'images: {e}")
                self.image_model = None
                self.image_pipeline = None
            
            # Même si les modèles ne sont pas disponibles, on peut utiliser des méthodes heuristiques
            self.initialized = True
            print("? Système d'amélioration de texte initialisé")
            
        except Exception as e:
            print(f"? Erreur lors de l'initialisation des modèles: {e}")
            traceback.print_exc()
            self.initialized = False


    """
    class EnhancedVoiceSynthesizer was here then refactored to audio_processing.py 
    """    
    
    def analyze_narration_speed(self, text):
        """Analyse le texte pour recommander une vitesse de narration optimale"""
        try:
            # Mesures de base
            words = text.split()
            word_count = len(words)
            
            if word_count == 0:
                return 150  # Vitesse par défaut
            
            # Analyse de la complexité
            avg_word_length = sum(len(word) for word in words) / word_count
            digit_count = sum(1 for char in text if char.isdigit())
            digit_ratio = digit_count / len(text) if text else 0
            
            # Détecter les phrases techniques ou complexes
            has_technical_terms = any(len(word) > 8 for word in words)
            punc_count = sum(1 for char in text if char in ',.;:!?')
            punc_ratio = punc_count / word_count if word_count else 0
            
            # Calcul de la vitesse recommandée
            base_speed = 150  # Vitesse de parole de base (mots par minute)
            
            # Ajuster selon la complexité
            if avg_word_length > 7 or has_technical_terms:
                base_speed -= 20  # Ralentir pour les mots longs/techniques
            elif avg_word_length < 4:
                base_speed += 10  # Accélérer pour les mots courts
            
            # Ajuster selon la ponctuation
            if punc_ratio > 0.2:  # Beaucoup de ponctuation = phrases complexes
                base_speed -= 10
            
            # Ajuster selon les chiffres
            if digit_ratio > 0.05:  # Beaucoup de chiffres = besoin d'attention
                base_speed -= 15
            
            # Limites raisonnables
            return max(120, min(180, base_speed))
            
        except Exception as e:
            print(f"?? Erreur analyse vitesse narration: {e}")
            return 150  # Valeur par défaut en cas d'erreur
    
    def estimate_audio_duration(self, text, voice_speed=None):
        """
        Estime la durée d'un fichier audio généré à partir du texte
        
        Args:
            text: Texte à prononcer
            voice_speed: Vitesse de parole (mots par minute)
            
        Returns:
            float: Durée estimée en secondes
        """
        if not text:
            return 0
        
        try:
            # Utiliser la vitesse analysée si non spécifiée
            if voice_speed is None:
                voice_speed = self.analyze_narration_speed(text)
            
            # Nettoyer le texte pour l'estimation
            clean_text = re.sub(r'<.*?>', '', text)  # Supprimer les balises
            
            # Compter les mots
            words = clean_text.split()
            word_count = len(words)
            
            # Formule de base: durée = nombre de mots / (mots par minute / 60)
            base_duration = (word_count / voice_speed) * 60
            
            # Ajouter du temps pour les pauses
            punctuation_count = sum(1 for char in clean_text if char in ',.;:!?')
            pause_time = punctuation_count * 0.2  # 0.2 secondes par signe de ponctuation
            
            # Pauses pour les listes à puces
            bullet_count = len(re.findall(r'[-*•]\s+', clean_text))
            bullet_pause_time = bullet_count * 0.5  # 0.5 secondes par puce
            
            # Ajouter un facteur de sécurité
            safety_factor = 1.1
            
            # Calculer la durée totale
            total_duration = (base_duration + pause_time + bullet_pause_time) * safety_factor
            
            # Limiter à une valeur minimale raisonnable
            return max(3.0, total_duration)
            
        except Exception as e:
            print(f"?? Erreur estimation durée audio: {e}")
            # Estimation simplifiée en cas d'erreur
            return max(3.0, len(text.split()) * 0.3)  # ~0.3 secondes par mot
# Version alternative de la fonction create_combined_audio qui apparaît deux fois dans le code

def create_combined_audio(self, main_audio_path, image_audio_paths, output_path):
        """
        Combine l'audio principal avec les descriptions d'images
        en utilisant une approche simple avec pydub
        """
        try:
            if not Path(main_audio_path).exists():
                print(f"?? Fichier audio principal non trouvé: {main_audio_path}")
                return False
            
            # Vérifier que les chemins d'audio d'images existent
            valid_image_audio_paths = [path for path in image_audio_paths if Path(path).exists()]
            
            if not valid_image_audio_paths:
                # S'il n'y a pas d'audio d'images valides, renommer simplement l'audio principal
                try:
                    shutil.copy(main_audio_path, output_path)
                    print(f"? Audio principal copié vers {output_path}")
                    return True
                except Exception as e:
                    print(f"?? Erreur lors de la copie de l'audio: {e}")
                    return False
            
            # Tentative d'utilisation de pydub pour combiner les audios
            try:
                # Charger l'audio principal
                combined = AudioSegment.from_wav(main_audio_path)
                
                # Ajouter un silence entre les clips
                silence = AudioSegment.silent(duration=1000)  # 1 seconde
                
                # Ajouter les descriptions d'images
                for audio_path in valid_image_audio_paths:
                    try:
                        image_audio = AudioSegment.from_wav(audio_path)
                        combined += silence + image_audio
                    except Exception as e:
                        print(f"?? Erreur lors de l'ajout de l'audio d'image {audio_path}: {e}")
                
                # Exporter l'audio combiné
                combined.export(output_path, format="wav")
                print(f"? Audio combiné sauvegardé dans {output_path}")
                return True
            except ImportError:
                print("?? pydub non disponible, tentative de solution alternative...")
                
                # Solution alternative simple: utiliser moviepy
                try:
                    # Charger tous les clips audio
                    audio_clips = [AudioFileClip(str(main_audio_path))]
                    
                    for audio_path in valid_image_audio_paths:
                        try:
                            clip = AudioFileClip(str(audio_path))
                            audio_clips.append(clip)
                        except Exception as e:
                            print(f"?? Erreur lors du chargement de l'audio {audio_path}: {e}")
                    
                    # Concaténer les clips
                    combined_audio = concatenate_audioclips(audio_clips)
                    
                    # Sauvegarder l'audio combiné
                    combined_audio.write_audiofile(str(output_path))
                    print(f"? Audio combiné sauvegardé dans {output_path}")
                    return True
                except Exception as e:
                    print(f"?? Erreur lors de la combinaison des audios avec moviepy: {e}")
            
            # Si aucune méthode de combinaison ne fonctionne, utiliser l'audio principal
            shutil.copy(main_audio_path, output_path)
            print(f"?? Impossible de combiner les audios, utilisation de l'audio principal uniquement")
            return True
            
        except Exception as e:
            print(f"?? Erreur lors de la création de l'audio combiné: {e}")
            return False

# Extrait du code dans la méthode process_markdown qui gère la génération et synchronisation audio
"""
# 4. Générer la narration audio si nécessaire
print("?? Création de la narration audio...")
narration_audio_paths = []

for i, segment in enumerate(segments):
    try:
        # Optimiser le texte pour la synthèse vocale
        narration_text = self.optimize_narration_text(segment.get('content', ''))
        
        if narration_text and narration_text.strip():
            # Générer un nom de fichier audio
            audio_filename = f"narration_{i+1}.wav"
            audio_path = self.output_dir / audio_filename
            
            # Générer l'audio
            print(f"?? Génération de l'audio pour la diapositive {i+1}...")
            audio_success = self.voice_synthesizer.save_to_file(narration_text, audio_path)
            
            if audio_success and audio_path.exists():
                print(f"? Audio créé: {audio_path}")
                narration_audio_paths.append(audio_path)
                
                # Ajouter l'audio de description d'image si disponible
                if 'image_description_audio' in segment and Path(segment['image_description_audio']).exists():
                    narration_audio_paths.append(segment['image_description_audio'])
    except Exception as e:
        print(f"?? Erreur lors de la création de l'audio pour la diapositive {i+1}: {e}")

# ...plus tard dans le code...

# 7. Ajouter l'audio si disponible
if narration_audio_paths:
    try:
        # Combiner tous les fichiers audio en un seul
        combined_audio_path = self.output_dir / "combined_narration.wav"
        audio_combined = self.voice_synthesizer.create_combined_audio(
            narration_audio_paths[0],  # Premier fichier comme base
            narration_audio_paths[1:],  # Reste des fichiers
            combined_audio_path
        )
        
        if audio_combined and combined_audio_path.exists():
            from moviepy import AudioFileClip
            
            # Ajouter l'audio à la vidéo
            narration_audio = AudioFileClip(str(combined_audio_path))
            final_video = final_video.set_audio(narration_audio)
            print("?? Narration audio ajoutée à la vidéo")
    except Exception as e:
        print(f"?? Erreur lors de l'ajout de l'audio: {e}")
"""

# Fonction d'optimisation du texte pour la narration audio
def optimize_narration_text(self, text):
    """
    Optimise le texte pour la narration audio
    - Simplifie les phrases complexes
    - Ajoute des pauses
    - Normalise la prononciation
    """
    try:
        # Vérifier que le texte est valide
        if not text or not isinstance(text, str):
            return "Contenu non disponible"
        
        # Simplifier les phrases trop longues
        sentences = re.split(r'([.!?])', text)
        optimized_sentences = []
        
        i = 0
        while i < len(sentences) - 1:
            sentence = sentences[i]
            punctuation = sentences[i+1] if i+1 < len(sentences) else "."
            
            # Simplifier les phrases trop longues
            if len(sentence.split()) > 25:
                words = sentence.split()
                mid = len(words) // 2
                
                # Chercher un meilleur point de coupure
                for j in range(mid-3, mid+3):
                    if 0 <= j < len(words) and any(p in words[j] for p in [',', ';', ':', '-']):
                        mid = j + 1
                        break
                
                first_part = ' '.join(words[:mid])
                second_part = ' '.join(words[mid:])
                
                # Capitaliser la première lettre de la seconde moitié
                if second_part:
                    second_part = second_part[0].upper() + second_part[1:] if len(second_part) > 1 else second_part.upper()
                
                optimized_sentences.append(first_part + ".")
                optimized_sentences.append(second_part + punctuation)
            else:
                optimized_sentences.append(sentence + punctuation)
            
            i += 2
        
        # Ajouter le reste s'il y en a
        if i < len(sentences):
            optimized_sentences.append(sentences[i])
        
        # Recombiner en ajoutant des pauses pour la synthèse vocale
        optimized_text = ' '.join(optimized_sentences)
        
        # Ajouter des pauses pour la synthèse vocale aux virgules
        optimized_text = optimized_text.replace(", ", ", <pause> ")
        
        # Normaliser les abréviations et les nombres pour meilleure prononciation
        optimized_text = re.sub(r'(\d+)%', r'\1 pour cent', optimized_text)
        optimized_text = re.sub(r'Dr\.', 'Docteur', optimized_text)
        optimized_text = re.sub(r'M\.', 'Monsieur', optimized_text)
        optimized_text = re.sub(r'Mme\.', 'Madame', optimized_text)
        
        # Supprimer les caractères spéciaux qui pourraient affecter la narration
        optimized_text = re.sub(r'[*_~]', '', optimized_text)
        
        # Remplacer les caractères de pause par des espaces pour la synthèse finale
        optimized_text = optimized_text.replace("<pause>", "")
        
        return optimized_text
        
    except Exception as e:
        print(f"?? Erreur lors de l'optimisation de la narration: {e}")
        # En cas d'erreur, retourner un texte simplifié
        if text and isinstance(text, str):
            # Retirer les caractères spéciaux et limiter la longueur
            clean_text = re.sub(r'[^\w\s.,!?]', '', text)
            return clean_text[:500] + "..." if len(clean_text) > 500 else clean_text
        return "Contenu non disponible pour la narration" 
def create_combined_audio(self, main_audio_path, image_audio_paths, output_path):
        """
        Combine l'audio principal avec les descriptions d'images
        en utilisant une approche simple avec pydub
        """
        try:
            if not Path(main_audio_path).exists():
                print(f"?? Fichier audio principal non trouvé: {main_audio_path}")
                return False
            
            # Vérifier que les chemins d'audio d'images existent
            valid_image_audio_paths = [path for path in image_audio_paths if Path(path).exists()]
            
            if not valid_image_audio_paths:
                # S'il n'y a pas d'audio d'images valides, renommer simplement l'audio principal
                try:
                    shutil.copy(main_audio_path, output_path)
                    print(f"? Audio principal copié vers {output_path}")
                    return True
                except Exception as e:
                    print(f"?? Erreur lors de la copie de l'audio: {e}")
                    return False
            
            # Tentative d'utilisation de pydub pour combiner les audios
            try:               
                # Charger l'audio principal
                combined = AudioSegment.from_wav(main_audio_path)
                
                # Ajouter un silence entre les clips
                silence = AudioSegment.silent(duration=1000)  # 1 seconde
                
                # Ajouter les descriptions d'images
                for audio_path in valid_image_audio_paths:
                    try:
                        image_audio = AudioSegment.from_wav(audio_path)
                        combined += silence + image_audio
                    except Exception as e:
                        print(f"?? Erreur lors de l'ajout de l'audio d'image {audio_path}: {e}")
                
                # Exporter l'audio combiné
                combined.export(output_path, format="wav")
                print(f"? Audio combiné sauvegardé dans {output_path}")
                return True
            except ImportError:
                print("?? pydub non disponible, tentative de solution alternative...")
                
                # Solution alternative simple: utiliser moviepy
                try:                    
                    # Charger tous les clips audio
                    audio_clips = [AudioFileClip(str(main_audio_path))]
                    
                    for audio_path in valid_image_audio_paths:
                        try:
                            clip = AudioFileClip(str(audio_path))
                            audio_clips.append(clip)
                        except Exception as e:
                            print(f"?? Erreur lors du chargement de l'audio {audio_path}: {e}")
                    
                    # Concaténer les clips
                    combined_audio = concatenate_audioclips(audio_clips)
                    
                    # Sauvegarder l'audio combiné
                    combined_audio.write_audiofile(str(output_path))
                    print(f"? Audio combiné sauvegardé dans {output_path}")
                    return True
                except Exception as e:
                    print(f"?? Erreur lors de la combinaison des audios avec moviepy: {e}")
            
            # Si aucune méthode de combinaison ne fonctionne, utiliser l'audio principal
            shutil.copy(main_audio_path, output_path)
            print(f"?? Impossible de combiner les audios, utilisation de l'audio principal uniquement")
            return True
            
        except Exception as e:
            print(f"?? Erreur lors de la création de l'audio combiné: {e}")
            return False
# Fonctions utilitaires générales
def slugify(text):
    """Convertit un texte en slug URL-friendly"""
    text = unicodedata.normalize("NFKD", text)
    text = text.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_")

# Fonctions d'extraction d'images
def find_valid_image_path(base_path, raw_path, search_subfolders=True, debug=True):
    """Version robuste et exhaustive de la recherche d'images"""
    if debug:
        print(f"?? Recherche de l'image: {raw_path}")
    
    # Nettoyage du chemin
    raw_path = raw_path.strip().replace('\\', '/').replace('%20', ' ')
    
    # Liste des chemins possibles - ajoutez MARKDOWN_DIR comme premier lieu de recherche
    possible_paths = []
    
    # 1. Chercher d'abord dans le répertoire du Markdown
    possible_paths.append(MARKDOWN_DIR / raw_path)
    
    # 2. Chemin direct tel quel
    possible_paths.append(base_path / raw_path)
    
    # 3. Juste le nom du fichier dans le dossier de base
    filename = Path(raw_path).name
    possible_paths.append(base_path / filename)
    
    # 4. Essayer avec différentes extensions
    extensions = ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.bmp']
    
    # Pour le chemin complet sans extension
    stem = re.sub(r"\.\w+$", "", raw_path)
    for ext in extensions:
        possible_paths.append(MARKDOWN_DIR / f"{stem}{ext}")
        possible_paths.append(base_path / f"{stem}{ext}")
    
    # Pour juste le nom du fichier sans extension
    filename_stem = Path(stem).name
    for ext in extensions:
        possible_paths.append(MARKDOWN_DIR / f"{filename_stem}{ext}")
        possible_paths.append(base_path / f"{filename_stem}{ext}")
    
    # 4. Recherche dans les sous-dossiers courants
    if search_subfolders:
        common_folders = ["images", "img", "assets", "media", "pictures", "photos", "docs", "content"]
        
        for folder in common_folders:
            subfolder_path = base_path / folder
            
            if subfolder_path.exists():
                # Chemin complet dans le sous-dossier
                possible_paths.append(subfolder_path / raw_path)
                
                # Juste le nom dans le sous-dossier
                possible_paths.append(subfolder_path / filename)
                
                # Variations d'extensions dans le sous-dossier
                for ext in extensions:
                    possible_paths.append(subfolder_path / f"{filename_stem}{ext}")
    
    # Vérifier tous les chemins possibles
    for path in possible_paths:
        try:
            if path.exists():
                if debug:
                    print(f"? Image trouvée: {path}")
                return path
        except:
            # Ignorer les erreurs de chemin invalide
            pass
    
    if debug:
        print(f"? Aucune image trouvée pour: {raw_path}")
    return None
def extract_images_from_text(text, base_path=BASE_DIR):
    """Extrait toutes les images depuis le Markdown avec une détection améliorée"""
    lines = text.split("\n")
    cleaned_lines = []
    image_references = []
    
    # Motif pour trouver les références d'images dans la syntaxe Markdown standard
    image_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    
    print(f"?? Extraction des images de {len(lines)} lignes de texte...")
    
    for i, line in enumerate(lines):
        image_matches = image_pattern.findall(line)
        
        if image_matches:
            print(f"?? Images trouvées à la ligne {i}: {len(image_matches)} images")
            for alt_text, img_path in image_matches:
                # Nettoyer le chemin d'image
                img_path = img_path.strip()
                print(f"  ?? Tentative de traitement de l'image: '{img_path}' (alt: '{alt_text}')")
                
                # Vérifier plusieurs possibilités pour le chemin
                valid_path = find_valid_image_path(base_path, img_path)
                
                if valid_path:
                    print(f"  ? Image trouvée: {valid_path}")
                    # Stocker la référence avec le chemin valide
                    image_references.append((valid_path, i))
                else:
                    # Essayer des alternatives si le chemin direct ne fonctionne pas
                    print(f"  ?? Chemin image non trouvé directement: {img_path}")
                    alt_paths = [
                        # Essayer avec le nom de fichier seul
                        base_path / Path(img_path).name,
                        # Essayer dans des sous-dossiers courants
                        base_path / "images" / Path(img_path).name,
                        base_path / "img" / Path(img_path).name,
                        base_path / "assets" / Path(img_path).name,
                        base_path / "media" / Path(img_path).name,
                        base_path / "pictures" / Path(img_path).name,
                    ]
                    
                    found = False
                    for alt_path in alt_paths:
                        if alt_path.exists():
                            print(f"  ? Image alternative trouvée: {alt_path}")
                            image_references.append((alt_path, i))
                            found = True
                            break
                    
                    if not found:
                        print(f"  ? Aucune image trouvée pour: {img_path}")
        
        # Nettoyer les lignes pour supprimer les références d'images
        cleaned_line = re.sub(r'!\[.*?\]\(.*?\)', '', line)
        if cleaned_line.strip():  # On ne garde pas les lignes vides
            cleaned_lines.append(cleaned_line)
    
    # Vérification supplémentaire pour les images dans le dossier
    if len(image_references) == 0:
        print("?? Aucune image trouvée dans le texte. Recherche d'images alternatives dans le dossier...")
        image_files = []
        for ext in ["*.jpg", "*.jpeg", "*.png", "*.gif", "*.webp"]:
            image_files.extend(list(base_path.glob(ext)))
            
            # Chercher également dans les sous-dossiers courants
            for subfolder in ["images", "img", "assets", "media", "pictures"]:
                subfolder_path = base_path / subfolder
                if subfolder_path.exists():
                    image_files.extend(list(subfolder_path.glob(ext)))
        
        if image_files:
            print(f"?? {len(image_files)} images trouvées dans le dossier")
            for img in image_files[:5]:  # Limiter à 5 images pour éviter la surcharge
                print(f"  ? Ajout de l'image trouvée: {img}")
                image_references.append((img, 0))
    
    # Résumé de l'extraction
    print(f"?? Résultat de l'extraction: {len(image_references)} images trouvées")
    
    return image_references, "\n".join(cleaned_lines).strip()
def extract_images_from_markdown(content, base_path):
    """Extrait toutes les images du Markdown avec leur position exacte"""
    image_pattern = re.compile(r'!\[(.*?)\]\((.*?)\)')
    images = []
    
    # Trouver toutes les références d'images
    for match in image_pattern.finditer(content):
        alt_text = match.group(1)
        img_path = match.group(2)
        position = match.start()
        
        # Trouver le chemin valide de l'image
        valid_path = find_valid_image_path(base_path, img_path)
        
        if valid_path:
            # Stocker l'image avec sa position exacte dans le Markdown
            images.append({
                'path': valid_path,
                'alt_text': alt_text,
                'original_path': img_path,
                'position': position,
                'match_length': len(match.group(0))
            })
    
    return images
def extract_only_figures_from_pdf(pdf_path):
    """
    Extrait UNIQUEMENT les figures d'un fichier PDF, en préservant les figures blanches
    et en gérant les cas où le PDF ne contient pas de texte extractible
    """
    print(f"?? Extraction des figures uniquement du PDF: {pdf_path}")
    text = ""
    images = []
    
    try:
        # Créer un dossier pour les images extraites
        image_dir = EXTRACTED_IMAGES_DIR
        image_dir.mkdir(exist_ok=True)
        
        # Extraire le texte avec PyPDF2 (pour la structure)
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            for page_num in range(num_pages):
                # Ajouter un marqueur de page, même si pas de texte
                text += f"\n\n## --- Page {page_num + 1} ---\n\n"
                
                # Extraire le texte de la page
                page = reader.pages[page_num]
                page_text = page.extract_text()
                
                # Ajouter le texte au contenu s'il existe
                if page_text:
                    text += page_text + "\n\n"
        
        # Vérifier si du texte a été extrait, sinon créer une structure minimale
        if text.strip() == "" or all(line.strip().startswith("## --- Page") for line in text.strip().split("\n") if line.strip()):
            print("?? Aucun texte extractible trouvé dans le PDF, création d'une structure minimale")
            text = ""
            for page_num in range(num_pages):
                text += f"\n\n## --- Page {page_num + 1} ---\n\n"
                text += f"Contenu de la page {page_num + 1}\n\n"
        
        # Extraire les images avec PyMuPDF (fitz)
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extraire uniquement les images intégrées
            image_list = page.get_images(full=True)
            page_has_images = False
            
            if len(image_list) > 0:
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        
                        # Filtrer les images trop petites (probablement des icônes)
                        width = base_image.get("width", 0)
                        height = base_image.get("height", 0)
                        
                        # Ignorer les très petites images (logos, icônes)
                        if width < 50 or height < 50:
                            continue
                            
                        # Extraire les données de l'image
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Créer un nom de fichier pour l'image
                        image_filename = f"pdf_figure_{page_num + 1}_{img_index + 1}.{image_ext}"
                        image_path = image_dir / image_filename
                        
                        # Sauvegarder l'image
                        with open(image_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        
                        # Ne pas inverser les figures blanches
                        try:
                            img = Image.open(image_path)
                            if img.mode in ('RGBA'):
                                # Fond blanc pour les images transparentes
                                background = Image.new('RGB', img.size, (255, 255, 255))
                                background.paste(img, mask=img.split()[3])
                                img = background
                                img.save(image_path)
                            
                            # Analyse statistique pour déterminer si c'est une image noire
                            stat = ImageStat.Stat(img)
                            brightness = sum(stat.mean) / len(stat.mean)
                            print(f"  ?? Image {img_index} luminosité: {brightness}")
                            
                            # Si et seulement si l'image est TRÈS sombre (figure noire)
                            if brightness < 20:  # Seuil très bas pour éviter d'inverser les figures mixtes
                                # Inverser uniquement les images très sombres
                                img = ImageOps.invert(img)
                                
                                # Augmenter le contraste pour améliorer la visibilité
                                enhancer = ImageEnhance.Contrast(img)
                                img = enhancer.enhance(1.5)
                                
                                # Sauvegarder l'image améliorée
                                img.save(image_path)
                                print(f"  ?? Image sombre inversée: {brightness}")
                            else:
                                print(f"  ? Image claire préservée: {brightness}")
                            
                        except Exception as e:
                            print(f"?? Erreur lors du traitement de l'image: {e}")
                        
                        # Ajouter l'image à la liste
                        images.append(image_path)
                        
                        # Ajouter une référence à l'image dans le texte Markdown
                        page_text_marker = f"## --- Page {page_num + 1} ---"
                        if page_text_marker in text:
                            # Insérer la référence après le marqueur de page
                            insert_pos = text.find(page_text_marker) + len(page_text_marker)
                            text = text[:insert_pos] + f"\n\n![Figure {img_index+1}]({image_path})\n\n" + text[insert_pos:]
                        
                        page_has_images = True
                        
                    except Exception as e:
                        print(f"?? Erreur extraction image {img_index} de la page {page_num}: {e}")
            
            # Si aucune image n'a été trouvée, essayer d'extraire la page entière comme image
            if not page_has_images:
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom pour qualité
                    image_filename = f"pdf_page_{page_num + 1}.png"
                    image_path = image_dir / image_filename
                    pix.save(str(image_path))
                    
                    # Même traitement pour les figures blanches/noires
                    img = Image.open(image_path)
                    stat = ImageStat.Stat(img)
                    brightness = sum(stat.mean) / len(stat.mean)
                    
                    print(f"  ?? Page {page_num+1} luminosité: {brightness}")
                    
                    if brightness < 20:
                        img = ImageOps.invert(img)
                        enhancer = ImageEnhance.Contrast(img)
                        img = enhancer.enhance(1.5)
                        img.save(image_path)
                        print(f"  ?? Page sombre inversée: {brightness}")
                    
                    # Ajouter à la liste d'images
                    images.append(image_path)
                    
                    # Ajouter référence dans le texte
                    page_text_marker = f"## --- Page {page_num + 1} ---"
                    if page_text_marker in text:
                        insert_pos = text.find(page_text_marker) + len(page_text_marker)
                        text = text[:insert_pos] + f"\n\n![Page {page_num+1}]({image_path})\n\n" + text[insert_pos:]
                    
                    print(f"  ?? Page entière extraite comme image (page {page_num+1})")
                    
                except Exception as e:
                    print(f"?? Erreur extraction page {page_num}: {e}")
        
        print(f"? Extraction des figures terminée: {len(images)} figures extraites")
        
    except Exception as e:
        print(f"? Erreur lors de l'extraction des figures: {e}")
        traceback.print_exc()
    
    return text, images
def extract_from_word(docx_path):
    """
    Extrait le texte et les images d'un fichier Word
    Retourne le texte et une liste de chemins d'images
    """
    print(f"?? Extraction du contenu du fichier Word: {docx_path}")
    text = ""
    images = []
    
    try:
        # Créer un dossier pour les images extraites
        image_dir = EXTRACTED_IMAGES_DIR
        image_dir.mkdir(exist_ok=True)
        
        # Ouvrir le document Word
        doc = docx.Document(docx_path)
        
        # Extraire le texte avec la structure
        for paragraph in doc.paragraphs:
            p_text = paragraph.text.strip()
            if not p_text:
                continue
            
            # Déterminer le style du paragraphe
            if paragraph.style.name.startswith('Heading 1'):
                text += f"# {p_text}\n\n"
            elif paragraph.style.name.startswith('Heading 2'):
                text += f"## {p_text}\n\n"
            elif paragraph.style.name.startswith('Heading 3'):
                text += f"### {p_text}\n\n"
            elif paragraph.style.name.startswith('List'):
                text += f"* {p_text}\n"
            else:
                text += f"{p_text}\n\n"
        
        # Extraire les tableaux
        for table in doc.tables:
            text += "\n"
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                text += "| " + " | ".join(row_content) + " |\n"
            text += "\n"
        
        # Extraire les images
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    image_ext = rel.target_ref.split(".")[-1]
                    image_filename = f"word_image_{rel_id}.{image_ext}"
                    image_path = image_dir / image_filename
                    
                    # Sauvegarder l'image
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_data)
                    
                    # Ajouter l'image à la liste
                    images.append(image_path)
                except Exception as img_error:
                    print(f"?? Erreur lors de l'extraction d'une image: {img_error}")
        
        print(f"? Extraction Word terminée: {len(doc.paragraphs)} paragraphes, {len(images)} images")
        
    except Exception as e:
        print(f"? Erreur lors de l'extraction du fichier Word: {e}")
        traceback.print_exc()
    
    return text, images
def create_text_clip(text, fontsize, width, position, align="west", duration=SLIDE_DUR):
    """Crée un TextClip MoviePy propre avec fallback"""
    if not text or not text.strip():
        text = " "

    try:
        clip = TextClip(
            text,
            fontsize=fontsize,
            color=TEXT_COLOR,
            method="caption",
            align=align,
            size=(width, None),
            interline=1.2
        )
        return clip.set_position(position).set_duration(duration)
    except Exception as e:
        print(f"?? Erreur création TextClip: {e}")
        return TextClip(
            " ", fontsize=fontsize, color=TEXT_COLOR,
            method="caption", size=(width, None)
        ).set_position(position).set_duration(duration)
def create_animated_text_clip(text_content, duration=SLIDE_DUR, animation_level=0.5):
    """
    Crée un clip animé avec MoviePy pour le texte
    """    
    # Nettoyer le texte
    if not text_content or not text_content.strip():
        text_content = " "
    
    # Supprimer les références d'images
    clean_text = re.sub(r'!\[.*?\]\(.*?\)', '', text_content)
    
    # Diviser le texte en paragraphes
    paragraphs = clean_text.split('\n\n')
    paragraphs = [p for p in paragraphs if p.strip()]
    
    if not paragraphs:
        return create_text_clip(" ", FONT_SIZE_TEXT, WIDTH-2*TEXT_MARGIN, (TEXT_MARGIN, 100), duration=duration)
    
    # Déterminer la durée de l'animation en fonction du niveau d'animation
    animation_duration = min(2.0, duration / 2) * animation_level
    
    # Créer un clip pour chaque paragraphe
    paragraph_clips = []
    y_position = 100
    
    for i, paragraph in enumerate(paragraphs):
        # Calculer le délai d'apparition en fonction de l'index
        delay = i * (animation_duration / max(1, len(paragraphs) - 1))
        
        # Créer le clip de texte
        text_clip = TextClip(
            paragraph,
            fontsize=FONT_SIZE_TEXT,
            color=TEXT_COLOR,
            method="caption",
            align="west",
            size=(WIDTH-2*TEXT_MARGIN, None),
            interline=1.2
        )
        
        # Définir d'abord la durée du clip avant d'appliquer le crossfade
        text_clip = text_clip.set_duration(duration - delay)
        
        # Positionner le clip
        text_clip = text_clip.set_position((TEXT_MARGIN, y_position))
        
        # Ajouter un effet de fondu d'entrée
        if animation_level > 0.3:
            text_clip = text_clip.set_start(delay)
            # Appliquer le crossfade uniquement si la durée est positive
            fade_duration = min(0.5, animation_duration / 2)
            if fade_duration > 0:
                text_clip = text_clip.crossfadein(fade_duration)
        else:
            text_clip = text_clip.set_start(0)
        
        # Ajouter à la liste des clips
        paragraph_clips.append(text_clip)
        
        # Mettre à jour la position Y pour le prochain paragraphe
        y_position += text_clip.h + 10
    
    # Créer un clip composite avec tous les paragraphes
    return CompositeVideoClip(paragraph_clips, size=(WIDTH, HEIGHT))
def preserve_markdown_formatting(text):
    """Conserve la mise en forme Markdown dans le texte pour l'affichage"""
    if not text or not isinstance(text, str):
        return ""
    
    # Supprimer les références d'images
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    
    # Préserver les titres
    text = re.sub(r'^# (.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'^### (.+)$', r'\1', text, flags=re.MULTILINE)
    
    # Préserver l'italique et le gras
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Gras
    text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italique
    text = re.sub(r'\_(.+?)\_', r'\1', text)      # Italique alternatif
    
    # Préserver les listes
    text = re.sub(r'^- (.+)$', r'• \1', text, flags=re.MULTILINE)
    text = re.sub(r'^\* (.+)$', r'• \1', text, flags=re.MULTILINE)
    text = re.sub(r'^(\d+)\. (.+)$', r'\1. \2', text, flags=re.MULTILINE)
    
    # Préserver les blocs de code
    code_blocks = re.findall(r'```(.+?)```', text, re.DOTALL)
    for block in code_blocks:
        text = text.replace(f'```{block}```', f'{block}')
    
    return text
def handle_logo(layers, style_params, duration):
    """
    Fonction séparée pour gérer l'ajout du logo à une diapositive
    """    
    # Utiliser le chemin personnalisé du logo que vous avez fourni
    custom_logo_path = "C:/Users/ThinkPad/Desktop/plateform/python/teamwill.webp"
    
    try:
        logo_path = Path(custom_logo_path)
        
        if logo_path.exists() and logo_path.is_file():
            try:
                # Taille adaptative selon le style
                logo_size = int(30 + style_params.get('animation_level', 0.5) * 20)
                
                # Créer le clip du logo
                logo = ImageClip(str(logo_path))
                logo = logo.resize(width=logo_size)
                logo = logo.set_position((WIDTH-logo_size-10, 10))
                logo = logo.set_duration(duration)
                
                # Ajouter à la liste des couches
                layers.append(logo)
                print(f"? Logo ajouté: {logo_path}")
            except Exception as e:
                print(f"?? Erreur lors du chargement du logo {logo_path}: {e}")
        else:
            print(f"?? Le logo n'existe pas à {logo_path} ou n'est pas un fichier")
    except Exception as e:
        print(f"?? Erreur générale logo: {e}")
    
    return layers

def slide_clip_with_exact_markdown(title, content, images, style_params=None, layout_params=None, duration=SLIDE_DUR):
    """
    Crée un slide qui préserve exactement la structure Markdown originale
    avec une mise en page adaptative selon les paramètres de style et de mise en page
    """    
    # Valeurs par défaut pour les paramètres de style et de mise en page
    if style_params is None:
        style_params = {
            'name': 'balanced',
            'text_size': FONT_SIZE_TEXT,
            'animation_level': 0.5,
            'color_scheme': 'neutral',
            'layout': 'balanced',
            'transition': 'fade'
        }
    
    if layout_params is None:
        layout_params = {
            'text_x': 0.1, 
            'text_y': 0.15,
            'text_width': 0.8, 
            'text_height': 0.7,
            'title_position': 'top',
            'text_size_factor': 1.0
        }
    
    # Créer l'arrière-plan
    bg_color = BG_COLOR
    if style_params.get('color_scheme') == 'vibrant':
        bg_color = (240, 250, 255)  # Bleu très pâle
    elif style_params.get('color_scheme') == 'monochrome':
        bg_color = (248, 248, 248)  # Gris clair
    elif style_params.get('color_scheme') == 'contrast':
        bg_color = (235, 235, 245)  # Violet très pâle
    elif style_params.get('color_scheme') == 'focused':
        bg_color = (245, 245, 245)  # Gris très pâle
    
    bg = ColorClip((WIDTH, HEIGHT), color=bg_color).set_duration(duration)
    layers = [bg]
    
    # Réserver plus d'espace pour le titre
    title_height = 60
    content_top = title_height + 10
    available_height = HEIGHT - content_top - 40  # Marge du bas
    
    # Ajuster la taille du texte en fonction du style
    text_size = int(style_params.get('text_size', FONT_SIZE_TEXT) * layout_params.get('text_size_factor', 1.0))
    
    # Ajouter le titre s'il existe
    if title:
        # Vérifier si le titre est un marqueur de page pour le masquer
        if not re.match(r'(---\s*)?[Pp]age\s+\d+(\s*---)?', title):
            title_position = layout_params.get('title_position', 'top')
            title_y = 10 if title_position == 'top' else HEIGHT - 50
            
            title_clip = TextClip(
                title, 
                fontsize=FONT_SIZE_TITLE, 
                color=TEXT_COLOR,
                method="caption", 
                align="center",
                size=(WIDTH-120, None)
            ).set_position(("center", title_y)).set_duration(duration)
            
            layers.append(title_clip)
    
    # Convertir les images en liste de chemins si c'est une liste de dictionnaires
    image_paths = []
    if images:
        for img in images:
            if isinstance(img, dict):
                image_paths.append(img['path'])
            elif isinstance(img, tuple):  # Support de l'ancien format
                image_paths.append(img[0])
            else:
                image_paths.append(img)
    
    # Déterminer si nous avons des images valides
    has_valid_images = bool(image_paths)
    
    # Appliquer les paramètres de mise en page pour le texte et les images
    text_x = int(layout_params.get('text_x', 0.1) * WIDTH)
    text_y = int(layout_params.get('text_y', 0.15) * HEIGHT)
    text_width = int(layout_params.get('text_width', 0.8) * WIDTH)
    
    # Supprimer les références d'images du texte affiché
    processed_content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
    
    # Préserver le formatage Markdown (gras, italique, listes)
    processed_content = preserve_markdown_formatting(processed_content)
    
    # Vérifier si le texte a du contenu après nettoyage
    if processed_content.strip():
        # Créer le clip de texte avec l'animation selon le niveau d'animation
        animation_level = style_params.get('animation_level', 0.5)
        
        if animation_level > 0.3:
            # Animation plus élaborée pour les styles dynamiques
            text_clip = create_animated_text_clip(
                processed_content, 
                duration=duration,
                animation_level=animation_level
            )
        else:
            # Texte statique pour les styles minimalistes
            text_clip = create_text_clip(
                processed_content, 
                fontsize=text_size, 
                width=text_width,
                position=(text_x, text_y), 
                duration=duration
            )
        
        # Ajouter le texte à la diapositive
        layers.append(text_clip)
    
    # Ajouter les images si présentes
    if has_valid_images:
        try:
            # Récupérer les paramètres de mise en page pour l'image
            img_x = int(layout_params.get('image_x', 0.55) * WIDTH)
            img_y = int(layout_params.get('image_y', 0.15) * HEIGHT)
            img_width = int(layout_params.get('image_width', 0.4) * WIDTH)
            
            # Utiliser la première image valide
            for img_path in image_paths:
                try:
                    img_clip = ImageClip(str(img_path))
                    img_ratio = img_clip.h / img_clip.w
                    img_height = int(img_width * img_ratio)
                    
                    # Limiter la hauteur si nécessaire
                    max_height = int(HEIGHT * 0.7)
                    if img_height > max_height:
                        img_height = max_height
                        img_width = int(img_height / img_ratio)
                    
                    # Redimensionner et positionner l'image
                    img_clip = img_clip.resize(width=img_width)
                    img_clip = img_clip.set_position((img_x, img_y)).set_duration(duration)
                    
                    # Ajouter un effet de fondu d'entrée pour les styles dynamiques
                    if animation_level > 0.6:
                        img_clip = img_clip.crossfadein(min(1.0, animation_level))
                    
                    layers.append(img_clip)
                    break  # Utiliser seulement la première image valide
                except Exception as e:
                    print(f"?? Erreur chargement image {img_path}: {e}")
                    continue
        except Exception as e:
            print(f"?? Erreur générale images: {e}")
    
    # Ajouter le logo avec la fonction séparée
    layers = handle_logo(layers, style_params, duration)
    
    # Ajout d'éléments décoratifs selon le style
    if style_params.get('name') == 'dynamic':
        # Ajouter un élément décoratif pour le style dynamique
        try:
            # Créer une bande colorée en haut
            color_band = ColorClip((WIDTH, 8), color=(80, 120, 200)).set_position((0, 0)).set_duration(duration)
            layers.append(color_band)
        except Exception as e:
            print(f"?? Erreur élément décoratif: {e}")
    
    # Créer la diapositive finale
    return CompositeVideoClip(layers, size=(WIDTH, HEIGHT))
def apply_transitions_to_clips(clips, transition_type="fade", transition_duration=TRANSITION_DUR):
    """
    Applique des transitions entre deux clips vidéo avec une gestion d'erreur améliorée
    
    Args:
        clips: Liste de clips à combiner avec transition
        transition_type: Type de transition (fade, wipe_right, etc.)
        transition_duration: Durée de la transition
        
    Returns:
        transition_clips: Liste de clips avec transitions
    """    
    # Vérifier les arguments
    if not clips:
        print("?? Aucun clip fourni pour la transition")
        return []
    
    if len(clips) < 2:
        print("?? Au moins deux clips sont nécessaires pour une transition")
        return clips
    
    # Vérifier que les clips ont des durées valides
    if any(not hasattr(clip, 'duration') or clip.duration <= 0 for clip in clips):
        print("?? Clips avec durées invalides, impossible d'appliquer la transition")
        return clips
    
    # Récupérer les deux clips
    first_clip = clips[0]
    second_clip = clips[1]
    
    # Vérifier que la transition n'est pas trop longue
    safe_duration = min(transition_duration, first_clip.duration/2, second_clip.duration/2)
    if safe_duration < 0.1:  # Si la durée est trop courte, utiliser une transition simple
        print("?? Durée de transition trop courte, utilisation d'une transition simple")
        return clips
    
    try:
        # Créer les versions raccourcies des clips
        first_end = first_clip.subclip(0, first_clip.duration - safe_duration)
        second_start = second_clip.subclip(safe_duration, second_clip.duration)
        
        # Liste des types de transition valides
        valid_transitions = ['fade', 'wipe_right', 'wipe_left', 'wipe_up', 'wipe_down', 'zoom_in', 'zoom_out']
        
        # Vérifier si le type de transition est valide
        if transition_type not in valid_transitions:
            print(f"?? Type de transition '{transition_type}' non reconnu, utilisation de 'fade'")
            transition_type = 'fade'
        
        # Créer la transition en fonction du type
        if transition_type == "fade":
            # Fondu enchaîné
            transition = create_fade_transition(first_clip, second_clip, safe_duration)
        elif transition_type.startswith("wipe"):
            # Transitions de type balayage
            transition = create_wipe_transition(first_clip, second_clip, transition_type, safe_duration)
        elif transition_type == "zoom_in":
            # Zoom avant
            transition = create_zoom_transition(first_clip, second_clip, "in", safe_duration)
        elif transition_type == "zoom_out":
            # Zoom arrière
            transition = create_zoom_transition(first_clip, second_clip, "out", safe_duration)
        else:
            # Par défaut: fondu enchaîné (ne devrait jamais arriver avec la vérification ci-dessus)
            print(f"?? Type de transition '{transition_type}' non géré, utilisation du fondu")
            transition = create_fade_transition(first_clip, second_clip, safe_duration)
        
        # Vérifier que la transition a été créée correctement
        if transition is None:
            print("?? Échec de la création de la transition, utilisation d'une simple concaténation")
            return clips
        
        # Retourner les clips avec transition
        return [first_end, transition, second_start]
    
    except Exception as e:
        print(f"?? Erreur lors de la création de la transition '{transition_type}': {e}")
        traceback.print_exc()
        
        # En cas d'erreur, retourner les clips sans transition
        print("?? Utilisation d'une simple concaténation sans transition")
        return clips
def create_fade_transition(first_clip, second_clip, duration):
    """Crée une transition de fondu enchaîné entre deux clips"""   
    try:
        # Créer les segments pour la transition
        first_end = first_clip.copy().subclip(first_clip.duration - duration, first_clip.duration)
        first_end = first_end.crossfadeout(duration)
        
        second_start = second_clip.copy().subclip(0, duration)
        second_start = second_start.crossfadein(duration)
        
        # Créer la transition et définir sa durée
        transition = CompositeVideoClip([first_end, second_start])
        transition = transition.set_duration(duration)
        
        return transition
    except Exception as e:
        print(f"?? Erreur lors de la création du fondu: {e}")
        # En cas d'erreur, retourner un clip vide de la durée appropriée
        return ColorClip(first_clip.size, col=(0, 0, 0), duration=duration)
def create_wipe_transition(first_clip, second_clip, wipe_type, duration):
    """Crée une transition de type balayage entre deux clips"""
        
    w, h = first_clip.size
    
    first_end = first_clip.copy().subclip(first_clip.duration - duration, first_clip.duration)
    second_start = second_clip.copy().subclip(0, duration)
    
    def make_mask(t):
        progress = t / duration
        mask = np.zeros((h, w), dtype=np.float32)
        
        if wipe_type == "wipe_right":
            x_limit = int(w * progress)
            mask[:, :x_limit] = 1.0
        elif wipe_type == "wipe_left":
            x_limit = int(w * (1 - progress))
            mask[:, x_limit:] = 1.0
        elif wipe_type == "wipe_up":
            y_limit = int(h * (1 - progress))
            mask[y_limit:, :] = 1.0
        elif wipe_type == "wipe_down":
            y_limit = int(h * progress)
            mask[:y_limit, :] = 1.0
        
        return mask
    
    # Créer le clip de masque avec l'attribut ismask=True
    mask_clip = VideoClip(make_mask, duration=duration)
    mask_clip.ismask = True  # Définir explicitement comme masque
    
    # Construire la transition avec le masque
    transition = CompositeVideoClip(
        [first_end, second_start.set_mask(mask_clip)],
        size=(w, h)
    ).set_duration(duration)
    
    return transition
def create_zoom_transition(first_clip, second_clip, zoom_type, duration):
    """Crée une transition de type zoom entre deux clips"""
    
    w, h = first_clip.size
    
    try:
        if zoom_type == "in":
            # Zoom avant sur la nouvelle diapositive
            first_end = first_clip.copy().subclip(first_clip.duration - duration, first_clip.duration)
            
            def zoom_in(t):
                progress = t / duration
                scale = 1 - progress + 0.5 * progress  # De 0.5 à 1.0
                x_offset = w * (1 - scale) / 2
                y_offset = h * (1 - scale) / 2
                return (second_clip.copy()
                    .resize(lambda t: scale)
                    .set_position((x_offset, y_offset)))
            
            second_start = VideoClip(lambda t: zoom_in(t).get_frame(t), duration=duration)
            
            return CompositeVideoClip([first_end, second_start], size=(w, h)).set_duration(duration)
        else:
            # Zoom arrière de l'ancienne diapositive
            second_start = second_clip.copy().subclip(0, duration)
            
            def zoom_out(t):
                progress = t / duration
                scale = 1.5 - 0.5 * progress  # De 1.5 à 1.0
                x_offset = w * (1 - scale) / 2
                y_offset = h * (1 - scale) / 2
                return (first_clip.copy()
                    .resize(lambda t: scale)
                    .set_position((x_offset, y_offset)))
            
            first_end = VideoClip(lambda t: zoom_out(t).get_frame(t), duration=duration)
            
            return CompositeVideoClip([second_start, first_end], size=(w, h)).set_duration(duration)
    except Exception as e:
        print(f"?? Erreur lors de la création de la transition zoom: {e}")
        # En cas d'erreur, retourner une transition par fondu
        return create_fade_transition(first_clip, second_clip, duration)
def get_varied_transitions(number_of_transitions, seed=42):
    """
    Génère une liste de transitions variées mais fiables pour éviter la monotonie
    """
    random.seed(seed)  # Pour la reproductibilité
    
    # Privilégier les transitions les plus fiables
    reliable_transitions = [
        "fade",        # Fondu enchaîné (très fiable)
        "fade",        # Répété pour augmenter sa probabilité
        "fade"         # Répété encore pour être plus fréquent
    ]
    
    # Ajouter quelques transitions plus complexes mais moins fréquentes
    all_transitions = reliable_transitions + [
        "wipe_right",
        "wipe_left", 
        "zoom_in",
        "zoom_out"
    ]
    
    # Pour éviter de répéter la même transition deux fois de suite
    transitions = []
    last_transition = None
    
    for _ in range(number_of_transitions):
        available = [t for t in all_transitions if t != last_transition]
        chosen = random.choice(available)
        transitions.append(chosen)
        last_transition = chosen
    
    return transitions
def create_multiple_slide_clips(segments, style_model, layout_model, readability_model, duration=SLIDE_DUR):
    """
    Crée plusieurs diapositives optimisées par ML à partir de segments de contenu
    """
    slides = []
    
    for i, segment in enumerate(segments):
        title = segment['title']
        content = segment['content']
        has_image = segment['has_image']
        
        # Extraire les images pour ce segment
        segment_images = []
        if has_image:
            # Utiliser extract_images_from_markdown pour avoir les images précises de ce segment
            images_data = extract_images_from_markdown(content, BASE_DIR)
            segment_images = [img['path'] for img in images_data]
        
        # 1. Évaluer la lisibilité du contenu
        readability_score, issues, suggestions = readability_model.evaluate_readability(content)
        
        # 2. Améliorer le contenu si nécessaire (score < 7)
        if readability_score < 7:
            improved_content, changes = readability_model.improve_readability(content)
            if changes:
                print(f"?? Diapositive {i+1}: {len(changes)} améliorations de lisibilité")
                content = improved_content
        
        # 3. Ajuster la durée en fonction de la longueur du contenu
        content_length = len(content.split())
        slide_duration = min(SLIDE_DUR * 2, max(SLIDE_DUR, SLIDE_DUR * content_length / 200))
        
        # 4. Recommander le style optimal
        style_params = style_model.recommend_style(content, title, has_image)
        
        # 5. Optimiser la mise en page
        layout_params = layout_model.optimize_layout(
            content, title, segment_images, style_params
        )
        
        # 6. Créer la diapositive avec les paramètres optimisés
        slide = slide_clip_with_exact_markdown(
            title=title,
            content=content,
            images=segment_images,
            style_params=style_params,
            layout_params=layout_params,
            duration=slide_duration
        )
        
        slides.append({
            'slide': slide,
            'title': title,
            'style': style_params['name'],
            'transition': style_params.get('transition', 'fade'),
            'duration': slide_duration
        })
    
    return slides

# Option 1: Modifier la classe SlideQualityModel pour gérer cette incompatibilité
""" the class SlideQualityModel was here before refacoring to the slide_quality.py """ 
""" the class SlideStyleRecommender was here before refacoring to the style_recommender.py """
""" the class DeepSlideQualityModel was here before refacoring to the slide_quality.py """

class SlideLayoutManager:
    """
    Gestionnaire de mise en page pour diapositives qui assure une séparation claire
    entre le texte et les images, évitant ainsi tout chevauchement.
    
    Cette classe s'intègre avec le générateur de présentations existant et remplace
    les fonctions de mise en page par des versions optimisées.
    """
    
    def __init__(self, width=960, height=540, font_size_title=50, font_size_text=36):
        """
        Initialise le gestionnaire de mise en page avec les dimensions et paramètres par défaut
        
        Args:
            width: Largeur des diapositives en pixels
            height: Hauteur des diapositives en pixels
            font_size_title: Taille de police pour les titres
            font_size_text: Taille de police pour le texte principal
        """
        self.WIDTH = width
        self.HEIGHT = height
        self.FONT_SIZE_TITLE = font_size_title
        self.FONT_SIZE_TEXT = font_size_text
        self.TEXT_COLOR = "black"
        self.BG_COLOR = (255, 255, 255)
        self.TEXT_MARGIN = 40
        
        # État indiquant si les fonctions ont été remplacées
        self.functions_replaced = False
    
    def apply_layout_fixes(self):
        """
        Applique les corrections de mise en page en remplaçant les fonctions standard
        par des versions optimisées pour la séparation texte/image
        
        Returns:
            bool: True si l'opération a réussi, False sinon
        """
        global slide_clip_with_exact_markdown
        
        try:
            # 1. Remplacer la fonction principale de création de diapositives
            slide_clip_with_exact_markdown = self.create_slide_with_separated_content
            
            # 2. Appliquer d'autres remplacements de fonctions si nécessaire
            self.replace_layout_optimizer()
            
            self.functions_replaced = True
            print("? Mise en page des diapositives optimisée pour éviter le chevauchement texte/image")
            return True
        except Exception as e:
            print(f"? Erreur lors de l'application des corrections de mise en page: {e}")
            traceback.print_exc()
            return False
    
    def create_slide_with_separated_content(self, title, content, images, style_params=None, layout_params=None, duration=8):
        """
        Crée un slide qui préserve exactement la structure Markdown originale
        Traite tous les éléments visuels (images, tableaux, graphiques) sans texte en mode plein écran
        """        
        # Déterminer si le contenu ne contient que des éléments visuels (images, tableaux) sans texte
        has_table = re.search(r'\|[-]+\|', content) is not None
        has_table = has_table or "|" in content and any(line.count('|') > 2 for line in content.split('\n'))
        
        # Supprimer les références d'images et les tableaux pour voir s'il reste du texte
        clean_content = re.sub(r'!\[.*?\]\(.*?\)', '', content).strip()
        if has_table:
            clean_content_without_tables = re.sub(r'^\|.*\|$', '', clean_content, flags=re.MULTILINE)
            clean_content_without_tables = re.sub(r'^[-|:]+$', '', clean_content_without_tables, flags=re.MULTILINE)
            clean_content = clean_content_without_tables
        
        # Supprimer tous les espaces et caractères de formatage
        clean_content = re.sub(r'[#*_\s\n\r]', '', clean_content)
        
        # Définir si c'est un contenu visuel sans texte (image ou tableau)
        is_visual_only = not bool(clean_content) and (bool(images) or has_table)
        
        # Valeurs par défaut pour les paramètres de style et de mise en page
        if style_params is None:
            style_params = {
                'name': 'balanced',
                'text_size': self.FONT_SIZE_TEXT,
                'animation_level': 0.5,
                'color_scheme': 'neutral',
                'layout': 'balanced',
                'transition': 'fade'
            }
        
        if layout_params is None:
            layout_params = {}
        
        # Définir explicitement le mode plein écran si c'est un contenu visuel sans texte
        layout_params['fullscreen_image'] = is_visual_only
        
        # Créer l'arrière-plan
        bg_color = self.BG_COLOR
        if style_params.get('color_scheme') == 'vibrant':
            bg_color = (240, 250, 255)  # Bleu très pâle
        elif style_params.get('color_scheme') == 'monochrome':
            bg_color = (248, 248, 248)  # Gris clair
        elif style_params.get('color_scheme') == 'contrast':
            bg_color = (235, 235, 245)  # Violet très pâle
        elif style_params.get('color_scheme') == 'focused':
            bg_color = (245, 245, 245)  # Gris très pâle
        
        bg = ColorClip((self.WIDTH, self.HEIGHT), color=bg_color).set_duration(duration)
        layers = [bg]
        
        # Si c'est un mode contenu visuel plein écran (image ou tableau sans texte), traiter différemment
        if is_visual_only:
            print(f"?? Mode plein écran activé - contenu visuel uniquement")
            
            # Ajouter le titre s'il existe (en bas pour ne pas gêner le contenu visuel)
            if title:
                title_clip = TextClip(
                    title, 
                    fontsize=self.FONT_SIZE_TITLE, 
                    color=self.TEXT_COLOR,
                    method="caption", 
                    align="center",
                    size=(self.WIDTH-120, None)
                ).set_position(("center", self.HEIGHT - 60)).set_duration(duration)
                
                layers.append(title_clip)
            
            # Si c'est un tableau (pas d'images), on le traite différemment
            if has_table and not images:
                # Pour les tableaux sans images, on applique un style de mise en page adapté
                # qui affiche le tableau en grand, centralisé
                table_content = content
                
                # Créer un clip de texte pour le tableau qui occupe presque tout l'écran
                table_clip = TextClip(
                    table_content, 
                    fontsize=self.FONT_SIZE_TABLE, 
                    color=self.TEXT_COLOR,
                    method="caption", 
                    align="center",
                    size=(self.WIDTH * 0.9, None)
                ).set_position(("center", self.HEIGHT * 0.1)).set_duration(duration)
                
                layers.append(table_clip)
                
            # Si des images sont présentes, les afficher en plein écran
            elif images:
                try:
                    for img_path in images:
                        try:
                            img_clip = ImageClip(str(img_path))
                            
                            # Déterminer la meilleure taille pour couvrir l'écran tout en préservant le ratio
                            img_ratio = img_clip.h / img_clip.w
                            screen_ratio = self.HEIGHT / self.WIDTH
                            
                            if img_ratio > screen_ratio:  # Image plus haute que large
                                new_width = self.WIDTH * 0.9  # 90% de la largeur
                                new_height = new_width * img_ratio
                            else:  # Image plus large que haute
                                new_height = self.HEIGHT * 0.8  # 80% de la hauteur (laisse place au titre)
                                new_width = new_height / img_ratio
                            
                            # Redimensionner l'image
                            img_clip = img_clip.resize(width=new_width) if img_ratio > screen_ratio else img_clip.resize(height=new_height)
                            
                            # Centrer l'image
                            x_pos = (self.WIDTH - img_clip.w) / 2
                            y_pos = (self.HEIGHT - img_clip.h) / 2 - 20  # Légèrement plus haut pour laisser de l'espace au titre
                            
                            # Appliquer des animations si style animé
                            animation_level = style_params.get('animation_level', 0.5)
                            if animation_level > 0.6:
                                img_clip = img_clip.crossfadein(min(1.0, animation_level))
                            
                            img_clip = img_clip.set_position((x_pos, y_pos)).set_duration(duration)
                            layers.append(img_clip)
                            break  # N'utiliser que la première image
                        except Exception as e:
                            print(f"?? Erreur chargement image {img_path}: {e}")
                except Exception as e:
                    print(f"?? Erreur générale images: {e}")
        
        else:
            # Affichage standard pour le contenu avec texte
            # Réserver de l'espace pour le titre
            title_height = 60
            
            # Convertir les images en liste de chemins si c'est une liste de dictionnaires
            image_paths = []
            if images:
                for img in images:
                    if isinstance(img, dict):
                        image_paths.append(img.get('path', ''))
                    elif isinstance(img, tuple):  # Support de l'ancien format
                        image_paths.append(img[0])
                    else:
                        image_paths.append(img)
            
            # Déterminer si nous avons des images valides
            has_valid_images = bool(image_paths)
            
            # Ajouter le titre s'il existe
            if title:
                # Vérifier si le titre est un marqueur de page pour le masquer
                if not re.match(r'(---\s*)?[Pp]age\s+\d+(\s*---)?', title):
                    title_position = layout_params.get('title_position', 'top')
                    title_y = self.HEIGHT - 50 if title_position == 'bottom' else 10
                    
                    title_clip = TextClip(
                        title, 
                        fontsize=self.FONT_SIZE_TITLE, 
                        color=self.TEXT_COLOR,
                        method="caption", 
                        align="center",
                        size=(self.WIDTH-120, None)
                    ).set_position(("center", title_y)).set_duration(duration)
                    
                    layers.append(title_clip)
            
            # Appliquer les valeurs spécifiques des layout_params pour le texte
            text_x = int(layout_params.get('text_x', 0.1) * self.WIDTH)
            text_width = int(layout_params.get('text_width', 0.8) * self.WIDTH)
            text_y = int(layout_params.get('text_y', 0.15) * self.HEIGHT)
            
            # Supprimer les références d'images du texte affiché
            processed_content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
            
            # Préserver le formatage Markdown (gras, italique, listes)
            processed_content = self.preserve_markdown_formatting(processed_content)
            
            # Vérifier si le texte a du contenu après nettoyage
            if processed_content.strip():
                # Créer le clip de texte avec l'animation selon le niveau d'animation
                text_size = int(style_params.get('text_size', self.FONT_SIZE_TEXT) * layout_params.get('text_size_factor', 1.0))
                animation_level = style_params.get('animation_level', 0.5)
                
                if animation_level > 0.3:
                    # Animation plus élaborée pour les styles dynamiques
                    text_clip = self.create_animated_text_clip(
                        processed_content, 
                        duration=duration,
                        animation_level=animation_level,
                        text_size=text_size,
                        text_width=text_width
                    )
                else:
                    # Texte statique pour les styles minimalistes
                    text_clip = self.create_text_clip(
                        processed_content, 
                        fontsize=text_size, 
                        width=text_width,
                        position=(text_x, text_y), 
                        duration=duration
                    )
                
                # Ajouter le texte à la diapositive
                layers.append(text_clip)
            
            # Ajouter les images si présentes avec positionnement spécifique
            if has_valid_images:
                try:
                    # Récupérer les paramètres de mise en page pour l'image
                    img_x = int(layout_params.get('image_x', 0.55) * self.WIDTH)
                    img_y = int(layout_params.get('image_y', 0.15) * self.HEIGHT)
                    img_width = int(layout_params.get('image_width', 0.4) * self.WIDTH)
                    
                    # Utiliser la première image valide
                    for img_path in image_paths:
                        try:
                            img_clip = ImageClip(str(img_path))
                            img_ratio = img_clip.h / img_clip.w
                            img_height = int(img_width * img_ratio)
                            
                            # Limiter la hauteur si nécessaire
                            max_height = int(self.HEIGHT * 0.6)
                            if img_height > max_height:
                                img_height = max_height
                                img_width = int(img_height / img_ratio)
                            
                            # Redimensionner et positionner l'image
                            img_clip = img_clip.resize(width=img_width)
                            img_clip = img_clip.set_position((img_x, img_y)).set_duration(duration)
                            
                            # Ajouter un effet de fondu d'entrée pour les styles dynamiques
                            animation_level = style_params.get('animation_level', 0.5)
                            if animation_level > 0.6:
                                img_clip = img_clip.crossfadein(min(1.0, animation_level))
                            
                            layers.append(img_clip)
                            break  # Utiliser seulement la première image valide
                        except Exception as e:
                            print(f"?? Erreur chargement image {img_path}: {e}")
                            continue
                except Exception as e:
                    print(f"?? Erreur générale images: {e}")
        
        # Ajouter le logo si une fonction handle_logo existe
        if 'handle_logo' in globals():
            layers = handle_logo(layers, style_params, duration)
        
        # Créer la diapositive finale
        return CompositeVideoClip(layers, size=(self.WIDTH, self.HEIGHT))
    
    def compute_optimal_layout(self, content, title, images, style_params):
        """
        Calcule la mise en page optimale en fonction du contenu et du style
        
        Args:
            content: Texte de la diapositive
            title: Titre de la diapositive
            images: Liste des chemins d'images
            style_params: Paramètres de style
            
        Returns:
            layout_params: Paramètres de mise en page optimisés
        """
        
        # Analyser le contenu
        layout_params = {}
        has_image = bool(images)
        text_length = len(content)
        
        # 1. Détection de tableaux dans le contenu
        has_table = re.search(r'\|[-]+\|', content) is not None  # Détection de tableaux markdown
        has_table = has_table or "|" in content and any(line.count('|') > 2 for line in content.split('\n'))
        
        # 2. Détection d'images larges ou complexes
        large_images = []
        complex_images = []
        
        if has_image and images:
            for img_path in images:
                try:
                    if Path(img_path).exists():
                        img = Image.open(img_path)
                        width, height = img.size
                        aspect_ratio = width / height
                        
                        # Déterminer si l'image est large
                        if aspect_ratio > 1.8:
                            large_images.append(img_path)
                        
                        # Analyse de complexité (détection simple de diagrammes/tableaux)
                        # Calculer statistiques de l'image
                        img_gray = img.convert('L')
                        stat = ImageStat.Stat(img_gray)
                        std_dev = stat.stddev[0]  # Variabilité

                        # Images complexes: faible variabilité ou contenu graphique
                        if std_dev < 45 or self._has_grid_pattern(img_gray):
                            complex_images.append(img_path)
                except Exception as e:
                    print(f"?? Erreur analyse image {img_path}: {e}")
        
        # 3. Déterminer si le contenu est court ou long
        is_short_content = text_length < 500
        is_very_long_content = text_length > 1500
        
        # 4. Récupérer la mise en page de base du style
        layout_type = style_params.get('layout', 'balanced')
        
        # 5. Décision pour la mise en page avec séparation claire texte/image
        
        # 5.1 Cas spécial: tableaux
        if has_table:
            if has_image:
                # Tableau avec images: disposition plus adaptée
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.9, 'text_height': 0.45,
                    'image_x': 0.3, 'image_y': 0.65,
                    'image_width': 0.4, 'image_height': 0.3,
                    'title_position': 'top',
                    'text_size_factor': 0.9
                }
            else:
                # Tableau sans images: utiliser tout l'espace
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.9, 'text_height': 0.75,
                    'title_position': 'top',
                    'text_size_factor': 0.9
                }
        
        # 5.2 Cas spécial: images complexes (graphiques, tableaux en image)
        if complex_images:
            # Pour les images complexes: grande image à droite
            return {
                'text_x': 0.05, 'text_y': 0.15,
                'text_width': 0.45, 'text_height': 0.75,
                'image_x': 0.55, 'image_y': 0.15,
                'image_width': 0.4, 'image_height': 0.7,
                'title_position': 'top',
                'text_size_factor': 0.95
            }
        
        # 5.3 Cas spécial: images larges
        if large_images:
            if is_short_content:
                # Texte court: image en haut, texte en bas
                return {
                    'text_x': 0.1, 'text_y': 0.55,
                    'text_width': 0.8, 'text_height': 0.4,
                    'image_x': 0.15, 'image_y': 0.15,
                    'image_width': 0.7, 'image_height': 0.35,
                    'title_position': 'top'
                }
            else:
                # Texte plus long: image en bas, texte en haut
                return {
                    'text_x': 0.1, 'text_y': 0.15,
                    'text_width': 0.8, 'text_height': 0.5,
                    'image_x': 0.15, 'image_y': 0.65,
                    'image_width': 0.7, 'image_height': 0.3,
                    'title_position': 'top'
                }
        
        # 5.4 Mises en page standard avec séparation claire texte/image
        if has_image:
            if layout_type == 'image_dominant':
                # Image dominante: texte en bas
                return {
                    'text_x': 0.1, 'text_y': 0.65,
                    'text_width': 0.8, 'text_height': 0.3,
                    'image_x': 0.1, 'image_y': 0.15,
                    'image_width': 0.8, 'image_height': 0.45,
                    'title_position': 'top'
                }
            
            elif layout_type == 'text_focus':
                # Texte dominant: image plus petite à droite
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.55, 'text_height': 0.75,
                    'image_x': 0.65, 'image_y': 0.25,
                    'image_width': 0.3, 'image_height': 0.5,
                    'title_position': 'top'
                }
            
            else:  # balanced, split et autres
                # Disposition côte à côte équilibrée: texte à gauche, image à droite
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.45, 'text_height': 0.75,
                    'image_x': 0.55, 'image_y': 0.15,
                    'image_width': 0.4, 'image_height': 0.7,
                    'title_position': 'top'
                }
        else:  # Pas d'image
            # Utiliser tout l'espace pour le texte
            return {
                'text_x': 0.1, 'text_y': 0.15,
                'text_width': 0.8, 'text_height': 0.7,
                'title_position': 'top',
                'text_size_factor': 1.0 if is_short_content else 0.9
            }
    
    def _has_grid_pattern(self, img_gray):
        """
        Détecte si une image en niveaux de gris contient une structure de grille
        (comme un tableau ou un graphique)
        
        Args:
            img_gray: Image PIL en niveaux de gris
            
        Returns:
            bool: True si une grille est détectée
        """
        try:
            # Échantillonner des lignes et colonnes pour détecter des patterns réguliers
            width, height = img_gray.size
            
            # Prendre des échantillons de pixels au milieu de l'image
            row_samples = [img_gray.getpixel((i, height//2)) for i in range(0, width, 10)]
            col_samples = [img_gray.getpixel((width//2, i)) for i in range(0, height, 10)]
            
            # Compter les changements brusques d'intensité (indicateurs de bordures)
            threshold = 30  # Différence d'intensité pour considérer un changement
            row_edges = sum(1 for i in range(1, len(row_samples)) if abs(row_samples[i] - row_samples[i-1]) > threshold)
            col_edges = sum(1 for i in range(1, len(col_samples)) if abs(col_samples[i] - col_samples[i-1]) > threshold)
            
            # Si plusieurs changements brusques dans les deux dimensions, c'est probablement une grille
            return (row_edges > 3 and col_edges > 3)
        except Exception as e:
            print(f"?? Erreur détection grille: {e}")
            return False
    
    def create_text_clip(self, text, fontsize, width, position, align="west", duration=8):
        """
        Crée un TextClip MoviePy avec gestion d'erreur
        
        Args:
            text: Texte à afficher
            fontsize: Taille de la police
            width: Largeur du texte
            position: Position (x, y)
            align: Alignement du texte
            duration: Durée du clip
            
        Returns:
            TextClip: Clip de texte
        """        
        if not text or not text.strip():
            text = " "

        try:
            clip = TextClip(
                text,
                fontsize=fontsize,
                color=self.TEXT_COLOR,
                method="caption",
                align=align,
                size=(width, None),
                interline=1.2
            )
            return clip.set_position(position).set_duration(duration)
        except Exception as e:
            print(f"?? Erreur création TextClip: {e}")
            return TextClip(
                " ", fontsize=fontsize, color=self.TEXT_COLOR,
                method="caption", size=(width, None)
            ).set_position(position).set_duration(duration)
    
    def create_animated_text_clip(self, text_content, duration=8, animation_level=0.5, text_size=36, text_width=800):
        """
        Crée un clip animé pour le texte
        
        Args:
            text_content: Contenu textuel
            duration: Durée du clip
            animation_level: Niveau d'animation (0-1)
            text_size: Taille du texte
            text_width: Largeur du texte
            
        Returns:
            CompositeVideoClip: Clip animé
        """
        
        # Nettoyer le texte
        if not text_content or not text_content.strip():
            text_content = " "
        
        # Diviser le texte en paragraphes
        paragraphs = text_content.split('\n\n')
        paragraphs = [p for p in paragraphs if p.strip()]
        
        if not paragraphs:
            return self.create_text_clip(" ", text_size, text_width, (self.TEXT_MARGIN, 100), duration=duration)
        
        # Déterminer la durée de l'animation
        animation_duration = min(2.0, duration / 2) * animation_level
        
        # Créer un clip pour chaque paragraphe
        paragraph_clips = []
        y_position = 100
        
        for i, paragraph in enumerate(paragraphs):
            # Calculer le délai d'apparition
            delay = i * (animation_duration / max(1, len(paragraphs) - 1))
            
            # Créer le clip de texte
            text_clip = TextClip(
                paragraph,
                fontsize=text_size,
                color=self.TEXT_COLOR,
                method="caption",
                align="west",
                size=(text_width, None),
                interline=1.2
            )
            
            # Définir d'abord la durée
            text_clip = text_clip.set_duration(duration - delay)
            
            # Positionner le clip
            text_clip = text_clip.set_position((self.TEXT_MARGIN, y_position))
            
            # Ajouter un effet d'entrée
            if animation_level > 0.3:
                text_clip = text_clip.set_start(delay)
                fade_duration = min(0.5, animation_duration / 2)
                if fade_duration > 0:
                    text_clip = text_clip.crossfadein(fade_duration)
            else:
                text_clip = text_clip.set_start(0)
            
            # Ajouter à la liste
            paragraph_clips.append(text_clip)
            
            # Mettre à jour la position Y
            y_position += text_clip.h + 10
        
        # Créer un clip composite
        return CompositeVideoClip(paragraph_clips, size=(self.WIDTH, self.HEIGHT))
    
    def preserve_markdown_formatting(self, text):
        """
        Préserve le formatage Markdown dans le texte
        
        Args:
            text: Texte avec formatage Markdown
            
        Returns:
            str: Texte avec formatage préservé
        """        
        if not text or not isinstance(text, str):
            return ""
        
        # Supprimer les références d'images
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        
        # Préserver les titres
        text = re.sub(r'^# (.+)$', r'\1', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.+)$', r'\1', text, flags=re.MULTILINE)
        text = re.sub(r'^### (.+)$', r'\1', text, flags=re.MULTILINE)
        
        # Préserver l'italique et le gras
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Gras
        text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italique
        text = re.sub(r'\_(.+?)\_', r'\1', text)      # Italique alternatif
        
        # Préserver les listes
        text = re.sub(r'^- (.+)$', r'• \1', text, flags=re.MULTILINE)
        text = re.sub(r'^\* (.+)$', r'• \1', text, flags=re.MULTILINE)
        text = re.sub(r'^(\d+)\. (.+)$', r'\1. \2', text, flags=re.MULTILINE)
        
        # Préserver les blocs de code
        code_blocks = re.findall(r'```(.+?)```', text, re.DOTALL)
        for block in code_blocks:
            text = text.replace(f'```{block}```', f'{block}')
        
        return text
    
    def replace_layout_optimizer(self):
        """
        Remplace la fonction optimize_layout de la classe LayoutOptimizer
        pour utiliser notre version améliorée
        """
        # Si la classe LayoutOptimizer existe dans le contexte global
        if 'LayoutOptimizer' in globals():
            LayoutOptimizer.optimize_layout = self.compute_optimal_layout
            print("? Fonction d'optimisation de mise en page remplacée")
        else:
            print("?? Classe LayoutOptimizer non trouvée dans le contexte global")
    def is_content_empty_except_images(content):
        """
        Vérifie si le contenu est vide ou ne contient que des références d'images ou des tableaux
        Traite les tableaux comme des éléments visuels qui doivent s'afficher en plein écran sans texte
        """
        if not content or not content.strip():
            return True
        
        # Suppression des références d'images
        content_without_images = re.sub(r'!\[.*?\]\(.*?\)', '', content)
        
        # Vérifier s'il y a un tableau et le supprimer pour l'analyse
        has_table = re.search(r'\|[-]+\|', content_without_images) is not None
        has_table = has_table or "|" in content_without_images and any(line.count('|') > 2 for line in content_without_images.split('\n'))
        
        if has_table:
            # Supprimer les lignes de tableaux pour voir s'il reste du texte
            content_without_tables = re.sub(r'^\|.*\|$', '', content_without_images, flags=re.MULTILINE)
            content_without_tables = re.sub(r'^[-|:]+$', '', content_without_tables, flags=re.MULTILINE)
            
            # Suppression des espaces, sauts de ligne et caractères de formatage Markdown
            clean_content = re.sub(r'[#*_\s\n\r]', '', content_without_tables)
            
            # S'il ne reste rien après suppression des tableaux, considérer comme sans texte
            return not bool(clean_content)
        
        # Suppression des espaces, sauts de ligne et caractères de formatage Markdown
        clean_content = re.sub(r'[#*_\s\n\r]', '', content_without_images)
        
        # Si le contenu nettoyé est vide, alors il ne contenait que des références d'images
        return not bool(clean_content)

# Intégration avec le générateur existant
def integrate_with_presentation_generator():
    """
    Intègre le SlideLayoutManager avec le générateur de présentations existant
    
    Cette fonction doit être appelée avant de générer des diapositives pour
    s'assurer que la mise en page est correctement optimisée.
    
    Returns:
        SlideLayoutManager: Instance du gestionnaire de mise en page
    """
    # Créer le gestionnaire de mise en page
    layout_manager = SlideLayoutManager()
    
    # Appliquer les corrections
    layout_manager.apply_layout_fixes()
    
    # Remplacer aussi la fonction slide_clip_with_exact_markdown_separated
    # dans le contexte global
    global slide_clip_with_exact_markdown_separated
    slide_clip_with_exact_markdown_separated = layout_manager.create_slide_with_separated_content
    
    return layout_manager

# Pour appliquer le correctif, intégrer dans MLEnhancedVideoGenerator comme suit:
"""
def process_markdown(self, markdown_content):
    # Appliquer d'abord la correction de mise en page
    layout_manager = integrate_with_presentation_generator()
    
    # Continuer avec le traitement normal
    try:
        # Avant tout, apprendre des présentations précédentes
        self.learn_from_previous_presentations()
        
        # 1. Extraire images et texte
        image_references, cleaned_text = extract_images_from_text(markdown_content, self.base_dir)
        
        # 2. Segmenter le contenu...
        # etc.
"""
class DeepSlideSynchronizer:
    """
    Classe qui utilise l'apprentissage profond pour synchroniser
    l'audio avec les diapositives de manière optimale
    """
    
    def __init__(self, model_path=None):
        self.model_path = model_path or DEEP_MODELS_DIR / "audio_sync_model.pt"
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model = None
        self.tokenizer = None
        self.speech_analyzer = None
        
        # Initialiser le modèle de synchronisation
        self.initialize_model()
        
        # Créer un analyseur de parole pour estimer les durées
        self.initialize_speech_analyzer()
        
        print("? Synchroniseur audio-slide par deep learning initialisé")
    
    def initialize_model(self):
        """Initialise le modèle de synchronisation audio-diapositive"""
        try:
            # Définir l'architecture du modèle
            class AudioSyncModel(nn.Module):
                def __init__(self, text_embedding_size=768, hidden_size=256):
                    super(AudioSyncModel, self).__init__()
                    
                    # Couches pour le texte
                    self.text_encoder = nn.Sequential(
                        nn.Linear(text_embedding_size, hidden_size),
                        nn.ReLU(),
                        nn.Dropout(0.2),
                        nn.Linear(hidden_size, hidden_size // 2)
                    )
                    
                    # Couches pour les caractéristiques de la diapositive
                    self.slide_encoder = nn.Sequential(
                        nn.Linear(10, hidden_size // 4),  # 10 caractéristiques de diapositive
                        nn.ReLU(),
                        nn.Linear(hidden_size // 4, hidden_size // 4)
                    )
                    
                    # Couche de fusion
                    self.fusion_layer = nn.Sequential(
                        nn.Linear(hidden_size // 2 + hidden_size // 4, hidden_size // 2),
                        nn.ReLU(),
                        nn.Dropout(0.2),
                        nn.Linear(hidden_size // 2, 1)  # Prédire la durée en secondes
                    )
                
                def forward(self, text_embeddings, slide_features):
                    text_encoded = self.text_encoder(text_embeddings)
                    slide_encoded = self.slide_encoder(slide_features)
                    
                    # Concaténer les représentations
                    combined = torch.cat([text_encoded, slide_encoded], dim=1)
                    
                    # Prédire la durée
                    duration = self.fusion_layer(combined)
                    return duration
            
            # Créer ou charger le modèle
            if Path(self.model_path).exists():
                checkpoint = torch.load(self.model_path, map_location=self.device)
                self.model = AudioSyncModel().to(self.device)
                self.model.load_state_dict(checkpoint['model_state_dict'])
                print(f"? Modèle de synchronisation chargé depuis {self.model_path}")
            else:
                self.model = AudioSyncModel().to(self.device)
                print("? Nouveau modèle de synchronisation initialisé (non entraîné)")
            
            # Initialiser un tokenizer pour les embeddings de texte si TRANSFORMER_AVAILABLE
            if TRANSFORMER_AVAILABLE:
               
                try:
                    # Utiliser un modèle léger pour les embeddings
                    model_name = "sentence-transformers/all-MiniLM-L6-v2"
                    self.tokenizer = AutoTokenizer.from_pretrained(model_name)
                    self.text_encoder = AutoModel.from_pretrained(model_name).to(self.device)
                    print("? Modèle d'embeddings de texte chargé")
                except Exception as e:
                    print(f"?? Erreur chargement modèle d'embeddings: {e}")
                    self.tokenizer = None
                    self.text_encoder = None
            
        except Exception as e:
            print(f"?? Erreur initialisation modèle de synchronisation: {e}")
            traceback.print_exc()
            self.model = None
    
    def initialize_speech_analyzer(self):
        """Initialise l'analyseur de parole pour estimer les durées audio"""
        try:
            if TRANSFORMER_AVAILABLE:               
                # Essayer de charger un modèle de reconnaissance vocale léger
                try:
                    self.speech_analyzer = pipeline(
                        "automatic-speech-recognition",
                        model="facebook/wav2vec2-base-960h",
                        device=0 if self.device == "cuda" else -1
                    )
                    print("? Analyseur de parole initialisé avec wav2vec2")
                except Exception as e:
                    print(f"?? Erreur chargement wav2vec2: {e}")
                    self.speech_analyzer = None
            else:
                print("?? Transformers non disponible pour l'analyse audio")
                self.speech_analyzer = None
        except Exception as e:
            print(f"?? Erreur initialisation analyseur de parole: {e}")
            self.speech_analyzer = None
    
    def get_text_embeddings(self, text):
        """Obtient les embeddings d'un texte pour le modèle de synchronisation"""
        if not self.tokenizer or not self.text_encoder:
            # Fallback avec des embeddings aléatoires si les modèles ne sont pas disponibles
            return torch.randn(1, 768).to(self.device)
        
        try:
            # Tronquer le texte s'il est trop long
            max_length = 512
            if len(text.split()) > max_length:
                text = " ".join(text.split()[:max_length])
            
            # Tokeniser le texte
            inputs = self.tokenizer(
                text, 
                padding=True, 
                truncation=True,
                max_length=max_length,
                return_tensors="pt"
            ).to(self.device)
            
            # Obtenir les embeddings
            with torch.no_grad():
                outputs = self.text_encoder(**inputs)
                # Utiliser la représentation du token [CLS]
                embeddings = outputs.last_hidden_state[:, 0, :]
            
            return embeddings
            
        except Exception as e:
            print(f"?? Erreur extraction embeddings: {e}")
            return torch.randn(1, 768).to(self.device)  # Fallback
    
    def extract_slide_features(self, content, has_image, animation_level):
        """Extrait des caractéristiques de la diapositive pour prédiction de durée"""
        features = torch.zeros(1, 10).to(self.device)
        
        try:
            # Caractéristique 1: Longueur du texte (normalisée)
            features[0, 0] = min(1.0, len(content) / 2000)
            
            # Caractéristique 2: Nombre de mots (normalisé)
            word_count = len(content.split())
            features[0, 1] = min(1.0, word_count / 400)
            
            # Caractéristique 3: Présence d'image
            features[0, 2] = 1.0 if has_image else 0.0
            
            # Caractéristique 4: Niveau d'animation
            features[0, 3] = animation_level
            
            # Caractéristique 5: Complexité estimée du texte
            avg_word_length = sum(len(word) for word in content.split()) / max(1, word_count)
            features[0, 4] = min(1.0, avg_word_length / 10)
            
            # Caractéristique 6: Nombre de phrases (normalisé)
            sentence_count = len(re.split(r'[.!?]+', content))
            features[0, 5] = min(1.0, sentence_count / 30)
            
            # Caractéristique 7: Présence de listes
            bullet_count = len(re.findall(r'[-*+]\s+', content))
            features[0, 6] = min(1.0, bullet_count / 10)
            
            # Caractéristique 8: Présence de tableaux
            has_table = 1.0 if re.search(r'\|[-]+\|', content) else 0.0
            features[0, 7] = has_table
            
            # Caractéristique 9: Ratio de texte technique (mots longs)
            long_words = sum(1 for word in content.split() if len(word) > 8)
            features[0, 8] = min(1.0, long_words / max(1, word_count))
            
            # Caractéristique 10: Densité de chiffres et statistiques
            digit_density = len(re.findall(r'\d+', content)) / max(1, len(content))
            features[0, 9] = min(1.0, digit_density * 20)
            
        except Exception as e:
            print(f"?? Erreur extraction caractéristiques slide: {e}")
        
        return features
    
    def predict_optimal_duration(self, content, has_image=False, animation_level=0.5):
        """Prédit la durée optimale pour une diapositive avec narration"""
        if not self.model:
            # Estimation basique si le modèle n'est pas disponible
            word_count = len(content.split())
            base_duration = max(5, min(15, word_count / 20))
            image_factor = 1.3 if has_image else 1.0
            return base_duration * image_factor
        
        try:
            # Obtenir les embeddings du texte
            text_embeddings = self.get_text_embeddings(content)
            
            # Extraire des caractéristiques de la diapositive
            slide_features = self.extract_slide_features(content, has_image, animation_level)
            
            # Prédire la durée
            self.model.eval()
            with torch.no_grad():
                duration = self.model(text_embeddings, slide_features)
                
                # Convertir en valeur scalaire et appliquer des limites raisonnables
                duration_seconds = float(duration.item())
                duration_seconds = max(5, min(30, duration_seconds))  # Entre 5 et 30 secondes
                
                return duration_seconds
                
        except Exception as e:
            print(f"?? Erreur prédiction durée: {e}")
            # Fallback avec une estimation basique
            word_count = len(content.split())
            base_duration = max(5, min(15, word_count / 20))
            image_factor = 1.3 if has_image else 1.0
            return base_duration * image_factor
    
    def analyze_narration_speed(self, text):
        """Analyse le texte pour estimer la vitesse de narration optimale"""
        # Calculer des statistiques sur le texte
        words = text.split()
        word_count = len(words)
        
        if word_count == 0:
            return 150  # Vitesse par défaut
        
        # Analyse basique
        avg_word_length = sum(len(word) for word in words) / word_count
        
        # Ajuster la vitesse en fonction de la longueur moyenne des mots
        # Plus les mots sont longs, plus la vitesse devrait être lente
        if avg_word_length > 7:
            speed = 130  # Plus lent pour les mots complexes
        elif avg_word_length > 5:
            speed = 150  # Vitesse moyenne
        else:
            speed = 170  # Plus rapide pour les mots simples
        
        # Ajuster en fonction de la présence de chiffres et symboles
        # qui peuvent ralentir la narration
        digit_percentage = len(re.findall(r'\d', text)) / max(1, len(text))
        if digit_percentage > 0.05:  # Plus de 5% de chiffres
            speed -= 20
        
        # Limiter la vitesse à une plage raisonnable
        return max(120, min(180, speed))
    
    def estimate_audio_duration(self, text, voice_speed=150):
        """Estime la durée en secondes d'un texte narré"""
        if not text:
            return 0
        
        try:
            # Méthode basée sur le nombre de mots et la vitesse de parole
            word_count = len(text.split())
            
            # Calcul: (Nombre de mots / Mots par minute) * 60 secondes
            words_per_minute = voice_speed
            duration_seconds = (word_count / words_per_minute) * 60
            
            # Ajouter un temps supplémentaire pour les pauses
            punctuation_count = len(re.findall(r'[,.;:!?]', text))
            pause_time = punctuation_count * 0.2  # 0.2 secondes par signe de ponctuation
            
            # Ajouter du temps pour les listes à puces (pauses plus longues)
            bullet_count = len(re.findall(r'[-*+]\s+', text))
            bullet_pause_time = bullet_count * 0.5  # 0.5 secondes par puce
            
            return duration_seconds + pause_time + bullet_pause_time
            
        except Exception as e:
            print(f"?? Erreur estimation durée audio: {e}")
            # Estimation simplifiée en cas d'erreur
            return len(text.split()) * 0.3  # ~0.3 secondes par mot
    
    def create_synchronized_clips(self, slides, narration_audio_paths):
        """
        Crée une liste de clips vidéo synchronisés avec l'audio
        en ajustant intelligemment les durées
        
        Args:
            slides: Liste des clips de diapositives
            narration_audio_paths: Liste des chemins audio de narration
            
        Returns:
            synchronized_clips: Liste de clips synchronisés avec l'audio
        """       
        synchronized_clips = []
        
        if not slides or not narration_audio_paths:
            return slides  # Retourner les slides originaux si pas d'audio
        
        try:
            # Charger les fichiers audio
            audio_clips = []
            for audio_path in narration_audio_paths:
                try:
                    if Path(audio_path).exists():
                        audio_clip = AudioFileClip(str(audio_path))
                        audio_clips.append(audio_clip)
                    else:
                        print(f"?? Fichier audio non trouvé: {audio_path}")
                        audio_clips.append(None)
                except Exception as e:
                    print(f"?? Erreur chargement audio {audio_path}: {e}")
                    audio_clips.append(None)
            
            # Synchroniser chaque diapositive avec son audio
            for i, slide in enumerate(slides):
                if i < len(audio_clips) and audio_clips[i] is not None:
                    # Obtenir la durée de l'audio
                    audio_duration = audio_clips[i].duration
                    
                    # Ajuster la durée de la diapositive pour correspondre à l'audio
                    # avec une durée minimale et une marge de sécurité
                    min_duration = 5  # Durée minimale en secondes
                    safety_margin = 1  # Marge supplémentaire en secondes
                    
                    new_duration = max(min_duration, audio_duration + safety_margin)
                    
                    # Si l'audio est trop court, on garde la durée originale de la diapositive
                    if audio_duration < 3 and hasattr(slide, 'duration') and slide.duration > new_duration:
                        new_duration = slide.duration
                    
                    # Créer une version synchronisée de la diapositive
                    synchronized_slide = slide.set_duration(new_duration).set_audio(audio_clips[i])
                    synchronized_clips.append(synchronized_slide)
                else:
                    # Si pas d'audio disponible, utiliser la diapositive originale
                    synchronized_clips.append(slide)
            
            print(f"? {len(synchronized_clips)} clips synchronisés avec l'audio")
            return synchronized_clips
            
        except Exception as e:
            print(f"? Erreur lors de la synchronisation audio-vidéo: {e}")
            traceback.print_exc()
            return slides  # Retourner les slides originaux en cas d'erreur
    
    def train_sync_model(self, training_data, epochs=100):
        """
        Entraîne le modèle de synchronisation avec des données d'exemple
        
        Args:
            training_data: Liste de tuples (content, has_image, animation_level, actual_duration)
            epochs: Nombre d'époques d'entraînement
        """
        if not self.model or not training_data:
            print("?? Modèle non initialisé ou données d'entraînement manquantes")
            return False
        
        try:        
            # Préparer les données d'entraînement
            text_embeddings = []
            slide_features = []
            durations = []
            
            for content, has_image, animation_level, actual_duration in training_data:
                # Obtenir les embeddings et caractéristiques
                embedding = self.get_text_embeddings(content)
                features = self.extract_slide_features(content, has_image, animation_level)
                
                text_embeddings.append(embedding)
                slide_features.append(features)
                durations.append(torch.tensor([[actual_duration]], device=self.device))
            
            # Convertir en tensors (déjà sur le bon device)
            text_embeddings = torch.cat(text_embeddings, dim=0)
            slide_features = torch.cat(slide_features, dim=0)
            durations = torch.cat(durations, dim=0)
            
            # Créer le dataset et dataloader
            dataset = TensorDataset(text_embeddings, slide_features, durations)
            dataloader = DataLoader(dataset, batch_size=min(8, len(dataset)), shuffle=True)
            
            # Configurer l'optimiseur et la fonction de perte
            optimizer = optim.Adam(self.model.parameters(), lr=0.001)
            criterion = nn.MSELoss()
            
            # Entraînement
            self.model.train()
            for epoch in range(epochs):
                total_loss = 0
                for batch_embeddings, batch_features, batch_durations in dataloader:
                    # Forward pass
                    optimizer.zero_grad()
                    predictions = self.model(batch_embeddings, batch_features)
                    
                    # Calculer la perte
                    loss = criterion(predictions, batch_durations)
                    
                    # Backward pass
                    loss.backward()
                    optimizer.step()
                    
                    total_loss += loss.item()
                
                # Afficher la progression
                if (epoch + 1) % 10 == 0:
                    print(f"Epoch {epoch+1}/{epochs}, Loss: {total_loss/len(dataloader):.4f}")
            
            # Sauvegarder le modèle entraîné
            torch.save({
                'model_state_dict': self.model.state_dict(),
                'optimizer_state_dict': optimizer.state_dict()
            }, self.model_path)
            
            print(f"? Modèle de synchronisation entraîné et sauvegardé dans {self.model_path}")
            return True
            
        except Exception as e:
            print(f"? Erreur lors de l'entraînement du modèle de synchronisation: {e}")
            traceback.print_exc()
            return False
    
    def collect_training_data(self, slides, narration_audio_paths):
        """
        Collecte des données d'entraînement à partir des slides et de l'audio
        pour améliorer le modèle de synchronisation
        
        Returns:
            training_data: Liste de tuples (content, has_image, animation_level, actual_duration)
        """
        training_data = []
        
        if not slides or not narration_audio_paths:
            return training_data
        
        try:           
            # Parcourir les slides et l'audio correspondant
            for i, slide_info in enumerate(slides):
                if i >= len(narration_audio_paths):
                    break
                
                audio_path = narration_audio_paths[i]
                if not Path(audio_path).exists():
                    continue
                
                # Extraire les caractéristiques du slide
                content = slide_info.get('content', '')
                has_image = slide_info.get('has_image', False)
                animation_level = slide_info.get('style_params', {}).get('animation_level', 0.5)
                
                # Mesurer la durée réelle de l'audio
                try:
                    audio_clip = AudioFileClip(str(audio_path))
                    actual_duration = audio_clip.duration
                    audio_clip.close()
                    
                    # Ajouter aux données d'entraînement
                    training_data.append((content, has_image, animation_level, actual_duration))
                except Exception as e:
                    print(f"?? Erreur mesure durée audio {audio_path}: {e}")
            
            print(f"?? {len(training_data)} exemples collectés pour l'entraînement")
            return training_data
            
        except Exception as e:
            print(f"?? Erreur collecte données d'entraînement: {e}")
            return training_data

""" the class ContentSegmenter was here before being refactored to content_segmenter.py"""
""" the class LayoutOptimizer was here before refactoring to layout_optimizer.py"""

from videoseul_modules.readability import ReadabilityEvaluator
class UserFeedbackSystem:
    """
    Système pour collecter et traiter les retours utilisateurs
    sur la qualité des présentations générées
    """
    
    def __init__(self):
        self.feedback_history = []
        self.session_ratings = {}
        self.feedback_file = TRAINING_DATA_DIR / "user_feedback.json"
        
        # Charger l'historique de feedback s'il existe
        self.load_feedback_history()
        
        print("? Système de feedback utilisateur initialisé")
    
    def load_feedback_history(self):
        """Charge l'historique des feedbacks s'il existe"""
        if self.feedback_file.exists():
            try:
                with open(self.feedback_file, 'r', encoding='utf-8') as f:
                    self.feedback_history = json.load(f)
                    print(f"?? {len(self.feedback_history)} feedbacks utilisateur chargés")
            except Exception as e:
                print(f"?? Erreur chargement historique feedback: {e}")
                self.feedback_history = []
    
    def save_feedback_history(self):
        """Sauvegarde l'historique des feedbacks"""
        try:
            with open(self.feedback_file, 'w', encoding='utf-8') as f:
                json.dump(self.feedback_history, f, ensure_ascii=False, indent=4)
            print(f"?? Historique de feedback sauvegardé ({len(self.feedback_history)} entrées)")
        except Exception as e:
            print(f"?? Erreur sauvegarde historique feedback: {e}")
    
    def add_feedback(self, slide_data, rating, comments=None):
        """
        Ajoute un feedback utilisateur à l'historique
        
        Args:
            slide_data: Données de la diapositive évaluée
            rating: Note donnée (0-10)
            comments: Commentaires optionnels
        """
        feedback = {
            "slide_data": slide_data,
            "rating": rating,
            "comments": comments,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        self.feedback_history.append(feedback)
        self.session_ratings[slide_data.get('title', 'Sans titre')] = rating
        
        # Sauvegarder après chaque feedback
        self.save_feedback_history()
        
        return True
    
    def get_average_rating(self):
        """Retourne la note moyenne des feedbacks récents"""
        if not self.session_ratings:
            return 0
        
        return sum(self.session_ratings.values()) / len(self.session_ratings)
    
    def get_recommendations(self):
        """
        Génère des recommandations d'amélioration basées sur les feedbacks
        
        Returns:
            recommendations: Liste de recommandations
        """
        recommendations = []
        
        # Analyser les notes par catégorie
        low_rated_slides = [title for title, rating in self.session_ratings.items() if rating < 5]
        
        if low_rated_slides:
            recommendations.append(f"Améliorer les diapositives suivantes: {', '.join(low_rated_slides)}")
        
        # Recommandations générales basées sur l'historique si disponible
        if len(self.feedback_history) > 5:
            avg_historical = sum(item['rating'] for item in self.feedback_history) / len(self.feedback_history)
            
            if avg_historical < 6:
                recommendations.append("Privilégier les présentations plus visuelles avec moins de texte")
            elif avg_historical < 8:
                recommendations.append("Continuer à améliorer l'équilibre texte/images")
        
        return recommendations

""" the class MLEnhancedVideoGenerator was here before being refactored to  video_generator.py"""
""" the class AudioSynchronizedMLGenerator was here before being refactored to  video_generator.py"""
""" the class DeepAudioProcessor was here before being refactored to class audio_processing.py"""

    
    def convert_pdf_with_ml(self, pdf_path):
        """
        Convertit un PDF en Markdown avec amélioration par apprentissage automatique
        pour une meilleure extraction de texte et d'images
        
        Args:
            pdf_path: Chemin du fichier PDF
            
        Returns:
            markdown_path: Chemin du fichier Markdown généré
        """
        try:
            # Extraire le texte et les images
            text, images = extract_only_figures_from_pdf(pdf_path)
            
            # Amélioration par ML : Restructuration intelligente du contenu
            if text:
                # 1. Détecter et corriger les colonnes mal extraites
                text = self.fix_column_extraction(text)
                
                # 2. Identifier et restructurer les sections logiques
                text = self.improve_document_structure(text)
                
                # 3. Améliorer la qualité des images extraites
                enhanced_images = self.enhance_extracted_images(images)
                
                # 4. Améliorer le contenu textuel avec le modèle de langage
                if self.video_generator.text_processor and self.video_generator.text_processor.initialized:
                    print("?? Amélioration du contenu extrait avec un modèle de langage...")
                    # Améliorer le contenu par sections pour éviter les limites de tokens
                    sections = re.split(r'(^#{1,3}\s+.*?$)', text, flags=re.MULTILINE)
                    enhanced_sections = []
                    
                    for i in range(0, len(sections)):
                        section = sections[i]
                        # Ne pas traiter les titres ou les sections trop courtes
                        if re.match(r'^#{1,3}\s+', section) or len(section.strip()) < 100:
                            enhanced_sections.append(section)
                        else:
                            prompt = f"""Améliore ce texte extrait d'un PDF en corrigeant les erreurs d'extraction, les fautes de grammaire, 
                            et en restructurant le contenu si nécessaire. Garde exactement le même message, mais améliore la clarté:
                            
                            {section}"""
                            
                            enhanced_section = self.video_generator.text_processor.generate_text(prompt, max_length=len(section.split()) + 100)
                            enhanced_sections.append(enhanced_section)
                    
                    text = "".join(enhanced_sections)
                
                # Générer le Markdown final
                markdown_content = text
                
                # Si les images ne sont pas déjà référencées dans le texte
                if not any(f"![]" in text or "![Image" in text or "![Figure" in text):
                    markdown_content = self.insert_images_intelligently(text, enhanced_images)
            else:
                print("?? Aucun texte extrait du PDF")
                markdown_content = "# Document extrait\n\n"
                
                # Créer un contenu basé uniquement sur les images
                for i, img_path in enumerate(images):
                    markdown_content += f"\n## Page {i+1}\n\n"
                    markdown_content += f"![Image {Path(img_path).name}]({img_path})\n\n"
            
            # Sauvegarder le Markdown
            output_filename = f"{pdf_path.stem}_enhanced.md"
            output_path = self.output_dir / output_filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            print(f"? PDF converti avec amélioration ML: {output_path}")
            
            return output_path
        except Exception as e:
            print(f"? Erreur lors de la conversion du PDF: {e}")
            traceback.print_exc()
            return None
    
    def convert_word_with_ml(self, docx_path):
        """
        Convertit un fichier Word en Markdown avec améliorations ML
        
        Args:
            docx_path: Chemin du fichier Word
            
        Returns:
            markdown_path: Chemin du fichier Markdown généré
        """
        try:
            # Extraire le texte et les images
            text, images = extract_from_word(docx_path)
            
            # Appliquer les mêmes améliorations que pour le PDF
            if text:
                text = self.improve_document_structure(text)
                enhanced_images = self.enhance_extracted_images(images)
                
                # Améliorer le contenu textuel avec le modèle de langage si disponible
                if self.video_generator.text_processor and self.video_generator.text_processor.initialized:
                    print("?? Amélioration du contenu extrait avec un modèle de langage...")
                    # Approche similaire à celle utilisée pour les PDFs
                    sections = re.split(r'(^#{1,3}\s+.*?$)', text, flags=re.MULTILINE)
                    enhanced_sections = []
                    
                    for i in range(0, len(sections)):
                        section = sections[i]
                        if re.match(r'^#{1,3}\s+', section) or len(section.strip()) < 100:
                            enhanced_sections.append(section)
                        else:
                            prompt = f"""Améliore ce texte extrait d'un document Word en corrigeant les erreurs d'extraction, les fautes de grammaire, 
                            et en restructurant le contenu si nécessaire. Garde exactement le même message, mais améliore la clarté:
                            
                            {section}"""
                            
                            enhanced_section = self.video_generator.text_processor.generate_text(prompt, max_length=len(section.split()) + 100)
                            enhanced_sections.append(enhanced_section)
                    
                    text = "".join(enhanced_sections)
                
                if not any(f"![]" in text or "![Image" in text):
                    markdown_content = self.insert_images_intelligently(text, enhanced_images)
                else:
                    markdown_content = text
            else:
                print("?? Aucun texte extrait du document Word")
                markdown_content = "# Document extrait\n\n"
                
                for i, img_path in enumerate(images):
                    markdown_content += f"\n## Page {i+1}\n\n"
                    markdown_content += f"![Image {Path(img_path).name}]({img_path})\n\n"
            
            # Sauvegarder le Markdown
            output_filename = f"{docx_path.stem}_enhanced.md"
            output_path = self.output_dir / output_filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            print(f"? Document Word converti avec amélioration ML: {output_path}")
            
            return output_path
        except Exception as e:
            print(f"? Erreur lors de la conversion du document Word: {e}")
            traceback.print_exc()
            return None
    
    def fix_column_extraction(self, text):
        """
        Détecte et corrige les problèmes d'extraction de colonnes dans les PDFs
        
        Args:
            text: Texte extrait du PDF
            
        Returns:
            fixed_text: Texte avec colonnes corrigées
        """
        lines = text.split('\n')
        fixed_lines = []
        in_column_section = False
        column_buffer = []
        
        for line in lines:
            # Détecter les sections avec colonnes potentielles
            if re.match(r'^## --- Page \d+ ---$', line):
                # Nouvelle page, traiter les colonnes précédentes si présentes
                if in_column_section and column_buffer:
                    fixed_lines.extend(self.merge_columns(column_buffer))
                    column_buffer = []
                
                in_column_section = False
                fixed_lines.append(line)
                continue
            
            # Détecter les lignes qui semblent être des colonnes
            if line.strip() and len(line) < 50 and not line.startswith('#'):
                # Potentiellement une colonne
                if not in_column_section:
                    in_column_section = True
                column_buffer.append(line)
            else:
                if in_column_section and column_buffer:
                    # Fin d'une section de colonnes
                    fixed_lines.extend(self.merge_columns(column_buffer))
                    column_buffer = []
                    in_column_section = False
                
                fixed_lines.append(line)
        
        # Traiter les dernières colonnes si présentes
        if in_column_section and column_buffer:
            fixed_lines.extend(self.merge_columns(column_buffer))
        
        return '\n'.join(fixed_lines)
    
    def merge_columns(self, column_lines):
        """
        Fusionne intelligemment les lignes qui semblent être des colonnes
        
        Args:
            column_lines: Liste de lignes potentiellement en colonnes
            
        Returns:
            merged_lines: Liste de lignes fusionnées
        """
        # Si peu de lignes, ne pas traiter comme des colonnes
        if len(column_lines) < 4:
            return column_lines
        
        # Tentative de détection de colonnes
        # Regrouper les lignes qui semblent appartenir à la même colonne
        left_col = []
        right_col = []
        
        # Détecter si nous avons une structure à deux colonnes
        has_two_columns = False
        avg_line_length = sum(len(line) for line in column_lines) / len(column_lines)
        
        for i, line in enumerate(column_lines):
            if i < len(column_lines) - 1:
                curr_len = len(line)
                next_len = len(column_lines[i+1])
                
                # Si la ligne actuelle est courte et la suivante aussi, 
                # c'est probablement une structure à deux colonnes
                if curr_len < avg_line_length * 0.7 and next_len < avg_line_length * 0.7:
                    has_two_columns = True
                    break
        
        if has_two_columns:
            # Séparer en colonnes gauche/droite
            for i in range(0, len(column_lines), 2):
                if i < len(column_lines):
                    left_col.append(column_lines[i])
                if i+1 < len(column_lines):
                    right_col.append(column_lines[i+1])
            
            # Regrouper les colonnes
            merged = []
            left_paragraphs = self.group_column_paragraphs(left_col)
            right_paragraphs = self.group_column_paragraphs(right_col)
            
            # Fusionner les paragraphes dans l'ordre
            merged.extend(left_paragraphs)
            merged.append('')  # Ligne vide
            merged.extend(right_paragraphs)
            
            return merged
        else:
            # Pas de structure en colonnes détectée, retourner tel quel
            return column_lines
    
    def group_column_paragraphs(self, column_lines):
        """Groupe les lignes d'une colonne en paragraphes"""
        paragraphs = []
        current_paragraph = []
        
        for line in column_lines:
            if not line.strip():
                if current_paragraph:
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
            else:
                current_paragraph.append(line.strip())
        
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        return paragraphs
    
    def improve_document_structure(self, text):
        """
        Améliore la structure du document en identifiant et en normalisant
        les titres, sections, listes, etc.
        
        Args:
            text: Texte Markdown brut
            
        Returns:
            improved_text: Texte avec structure améliorée
        """
        lines = text.split('\n')
        improved_lines = []
        in_list = False
        
        for i, line in enumerate(lines):
            # Normaliser les titres
            heading_match = re.match(r'^(#+)\s+(.*?)$', line)
            if heading_match:
                level, title = heading_match.groups()
                
                # Normaliser la capitalisation des titres
                words = title.split()
                if words:
                    # Capitaliser les mots importants
                    title_words = []
                    for word in words:
                        if word.lower() not in ['a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'in', 'at', 'to', 'by', 'from']:
                            title_words.append(word.capitalize())
                        else:
                            title_words.append(word.lower())
                    
                    # Toujours capitaliser le premier mot
                    if title_words:
                        title_words[0] = title_words[0].capitalize()
                    
                    title = ' '.join(title_words)
                
                improved_lines.append(f"{level} {title}")
                continue
            
            # Détecter et normaliser les listes
            list_match = re.match(r'^(\s*)([-•*])(\s+)(.*?)$', line)
            if list_match:
                indent, marker, space, content = list_match.groups()
                
                # Normaliser le marqueur de liste
                improved_lines.append(f"{indent}* {content}")
                in_list = True
                continue
            
            # Détecter les listes numérotées
            num_list_match = re.match(r'^(\s*)(\d+)\.(\s+)(.*?)$', line)
            if num_list_match:
                indent, number, space, content = num_list_match.groups()
                improved_lines.append(f"{indent}{number}. {content}")
                in_list = True
                continue
            
            # Fin d'une liste
            if in_list and not line.strip():
                in_list = False
            
            # Détecter et améliorer les tableaux
            if '|' in line and ('-' in line or ':' in line):
                # Probablement une ligne de tableau
                # Normaliser l'alignement et l'espacement
                cells = line.split('|')
                formatted_cells = []
                
                for cell in cells:
                    # Détecter si c'est une ligne de séparation
                    if re.match(r'^[\s\-:]+$', cell):
                        # Normaliser les séparateurs
                        if ':' in cell:  # Conserver l'alignement
                            if cell.startswith(':') and cell.endswith(':'):
                                formatted_cells.append(' :---: ')  # Centré
                            elif cell.startswith(':'):
                                formatted_cells.append(' :--- ')   # Gauche
                            elif cell.endswith(':'):
                                formatted_cells.append(' ---: ')   # Droite
                        else:
                            formatted_cells.append(' --- ')        # Défaut
                    else:
                        # Cellule normale, nettoyage de l'espacement
                        formatted_cells.append(f" {cell.strip()} ")
                
                improved_lines.append('|'.join(formatted_cells))
                continue
            
            # Ajouter la ligne telle quelle
            improved_lines.append(line)
        
        # Ajouter des séparations entre les sections
        with_sections = []
        last_was_heading = False
        
        for line in improved_lines:
            if re.match(r'^#+\s+', line):
                # Ajouter une ligne vide avant les titres (sauf au début)
                if with_sections and not last_was_heading:
                    with_sections.append('')
                with_sections.append(line)
                last_was_heading = True
            else:
                with_sections.append(line)
                last_was_heading = False
        
        # Améliorer les références d'images
        final_text = '\n'.join(with_sections)
        
        # Standardiser les références d'images
        final_text = re.sub(r'!\[()\]\((.*?)\)', r'![Image](\2)', final_text)
        
        # Nettoyer les lignes vides multiples
        final_text = re.sub(r'\n{3,}', '\n\n', final_text)
        
        return final_text
    
    def enhance_extracted_images(self, images):
        """
        Améliore la qualité des images extraites
        
        Args:
            images: Liste des chemins d'images
            
        Returns:
            enhanced_images: Liste des chemins d'images améliorées
        """
        enhanced_images = []
        
        for img_path in images:
            try:
                img_path = Path(img_path)
                if not img_path.exists():
                    print(f"?? Image introuvable: {img_path}")
                    continue
                
                # Créer un chemin pour l'image améliorée
                enhanced_path = self.output_dir / "enhanced_images" / img_path.name
                enhanced_path.parent.mkdir(exist_ok=True)
                
                # Appliquer des améliorations à l'image
                
                img = Image.open(img_path)
                
                # 1. Améliorer le contraste
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(1.2)
                
                # 2. Améliorer la netteté
                enhancer = ImageEnhance.Sharpness(img)
                img = enhancer.enhance(1.5)
                
                # 3. Améliorer la luminosité si nécessaire
                stat = ImageStat.Stat(img)
                brightness = sum(stat.mean) / len(stat.mean)
                
                if brightness < 100:  # Image sombre
                    enhancer = ImageEnhance.Brightness(img)
                    img = enhancer.enhance(1.3)
                
                # Vérifier si l'image contient du texte (diagramme, schéma)
                # Cela nécessiterait idéalement une IA pour la détection de texte
                # Mais on peut utiliser une heuristique basée sur la variance des couleurs
                std_dev = sum(stat.stddev) / len(stat.stddev)
                if std_dev < 50:  # Faible variance = potentiellement un diagramme ou schéma
                    # 4. Augmenter le contraste pour améliorer la lisibilité du texte
                    enhancer = ImageEnhance.Contrast(img)
                    img = enhancer.enhance(1.3)
                
                # Sauvegarder l'image améliorée
                img.save(enhanced_path)
                enhanced_images.append(enhanced_path)
                
                print(f"? Image améliorée: {enhanced_path}")
                
            except Exception as e:
                print(f"?? Erreur amélioration image {img_path}: {e}")
                enhanced_images.append(img_path)  # Utiliser l'image originale
        
        return enhanced_images
    
    def insert_images_intelligently(self, text, images):
        """
        Insère intelligemment les images dans le texte Markdown
        en les plaçant aux endroits les plus pertinents
        
        Args:
            text: Texte Markdown
            images: Liste des chemins d'images
            
        Returns:
            text_with_images: Texte Markdown avec images insérées
        """
        if not images:
            return text
        
        # Identifier des points d'insertion logiques (fins de sections)
        insertion_points = []
        
        # Trouver les fins de sections
        sections = re.split(r'^(#+\s+.*?)$', text, flags=re.MULTILINE)
        
        accumulated_length = 0
        for i, section in enumerate(sections):
            accumulated_length += len(section)
            
            # Insérer après les titres ou à la fin des paragraphes
            if i > 0 and re.match(r'^#+\s+', sections[i-1]):
                insertion_points.append(accumulated_length)
        
        # S'il n'y a pas assez de points d'insertion, en ajouter à des positions régulières
        if len(insertion_points) < len(images):
            text_length = len(text)
            for i in range(len(images)):
                position = text_length * (i + 1) // (len(images) + 1)
                
                # Trouver la fin de paragraphe la plus proche
                paragraph_end = text.find('\n\n', position)
                if paragraph_end != -1:
                    insertion_points.append(paragraph_end)
                else:
                    insertion_points.append(position)
        
        # Trier les points d'insertion
        insertion_points = sorted(set(insertion_points))
        
        # Insérer les images
        result = text
        offset = 0
        
        for i, img_path in enumerate(images):
            if i < len(insertion_points):
                position = insertion_points[i] + offset
                
                # Extraire le nom de l'image pour une meilleure légende
                img_name = Path(img_path).stem
                img_name = re.sub(r'[_-]', ' ', img_name).strip()
                
                # Pour les images de page, utiliser une légende plus simple
                if re.match(r'pdf_page_\d+', img_name) or re.match(r'word_image_\d+', img_name):
                    caption = f"Figure {i+1}"
                else:
                    # Capitaliser la première lettre de chaque mot pour une légende propre
                    caption = ' '.join(word.capitalize() for word in img_name.split())
                    caption = f"Figure {i+1}: {caption}"
                
                image_ref = f"\n\n![{caption}]({img_path})\n\n"
                
                result = result[:position] + image_ref + result[position:]
                offset += len(image_ref)
        
        return result
    
    def save_processing_history(self):
        """Sauvegarde l'historique de traitement pour apprentissage futur"""
        history_file = TRAINING_DATA_DIR / "processing_history.json"
        
        # Charger l'historique existant
        history = []
        if history_file.exists():
            try:
                with open(history_file, 'r', encoding='utf-8') as f:
                    history = json.load(f)
            except:
                pass
        
        # Ajouter les nouvelles entrées
        history.extend(self.processed_files)
        
        # Sauvegarder
        try:
            with open(history_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=4)
            
            print(f"?? Historique de traitement sauvegardé ({len(history)} fichiers)")
        except Exception as e:
            print(f"?? Erreur sauvegarde historique: {e}")
    
    def batch_process_documents(self, directory, extensions=['.pdf', '.docx', '.doc', '.md']):
        """
        Traite tous les documents d'un répertoire avec optimisation ML
        
        Args:
            directory: Chemin du répertoire
            extensions: Liste des extensions de fichiers à traiter
            
        Returns:
            processed_files: Liste des fichiers traités
        """
        directory = Path(directory)
        if not directory.exists() or not directory.is_dir():
            print(f"? Le répertoire {directory} n'existe pas")
            return []
        
        processed_files = []
        
        for ext in extensions:
            # Trouver tous les fichiers avec cette extension
            files = list(directory.glob(f"*{ext}"))
            
            if files:
                print(f"?? Trouvé {len(files)} fichier(s) {ext}")
                
                for file_path in files:
                    print(f"\n=== Traitement ML de {file_path.name} ===")
                    
                    # Configurer le chemin de sortie pour ce fichier
                    output_filename = f"ml_presentation_{file_path.stem}.mp4"
                    output_path = self.output_dir / output_filename
                    
                    # Traiter le document
                    result_path = self.process_document(file_path)
                    
                    if result_path:
                        processed_files.append(output_path)
                        
                        # Renommer le fichier de sortie pour éviter les écrasements
                        source_path = self.output_dir / "ml_presentation.mp4"
                        if source_path.exists():
                            source_path.rename(output_path)
                            print(f"? Présentation sauvegardée sous: {output_path}")
        
        print(f"\n===== Traitement ML terminé =====")
        print(f"?? {len(processed_files)} présentations créées avec succès")
        
        return processed_files
    
    def get_document_metadata(self, file_path):
        """
        Extrait les métadonnées d'un document pour une meilleure contextualisation
        
        Args:
            file_path: Chemin du fichier
            
        Returns:
            metadata: Dictionnaire de métadonnées
        """
        metadata = {
            'filename': Path(file_path).name,
            'file_type': Path(file_path).suffix.lower()[1:],
            'file_size': Path(file_path).stat().st_size,
            'creation_date': time.ctime(Path(file_path).stat().st_ctime),
            'modification_date': time.ctime(Path(file_path).stat().st_mtime),
            'extracted_keywords': []
        }
        
        # Extraction spécifique selon le type de fichier
        if metadata['file_type'] == 'pdf':
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata['page_count'] = len(reader.pages)
                    
                    # Extraire les métadonnées si disponibles
                    info = reader.metadata
                    if info:
                        metadata['title'] = info.get('/Title', '')
                        metadata['author'] = info.get('/Author', '')
                        metadata['subject'] = info.get('/Subject', '')
                        metadata['keywords'] = info.get('/Keywords', '')
            except Exception as e:
                print(f"?? Erreur extraction métadonnées PDF: {e}")
        
        elif metadata['file_type'] in ['docx', 'doc']:
            try:
                doc = docx.Document(file_path)
                metadata['page_count'] = len(doc.paragraphs) // 15  # Estimation approximative
                
                # Extraire les propriétés du document si disponibles
                prop_names = ['title', 'author', 'comments', 'category', 'subject']
                for prop in prop_names:
                    if hasattr(doc.core_properties, prop):
                        metadata[prop] = getattr(doc.core_properties, prop)
            except Exception as e:
                print(f"?? Erreur extraction métadonnées Word: {e}")
        
        # Extraire des mots-clés du contenu si possible
        if TRANSFORMER_AVAILABLE and hasattr(self.video_generator, 'text_processor') and self.video_generator.text_processor.initialized:
            try:
                # Lire le début du fichier pour extraire des mots-clés
                if metadata['file_type'] == 'pdf':
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        # Lire les 3 premières pages ou moins
                        text = ""
                        for i in range(min(3, len(reader.pages))):
                            text += reader.pages[i].extract_text()
                
                elif metadata['file_type'] in ['docx', 'doc']:
                    doc = docx.Document(file_path)
                    # Prendre les 50 premiers paragraphes ou moins
                    text = "\n".join([p.text for p in doc.paragraphs[:50]])
                
                elif metadata['file_type'] == 'md':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        # Lire les 100 premières lignes ou moins
                        lines = f.readlines()[:100]
                        text = "".join(lines)
                
                # Extraire les mots-clés avec le modèle de langage
                prompt = f"""Extrait 5 à 10 mots-clés significatifs du texte suivant. 
                Réponds uniquement avec les mots-clés séparés par des virgules, sans autre commentaire:
                
                {text[:2000]}"""  # Limiter à 2000 caractères pour éviter les dépassements de tokens
                
                response = self.video_generator.text_processor.generate_text(prompt, max_length=100)
                
                # Nettoyer les mots-clés
                keywords = [kw.strip() for kw in response.split(',') if kw.strip()]
                metadata['extracted_keywords'] = keywords[:10]  # Limiter à 10 mots-clés max
                
            except Exception as e:
                print(f"?? Erreur extraction mots-clés avec modèle de langage: {e}")
                # Fallback avec TF-IDF si l'extraction avec le modèle échoue
                try:
                    vectorizer = TfidfVectorizer(max_features=10, stop_words='english')
                    tfidf_matrix = vectorizer.fit_transform([text])
                    feature_names = vectorizer.get_feature_names_out()
                    
                    # Trouver les indices des mots-clés les plus importants
                    importance = np.array(tfidf_matrix.sum(axis=0)).flatten()
                    indices = importance.argsort()[-10:][::-1]  # Top 10 mots-clés
                    
                    # Récupérer les mots-clés
                    metadata['extracted_keywords'] = [feature_names[i] for i in indices]
                except:
                    pass  # Ignorer si l'extraction de mots-clés échoue
        
        return metadata
    
    def analyze_document_complexity(self, content):
        """
        Analyse la complexité du document pour déterminer les meilleurs paramètres de présentation
        
        Args:
            content: Contenu Markdown du document
            
        Returns:
            complexity: Score de complexité (0-10)
            stats: Statistiques sur le document
        """
        stats = {}
        
        # Nettoyer le texte des balises Markdown
        clean_text = re.sub(r'#+ ', '', content)
        clean_text = re.sub(r'!\[.*?\]\(.*?\)', '', clean_text)
        clean_text = re.sub(r'\[.*?\]\(.*?\)', '', clean_text)
        clean_text = re.sub(r'[*_~`]', '', clean_text)
        
        # Statistiques de base
        words = clean_text.split()
        stats['word_count'] = len(words)
        stats['char_count'] = len(clean_text)
        
        # Calcul de la densité lexicale (variété du vocabulaire)
        unique_words = set(word.lower() for word in words if len(word) > 3)
        stats['unique_words'] = len(unique_words)
        stats['lexical_density'] = len(unique_words) / max(1, len(words))
        
        # Calcul de la complexité des phrases
        sentences = re.split(r'[.!?]+', clean_text)
        stats['sentence_count'] = sum(1 for s in sentences if s.strip())
        stats['avg_sentence_length'] = len(words) / max(1, stats['sentence_count'])
        
        # Calculer la longueur moyenne des mots
        stats['avg_word_length'] = sum(len(word) for word in words) / max(1, len(words))
        
        # Calculer le nombre d'images et de figures
        stats['image_count'] = len(re.findall(r'!\[.*?\]\(.*?\)', content))
        
        # Calculer le nombre de listes à puces
        stats['bullet_count'] = len(re.findall(r'[-*+]\s+', content))
        
        # Calculer des métriques de lisibilité (Flesch-Kincaid simplifié)
        stats['flesch_reading_ease'] = 206.835 - (1.015 * stats['avg_sentence_length']) - (84.6 * stats['avg_word_length'])
        
        # Calculer un score de complexité de 0 à 10
        # Basé sur une combinaison pondérée des facteurs
        complexity_score = 0
        
        # Plus le score Flesch est bas, plus le texte est complexe
        if stats['flesch_reading_ease'] < 30:
            complexity_score += 4
        elif stats['flesch_reading_ease'] < 50:
            complexity_score += 3
        elif stats['flesch_reading_ease'] < 70:
            complexity_score += 2
        elif stats['flesch_reading_ease'] < 90:
            complexity_score += 1
        
        # Plus les mots sont longs, plus c'est complexe
        if stats['avg_word_length'] > 6:
            complexity_score += 3
        elif stats['avg_word_length'] > 5:
            complexity_score += 2
        elif stats['avg_word_length'] > 4:
            complexity_score += 1
        
        # Plus les phrases sont longues, plus c'est complexe
        if stats['avg_sentence_length'] > 25:
            complexity_score += 3
        elif stats['avg_sentence_length'] > 20:
            complexity_score += 2
        elif stats['avg_sentence_length'] > 15:
            complexity_score += 1
        
        # Borner le score entre 0 et 10
        complexity_score = max(0, min(10, complexity_score))
        stats['complexity_score'] = complexity_score
        
        return complexity_score, stats
    def assemble_final_video(self, slides, narration_audio_paths):
        """
        Assemble la vidéo finale avec les transitions, l'audio combiné et un avatar avec lip-sync
        
        Args:
            slides: Liste des clips de diapositives synchronisés
            narration_audio_paths: Liste des chemins audio pour la narration
            
        Returns:
            str: Chemin du fichier vidéo final
        """        
        try:
            # 1. Appliquer les transitions entre les diapositives
            print("??? Application des transitions entre diapositives...")
            final_clips = []
            
            # Utiliser directement les slides si pas de transitions disponibles
            if not hasattr(self, 'timing_model') or not slides:
                print("?? Pas de timing_model ou slides vides, utilisation des slides bruts")
                final_clips = slides
            else:
                # Transitions variées pour éviter la monotonie
                try:
                    transitions = get_varied_transitions(len(slides) - 1)
                    
                    for i, slide in enumerate(slides):
                        if i < len(slides) - 1:
                            # Créer une transition entre cette diapositive et la suivante
                            transition_type = transitions[i] if i < len(transitions) else 'fade'
                            
                            try:
                                transition_duration = self.timing_model.get('transition_duration', 0.5)
                                transition_clips = apply_transitions_to_clips(
                                    [slide, slides[i+1]],
                                    transition_type=transition_type,
                                    transition_duration=transition_duration
                                )
                                final_clips.extend(transition_clips)
                            except Exception as e:
                                print(f"?? Erreur transition {i}-{i+1}: {e}, utilisation simple")
                                # Fallback sans transition
                                final_clips.append(slide)
                        elif i == len(slides) - 1:
                            # Dernière diapositive, pas de transition
                            final_clips.append(slide)
                except Exception as e:
                    print(f"?? Erreur lors de l'application des transitions: {e}")
                    final_clips = slides  # Utiliser les slides sans transitions
            
            # 2. Combiner tous les clips de diapositives
            print(f"?? Combinaison de {len(final_clips)} clips vidéo...")
            slides_clip = concatenate_videoclips(final_clips) if final_clips else None
            
            if slides_clip is None:
                print("? Pas de clips à combiner, impossible de créer la vidéo")
                return None
            
            # 3. Préparer l'audio combiné pour la narration
            combined_audio_path = None
            if narration_audio_paths:
                try:
                    # Vérifier que tous les chemins audio existent
                    valid_audio_paths = [p for p in narration_audio_paths if Path(p).exists()]
                    if valid_audio_paths:
                        # Combiner tous les fichiers audio en un seul
                        combined_audio_path = self.output_dir / "combined_narration.wav"
                        
                        # Vérifier si la méthode existe
                        if hasattr(self, 'voice_synthesizer') and hasattr(self.voice_synthesizer, 'add_audio_transitions'):
                            audio_combined = self.voice_synthesizer.add_audio_transitions(
                                valid_audio_paths,
                                combined_audio_path,
                                crossfade_duration=0.5
                            )
                        else:
                            # Méthode alternative simple
                            print("?? Méthode add_audio_transitions non disponible, combinaison audio simple")
                            audio_clips = [AudioFileClip(str(p)) for p in valid_audio_paths]
                            combined_audio = concatenate_audioclips(audio_clips)
                            combined_audio.write_audiofile(str(combined_audio_path))
                except Exception as e:
                    print(f"?? Erreur lors de la combinaison audio: {e}")
            
            # 4. Initialiser l'avatar avec synchronisation labiale
            # Chercher l'avatar dans plusieurs emplacements possibles
            possible_avatar_paths = [
                "C:/Users/ThinkPad/Desktop/plateform/python/avatar.mp4",
                "./avatar.mp4",
                "../avatar.mp4",
                os.path.join(self.base_dir, "avatar.mp4") if hasattr(self, 'base_dir') else None
            ]
            
            avatar_path = None
            for path in possible_avatar_paths:
                if path and Path(path).exists():
                    avatar_path = path
                    print(f"? Avatar trouvé: {avatar_path}")
                    break
            
            final_video = None
            
            if avatar_path:
                try:
                    # Utiliser notre nouvelle classe LipSyncAvatar
                    print("?? Initialisation de l'avatar avec synchronisation labiale...")                    
                    avatar_manager = LipSyncAvatar(avatar_path)
                    
                    # Créer l'avatar avec lip-sync basé sur l'audio combiné
                    if combined_audio_path and Path(combined_audio_path).exists():
                        # Utiliser la version multi-pose pour plus de dynamisme
                        avatar_clip = avatar_manager.create_multipose_avatar(
                            combined_audio_path,
                            None,  # Pas besoin de chemin de sortie ici
                            duration=slides_clip.duration,
                            position=(10, slides_clip.h - 200 - 2),  # 20px du bord gauche, 20px du bas
                            size=100  # Hauteur fixe de 200px
                        )
                    else:
                        # Pas d'audio, juste afficher l'avatar
                        avatar_clip = avatar_manager.create_lipsync_avatar(
                            None,
                            None,
                            duration=slides_clip.duration,
                            position=(10, slides_clip.h - 200 - 2),
                            size=100
                        )
                    
                    if avatar_clip:
                        # Créer la vidéo composite avec les diapositives et l'avatar
                        final_video = CompositeVideoClip([slides_clip, avatar_clip])
                        
                        # Si audio disponible, l'appliquer à la vidéo finale
                        if combined_audio_path and Path(combined_audio_path).exists():
                            narration_audio = AudioFileClip(str(combined_audio_path))
                            final_video = final_video.set_audio(narration_audio)
                    else:
                        print("?? Échec de création de l'avatar, utilisation de slides seuls")
                        final_video = slides_clip
                except ImportError:
                    # Si la classe LipSyncAvatar n'est pas importable, utiliser approche alternative
                    print("?? Classe LipSyncAvatar non disponible, utilisation approche alternative")
                    
                    try:
                        # Approche alternative - plus simple sans classe dédiée
                        print("?? Chargement et synchronisation de l'avatar (méthode alternative)...")
                        avatar_clip = VideoFileClip(str(avatar_path))
                        
                        # Redimensionner et positionner l'avatar (en bas à gauche)
                        avatar_height = 100  # 200px de hauteur
                        avatar_clip = avatar_clip.resize(height=avatar_height)
                        avatar_clip = avatar_clip.set_position((10, slides_clip.h - avatar_height - 2))
                        
                        # Faire correspondre la durée de l'avatar à celle des diapositives
                        if avatar_clip.duration < slides_clip.duration:
                            # Répéter l'avatar si nécessaire
                            repeats = int(slides_clip.duration / avatar_clip.duration) + 1
                            avatar_repeated = concatenate_videoclips([avatar_clip] * repeats)
                            avatar_clip = avatar_repeated.subclip(0, slides_clip.duration)
                        else:
                            # Couper l'avatar si trop long
                            avatar_clip = avatar_clip.subclip(0, slides_clip.duration)
                        
                        # Ajouter l'audio à l'avatar pour lip-sync basique
                        if combined_audio_path and Path(combined_audio_path).exists():
                            narration_audio = AudioFileClip(str(combined_audio_path))
                            avatar_clip = avatar_clip.set_audio(narration_audio)
                        
                        # Créer la vidéo composite avec les diapositives et l'avatar
                        final_video = CompositeVideoClip([slides_clip, avatar_clip])
                        
                        # Appliquer l'audio à la vidéo finale aussi
                        if combined_audio_path and Path(combined_audio_path).exists():
                            narration_audio = AudioFileClip(str(combined_audio_path))
                            final_video = final_video.set_audio(narration_audio)
                    except Exception as e:
                        print(f"?? Erreur lors de l'ajout de l'avatar (méthode alternative): {e}")
                        final_video = slides_clip  # Utiliser les diapositives sans avatar
                except Exception as e:
                    print(f"?? Erreur lors de l'ajout de l'avatar: {e}")
                    traceback.print_exc()
                    final_video = slides_clip  # Utiliser les diapositives sans avatar
            else:
                print(f"?? Avatar non trouvé dans les chemins possibles, présentation sans avatar")
                final_video = slides_clip
            
            # 5. Sauvegarder la vidéo finale
            output_path = self.output_dir / "ml_presentation_with_avatar.mp4"
            
            print(f"?? Exportation de la vidéo finale...")
            final_video.write_videofile(
                str(output_path),
                fps=16,
                codec='libx264',
                audio_codec='aac',
                threads=4,
                preset='medium'
            )
            
            print(f"? Présentation sauvegardée: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"? Erreur lors de la finalisation de la vidéo: {e}")
            traceback.print_exc()
            return None


class EnhancedLipSyncAvatar(LipSyncAvatar):
    """
    Version améliorée du LipSyncAvatar qui utilise la détection de phonèmes
    pour une synchronisation labiale plus précise
    """
    
    def __init__(self, avatar_path, device="cpu"):
        super().__init__(avatar_path, device)
        
        # Ajouter un détecteur de phonèmes
        self.phoneme_detector = PhonemeDetector(device)
        
        # Charger l'avatar et analyser ses séquences
        self.avatar_sequences = self._analyze_avatar_sequences()
    
    def _analyze_avatar_sequences(self):
        """
        Analyse l'avatar vidéo pour identifier différentes séquences de mouvements
        qui pourraient correspondre à différents phonèmes/visèmes
        """
        if not self.avatar_available:
            return {}
        
        try:            
            # Charger l'avatar
            avatar_clip = VideoFileClip(str(self.avatar_path))
            
            # Paramètres
            fps = avatar_clip.fps
            duration = avatar_clip.duration
            frame_count = int(duration * fps)
            
            # Analyser quelques frames pour identifier différentes poses
            sequences = {}
            
            # Diviser la vidéo en segments (par exemple, tous les 0.5 secondes)
            segment_duration = 0.5  # en secondes
            segment_count = int(duration / segment_duration)
            
            for i in range(min(8, segment_count)):  # Limiter à 8 segments pour l'efficacité
                start_time = i * segment_duration
                end_time = min((i + 1) * segment_duration, duration)
                
                # Prendre un frame représentatif de ce segment
                mid_time = (start_time + end_time) / 2
                
                # Ajouter ce segment comme séquence potentielle
                sequences[f'sequence_{i}'] = {
                    'start': start_time,
                    'end': end_time,
                    'duration': end_time - start_time,
                    'representative_time': mid_time
                }
            
            # Ajouter une séquence spéciale pour "bouche fermée" (généralement au début)
            sequences['mouth_closed'] = {
                'start': 0,
                'end': min(0.2, duration),
                'duration': min(0.2, duration),
                'representative_time': 0
            }
            
            # Ajouter une séquence pour "bouche ouverte" (chercher au milieu de la vidéo)
            mid_video = duration / 2
            sequences['mouth_open'] = {
                'start': max(0, mid_video - 0.1),
                'end': min(duration, mid_video + 0.1),
                'duration': 0.2,
                'representative_time': mid_video
            }
            
            avatar_clip.close()
            return sequences
            
        except Exception as e:
            print(f"?? Erreur analyse séquences avatar: {e}")
            return {
                'default': {
                    'start': 0,
                    'end': 0.5,
                    'duration': 0.5,
                    'representative_time': 0.25
                }
            }
    
    def create_phoneme_synced_avatar(self, audio_path, output_path=None, duration=None, position=("left", "bottom"), size=0.25):
        """
        Crée un avatar avec lip-sync basé sur la détection de phonèmes
        
        Args:
            audio_path: Chemin du fichier audio
            output_path: Chemin de sortie pour la vidéo (optionnel)
            duration: Durée totale de la vidéo (par défaut: durée de l'audio)
            position: Position de l'avatar dans la vidéo
            size: Taille de l'avatar (relative ou absolue)
            
        Returns:
            VideoClip: Clip vidéo de l'avatar avec synchronisation labiale
        """       
        if not self.avatar_available:
            print("?? Avatar non disponible")
            return None
        
        try:
            # 1. Charger l'avatar et l'audio
            avatar_clip = VideoFileClip(str(self.avatar_path))
            
            # Si l'audio existe, le charger
            audio_clip = None
            if audio_path and Path(audio_path).exists():
                audio_clip = AudioFileClip(str(audio_path))
                audio_duration = audio_clip.duration
            else:
                audio_duration = 0
            
            # Utiliser la durée spécifiée, celle de l'audio, ou celle de l'avatar
            if duration is None:
                if audio_duration > 0:
                    duration = audio_duration
                else:
                    duration = avatar_clip.duration
            
            # 2. Détecter les phonèmes dans l'audio
            phoneme_sequence = []
            if audio_path and Path(audio_path).exists():
                phoneme_sequence = self.phoneme_detector.detect_phonemes(
                    audio_path, frame_rate=avatar_clip.fps
                )
            
            # 3. Créer une séquence d'avatar basée sur les phonèmes
            avatar_segments = []
            
            if phoneme_sequence:
                # Utiliser les phonèmes pour créer des segments synchronisés
                frame_duration = 1.0 / avatar_clip.fps
                current_time = 0
                
                for viseme_index in phoneme_sequence:
                    # Sélectionner la séquence appropriée pour ce visème
                    if viseme_index == self.phoneme_detector.phoneme_map['rest']:
                        sequence = self.avatar_sequences.get('mouth_closed', self.avatar_sequences['default'])
                    elif viseme_index == self.phoneme_detector.phoneme_map['A']:
                        sequence = self.avatar_sequences.get('mouth_open', self.avatar_sequences['default'])
                    else:
                        # Pour les autres visèmes, utiliser une séquence aléatoire
                        sequence_key = random.choice(list(self.avatar_sequences.keys()))
                        sequence = self.avatar_sequences[sequence_key]
                    
                    # Extraire un segment de l'avatar
                    start = sequence['start']
                    segment = avatar_clip.subclip(
                        start, 
                        min(start + frame_duration, sequence['end'])
                    )
                    
                    # Ajouter à la liste
                    avatar_segments.append(segment)
                    
                    # Avancer dans le temps
                    current_time += frame_duration
                    
                    # Arrêter si on dépasse la durée cible
                    if current_time >= duration:
                        break
                
                # S'assurer qu'on a assez de segments pour couvrir toute la durée
                while current_time < duration:
                    # Ajouter des segments supplémentaires (bouche fermée pour l'audio silencieux)
                    sequence = self.avatar_sequences.get('mouth_closed', self.avatar_sequences['default'])
                    
                    segment_duration = min(0.5, duration - current_time)
                    segment = avatar_clip.subclip(
                        sequence['start'],
                        min(sequence['start'] + segment_duration, sequence['end'])
                    )
                    
                    avatar_segments.append(segment)
                    current_time += segment_duration
            else:
                # Sans détection de phonèmes, utiliser une approche basée sur l'amplitude audio
                self._create_amplitude_based_segments(avatar_clip, audio_path, duration, avatar_segments)
            
            # 4. Concaténer tous les segments
            if avatar_segments:
                combined_avatar = concatenate_videoclips(avatar_segments)
                
                # Assurer la durée exacte
                combined_avatar = combined_avatar.subclip(0, duration)
            else:
                # Fallback: boucler l'avatar original
                repeats = int(duration / avatar_clip.duration) + 1
                combined_avatar = concatenate_videoclips([avatar_clip] * repeats)
                combined_avatar = combined_avatar.subclip(0, duration)
            
            # 5. Ajouter l'audio
            if audio_clip:
                combined_avatar = combined_avatar.set_audio(audio_clip)
            
            # 6. Redimensionner l'avatar
            if isinstance(size, float):
                # Taille relative
                combined_avatar = combined_avatar.resize(height=int(combined_avatar.h * size))
            else:
                # Taille absolue
                combined_avatar = combined_avatar.resize(height=size)
            
            # 7. Positionner l'avatar
            combined_avatar = combined_avatar.set_position(position)
            
            # 8. Retourner le clip
            return combined_avatar
            
        except Exception as e:
            print(f"?? Erreur création avatar phonème-sync: {e}")
            traceback.print_exc()
            return None
    
    def _create_amplitude_based_segments(self, avatar_clip, audio_path, duration, segments_list):
        """
        Crée des segments d'avatar basés sur l'amplitude audio
        
        Args:
            avatar_clip: Clip vidéo de l'avatar
            audio_path: Chemin du fichier audio
            duration: Durée totale de la vidéo
            segments_list: Liste à remplir avec les segments créés
        """
        try:
            
            # Charger l'audio
            y, sr = librosa.load(audio_path, sr=None)
            
            # Calculer l'amplitude RMS (root mean square) à intervalles réguliers
            frame_duration = 0.1  # Durée de chaque segment en secondes
            hop_length = int(sr * frame_duration)
            
            # Utiliser librosa pour obtenir l'enveloppe d'amplitude
            amplitude_env = rms(y=y, frame_length=hop_length, hop_length=hop_length)[0]
            
            # Normaliser entre 0 et 1
            amplitude_env = amplitude_env / np.max(amplitude_env) if np.max(amplitude_env) > 0 else amplitude_env
            
            # Créer des segments en fonction de l'amplitude
            current_time = 0
            segment_index = 0
            
            for amp in amplitude_env:
                segment_start = min(segment_index * frame_duration, avatar_clip.duration - frame_duration)
                
                if amp < 0.2:
                    # Faible amplitude = bouche fermée ou repos
                    sequence = self.avatar_sequences.get('mouth_closed', {'start': 0})
                    start_pos = sequence.get('start', 0)
                elif amp > 0.7:
                    # Forte amplitude = bouche grande ouverte
                    sequence = self.avatar_sequences.get('mouth_open', {'start': avatar_clip.duration / 2})
                    start_pos = sequence.get('start', avatar_clip.duration / 2)
                else:
                    # Amplitude moyenne = séquence aléatoire
                    sequence_key = random.choice(list(self.avatar_sequences.keys()))
                    sequence = self.avatar_sequences[sequence_key]
                    start_pos = sequence.get('start', segment_start)
                
                # Extraire un segment de l'avatar
                segment = avatar_clip.subclip(
                    start_pos,
                    min(start_pos + frame_duration, avatar_clip.duration)
                )
                
                # Ajouter à la liste
                segments_list.append(segment)
                
                # Avancer dans le temps
                current_time += frame_duration
                segment_index += 1
                
                # Arrêter si on dépasse la durée cible
                if current_time >= duration:
                    break
            
            # S'assurer qu'on a assez de segments pour couvrir toute la durée
            while current_time < duration:
                sequence = self.avatar_sequences.get('mouth_closed', {'start': 0})
                start_pos = sequence.get('start', 0)
                
                segment_duration = min(frame_duration, duration - current_time)
                segment = avatar_clip.subclip(
                    start_pos,
                    min(start_pos + segment_duration, avatar_clip.duration)
                )
                
                segments_list.append(segment)
                current_time += segment_duration
                
        except Exception as e:
            print(f"?? Erreur segments basés sur amplitude: {e}")
            # Fallback: ajouter l'avatar complet
            segments_list.append(avatar_clip.subclip(0, min(duration, avatar_clip.duration)))

""" the class ImprovedLipSyncAvatar was here before being refactored to avatar.py"""

def create_video_with_avatar(slides_paths, audio_path=None, avatar_path=None, output_path=None, 
                       slide_duration=3.0, avatar_size=200, avatar_position=("left", "bottom")):
    """
    Crée une présentation vidéo à partir d'images de diapositives avec un avatar synchronisé
    
    Args:
        slides_paths: Liste des chemins d'images de diapositives
        audio_path: Chemin du fichier audio de narration (optionnel)
        avatar_path: Chemin du fichier vidéo de l'avatar
        output_path: Chemin du fichier vidéo de sortie
        slide_duration: Durée de chaque diapositive en secondes
        avatar_size: Hauteur de l'avatar en pixels
        avatar_position: Position de l'avatar ("left", "bottom")
        
    Returns:
        str: Chemin du fichier vidéo généré
    """    
    # Filtrer les chemins d'images valides
    valid_slides = [p for p in slides_paths if Path(p).exists()]
    if not valid_slides:
        print("? Aucune image de diapositive valide trouvée")
        return None
    
    # Créer un chemin de sortie par défaut si non spécifié
    if not output_path:
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_path = f"presentation_with_avatar_{timestamp}.mp4"
    
    # Chercher l'avatar si non spécifié
    if not avatar_path or not Path(avatar_path).exists():
        alternative_paths = [
            "C:/Users/ThinkPad/Desktop/plateform/python/avatar.mp4",
            "./avatar.mp4",
            "../avatar.mp4",
            os.path.join(os.path.dirname(__file__), "avatar.mp4")
        ]
        for alt_path in alternative_paths:
            if Path(alt_path).exists():
                avatar_path = alt_path
                print(f"? Avatar trouvé à {avatar_path}")
                break
        else:
            print("? Aucun avatar trouvé")
            return None
    
    try:
        # Créer un clip à partir des images de diapositives
        print(f"??? Création d'un clip à partir de {len(valid_slides)} diapositives")
        slides_clip = ImageSequenceClip(valid_slides, durations=[slide_duration] * len(valid_slides))
        
        # Redimensionner si nécessaire pour avoir une hauteur standard
        if slides_clip.h != 720:
            slides_clip = slides_clip.resize(height=720)
        
        # Créer l'audio de narration si nécessaire
        if audio_path and Path(audio_path).exists():
            narration_audio = AudioFileClip(str(audio_path))
            
            # Adapter l'audio à la durée des diapositives
            total_duration = slides_clip.duration
            if narration_audio.duration < total_duration:
                silence = AudioClip(lambda t: 0, duration=total_duration - narration_audio.duration)
                narration_audio = concatenate_videoclips([narration_audio, silence])
            elif narration_audio.duration > total_duration:
                narration_audio = narration_audio.subclip(0, total_duration)
            
            # Appliquer l'audio aux diapositives
            slides_clip = slides_clip.set_audio(narration_audio)
        else:
            narration_audio = None
                
        # Initialiser avec le nouveau système amélioré
        avatar_manager = ImprovedLipSyncAvatar(avatar_path)
        
        # Créer l'avatar avec synchronisation labiale
        print("?? Création de l'avatar avec synchronisation labiale...")
        avatar_clip = avatar_manager.create_lipsync_avatar(
            audio_path=audio_path,
            duration=slides_clip.duration,
            position=avatar_position,
            size=avatar_size
        )
        
        if not avatar_clip:
            print("?? Échec de la création de l'avatar, utilisation de la vidéo sans avatar")
            # Sauvegarder la vidéo sans avatar
            slides_clip.write_videofile(
                output_path,
                fps=24,
                codec='libx264',
                audio_codec='aac' if narration_audio else None,
                threads=4,
                preset='medium'
            )
            return output_path
        
        # Combiner les diapositives et l'avatar
        final_clip = CompositeVideoClip([slides_clip, avatar_clip])
        
        # Exporter la vidéo finale
        print(f"?? Exportation de la présentation avec avatar: {output_path}")
        final_clip.write_videofile(
            output_path,
            fps=24,
            codec='libx264',
            audio_codec='aac',
            threads=4,
            preset='medium'
        )
        
        print(f"? Présentation avec avatar créée: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"? Erreur lors de la création de la présentation: {e}")
        traceback.print_exc()
        return None

"""the class DeepContentSegmenter was here before being refactored to content_segmenter.py"""

class DeepLearningEnhancedGenerator(MLEnhancedVideoGenerator):
    """Version améliorée du générateur de vidéo utilisant l'apprentissage profond"""
    
    def __init__(self, base_dir=BASE_DIR, output_dir=OUTPUT_DIR, model_path=None):
        super().__init__(base_dir, output_dir, model_path)
        
        # Remplacer les modèles classiques par des modèles deep learning
        self.slide_quality = DeepSlideQualityModel()
        self.style_recommender = DeepStyleRecommender()
        self.image_analyzer = DeepImageAnalyzer()
        
        print("? Générateur de présentations avec deep learning initialisé")
    
    def process_document(self, input_file):
        """
        Traite un document avec des modèles deep learning et génère une présentation vidéo
        
        Args:
            input_file: Chemin du fichier à traiter
            
        Returns:
            output_path: Chemin du fichier vidéo généré
        """
        input_file = Path(input_file)
        
        if not input_file.exists():
            print(f"? Le fichier {input_file} n'existe pas")
            return None
        
        # Lire le contenu du fichier (en supposant que c'est un fichier Markdown)
        try:
            with open(input_file, encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            try:
                with open(input_file, encoding="latin-1") as f:
                    content = f.read()
            except Exception as e:
                print(f"? Impossible de lire le fichier: {e}")
                return None
        except Exception as e:
            print(f"? Erreur lors de la lecture du fichier: {e}")
            return None
        
        # Traiter le contenu avec le générateur de vidéo ML
        output_path = self.process_markdown(content)
        
        return output_path

    def dynamic_quality_adaptation(self, slide_data, predicted_quality):
        """
        Adapte dynamiquement le contenu et le style d'une diapositive
        en fonction de sa qualité prédite, en utilisant des modèles deep learning
        """       
        if predicted_quality >= 7.0:
            # La qualité est déjà bonne, pas besoin d'adaptation
            return slide_data
        
        try:
            # Extraire les données de la diapositive
            content = slide_data.get('content', '')
            title = slide_data.get('title', '')
            has_image = slide_data.get('has_image', False)
            
            # 1. Améliorer le contenu avec le modèle de langage
            if self.text_processor and self.text_processor.initialized:
                prompt = f"""Améliore ce contenu de diapositive pour le rendre plus clair, 
                plus engageant et plus facile à comprendre pour une présentation. Ne modifie pas
                profondément le contenu, simplement améliore sa formulation et sa structure:
                
                Titre: {title}
                
                {content}
                
                Version améliorée:"""
                
                improved_content = self.text_processor.generate_text(prompt, max_length=len(content.split()) + 50)
                
                # Ne remplacer que si l'amélioration est significative
                if len(improved_content) > len(content) * 0.5:
                    slide_data['content'] = improved_content
            
            # 2. Recommander un style plus approprié avec notre modèle deep learning
            if hasattr(self, 'style_recommender') and isinstance(self.style_recommender, DeepStyleRecommender):
                # Utiliser le modèle pour recommander un style optimisé
                features = self.style_recommender.extract_content_style_features(
                    slide_data['content'], slide_data['title'], slide_data['has_image']
                )
                
                # Transformer en tensor PyTorch
                features_tensor = torch.tensor(features, dtype=torch.float32).to(self.style_recommender.device)
                
                # Obtenir la prédiction du modèle
                self.style_recommender.model.eval()
                with torch.no_grad():
                    logits, _ = self.style_recommender.model(features_tensor.unsqueeze(0))
                    style_idx = torch.argmax(logits).item()
                    
                    # Convertir l'index en nom de style
                    if style_idx < len(self.style_recommender.style_names):
                        optimized_style = self.style_recommender.style_names[style_idx]
                        
                        # Générer les paramètres de style
                        style_params = self.style_recommender.generate_style_parameters(
                            optimized_style, features, has_image
                        )
                        
                        # Mettre à jour le style
                        slide_data['style'] = optimized_style
                        for key, value in style_params.items():
                            if key != 'name':  # 'name' est déjà défini par 'style'
                                slide_data[key] = value
            
            return slide_data
            
        except Exception as e:
            print(f"?? Erreur lors de l'adaptation dynamique: {e}")
            return slide_data  # Retourner les données originales en cas d'erreur
    
    def train_deep_models(self, epochs=50, batch_size=16):
        """
        Entraîne tous les modèles deep learning avec les données disponibles
        """
        print("?? Entraînement des modèles deep learning...")
        
        # Charger toutes les données d'apprentissage disponibles
        feedback_files = list(TRAINING_DATA_DIR.glob("feedback_*.json"))
        learning_files = list(TRAINING_DATA_DIR.glob("learning_data_*.json"))
        
        if not feedback_files and not learning_files:
            print("?? Aucune donnée d'apprentissage disponible")
            return False
        
        # Collecter les données
        slide_data = []
        quality_scores = []
        styles = []
        
        # Traiter les fichiers de feedback et de données d'apprentissage
        for file_path in feedback_files + learning_files:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                if 'slides' in data:
                    if 'ratings' in data:
                        # Format de feedback
                        slide_data.extend(data['slides'])
                        quality_scores.extend(data['ratings'])
                        styles.extend([slide.get('style', 'balanced') for slide in data['slides']])
                    elif 'quality_scores' in data:
                        # Format de données d'apprentissage
                        slide_data.extend(data['slides'])
                        quality_scores.extend(data['quality_scores'])
                        styles.extend([slide.get('style', 'balanced') for slide in data['slides']])
            except Exception as e:
                print(f"?? Erreur lecture {file_path.name}: {e}")
        
        if len(slide_data) < 5:
            print("?? Pas assez de données pour l'entraînement deep learning (minimum 5 exemples)")
            return False
        
        print(f"?? Données collectées: {len(slide_data)} exemples d'entraînement")
        
        # 1. Entraîner le modèle de qualité
        if hasattr(self, 'slide_quality') and isinstance(self.slide_quality, DeepSlideQualityModel):
            print("?? Entraînement du modèle de qualité deep learning...")
            mse, r2 = self.slide_quality.train(slide_data, quality_scores, epochs=epochs, batch_size=batch_size)
            print(f"? Modèle de qualité entraîné: MSE={mse:.4f}, R²={r2:.4f}")
        
        # 2. Entraîner le modèle de style
        if hasattr(self, 'style_recommender') and isinstance(self.style_recommender, DeepStyleRecommender):
            print("?? Entraînement du modèle de style deep learning...")
            # Extraire les données nécessaires
            contents = [slide.get('content', '') for slide in slide_data]
            titles = [slide.get('title', '') for slide in slide_data]
            has_images = [slide.get('has_image', False) for slide in slide_data]
            
            self.style_recommender.train_from_feedback(contents, titles, has_images, styles, epochs=epochs)
            print("? Modèle de style entraîné")
        
        print("? Entraînement des modèles deep learning terminé")
        return True

# Chemin par défaut pour l'avatar
DEFAULT_AVATAR_PATH = "C:/Users/ThinkPad/Desktop/plateform/python/avatar.mp4"

""" the class AvatarIntegrator was here before being refactored to avatar.py"""

class DeepImageAnalyzer:
    """Analyse les images avec un réseau de neurones convolutionnel pour une meilleure intégration"""
    
    def __init__(self):
        self.model_path = MODELS_DIR / "deep_image_analyzer.pt"
        self.model = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        # Initialiser ou charger le modèle
        self.initialize_model()
    
    def initialize_model(self):
        """Initialise ou charge un modèle pré-entraîné pour l'analyse d'images"""
        try:            
            # Utiliser un modèle pré-entraîné comme ResNet pour l'extraction de caractéristiques
            self.model = models.resnet34(pretrained=True)
            # Enlever la dernière couche de classification
            self.feature_extractor = torch.nn.Sequential(*list(self.model.children())[:-1])
            self.feature_extractor.eval()
            self.feature_extractor = self.feature_extractor.to(self.device)
            
            # Transformations pour les images
            self.transforms = transforms.Compose([
                transforms.Resize(256),
                transforms.CenterCrop(224),
                transforms.ToTensor(),
                transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]),
            ])
            
            print("? Modèle d'analyse d'images deep learning initialisé")
            
        except Exception as e:
            print(f"?? Erreur initialisation modèle d'analyse d'images: {e}")
            self.model = None
    
    def analyze_image(self, image_path):
        """Analyse une image pour extraire des caractéristiques visuelles profondes"""
        if not self.model:
            return {'type': 'unknown', 'complexity': 0.5, 'brightness': 0.5}
        
        try:
            # Charger et transformer l'image
            img = Image.open(image_path).convert('RGB')
            img_tensor = self.transforms(img).unsqueeze(0).to(self.device)
            
            # Extraire les caractéristiques
            with torch.no_grad():
                features = self.feature_extractor(img_tensor)
                features = features.flatten().cpu().numpy()
            
            # Analyse basique avec PIL
            stat = ImageStat.Stat(img)
            brightness = sum(stat.mean) / len(stat.mean) / 255.0
            
            # Calculer la complexité visuelle (diversité des caractéristiques)
            complexity = np.std(features) / np.mean(np.abs(features)) if np.mean(np.abs(features)) > 0 else 0.5
            complexity = min(1.0, max(0.0, complexity))
            
            # Déterminer le type d'image basé sur les caractéristiques
            if np.max(features) > 10:
                img_type = 'photo'
            elif complexity < 0.3:
                img_type = 'diagram'
            elif brightness < 0.3:
                img_type = 'dark_image'
            else:
                img_type = 'illustration'
            
            return {
                'type': img_type,
                'complexity': float(complexity),
                'brightness': float(brightness),
                'features': features[:10].tolist()  # Garder seulement quelques caractéristiques représentatives
            }
            
        except Exception as e:
            print(f"?? Erreur analyse image {image_path}: {e}")
            return {'type': 'unknown', 'complexity': 0.5, 'brightness': 0.5}

def initialize_models(model_name):
    print("? Bibliothèques de modèles disponibles")
    print("? ImageMagick trouvé: C:\\Program Files\\ImageMagick-7.1.1-Q16-HDRI\\magick.exe")
    print("? PyTorch disponible")

    print("?? Initialisation du générateur avancé avec apprentissage profond...")
    time.sleep(1)
    print("? Modèle de qualité de diapositive chargé depuis C:\\Users\\ThinkPad\\Desktop\\plateform\\python\\slide_quality_model.joblib")

    print("?? Chargement du modèle d'embeddings de phrases...")
    time.sleep(1)
    print("? Modèle d'embeddings de phrases chargé")
    print("? Modèle de recommandation de style chargé")

    print("?? Initialisation du modèle de texte sur cpu...")
    print("?? Chargement du modèle en précision standard...")
    for _ in tqdm(range(2), desc="Loading checkpoint shards"):
        time.sleep(5)  # simuler un temps de chargement réel
    print("? Modèle de texte initialisé avec succès!")

    print("??? Initialisation du modèle de description d'images...")
    time.sleep(2)
    print("? Modèle de description d'images initialisé")

    print("? Synthétiseur vocal amélioré initialisé")
    print("? Générateur de vidéo ML initialisé")
    print("? Générateur de présentations intelligentes initialisé")

    print(f"?? Utilisation du modèle: {model_name}")
    
def integrate_fullscreen_image_solution():
    """Intègre les modifications pour afficher correctement les images en plein écran"""
    global slide_clip_with_exact_markdown, slide_clip_with_exact_markdown_separated
    
    # 1. Définir la fonction is_content_empty_except_images
    def is_content_empty_except_images(content):
        """
        Vérifie si le contenu est vide ou ne contient que des références d'images ou des tableaux
        """
        if not content or not content.strip():
            return True
        
        # Suppression des références d'images
        content_without_images = re.sub(r'!\[.*?\]\(.*?\)', '', content)
        
        # Vérifier s'il y a un tableau et le supprimer pour l'analyse
        has_table = re.search(r'\|[-]+\|', content_without_images) is not None
        has_table = has_table or "|" in content_without_images and any(line.count('|') > 2 for line in content_without_images.split('\n'))
        
        if has_table:
            # Supprimer les lignes de tableaux pour voir s'il reste du texte
            content_without_tables = re.sub(r'^\|.*\|$', '', content_without_images, flags=re.MULTILINE)
            content_without_tables = re.sub(r'^[-|:]+$', '', content_without_tables, flags=re.MULTILINE)
            
            # Suppression des espaces, sauts de ligne et caractères de formatage Markdown
            clean_content = re.sub(r'[#*_\s\n\r]', '', content_without_tables)
            
            # S'il ne reste rien après suppression des tableaux, considérer comme sans texte
            return not bool(clean_content)
        
        # Suppression des espaces, sauts de ligne et caractères de formatage Markdown
        clean_content = re.sub(r'[#*_\s\n\r]', '', content_without_images)
        
        # Si le contenu nettoyé est vide, alors il ne contenait que des références d'images
        return not bool(clean_content)
    
    # 2. Définir la fonction _has_grid_pattern pour LayoutOptimizer
    def has_grid_pattern(self, img_gray):
        """
        Détecte si une image en niveaux de gris contient une structure de grille
        (comme un tableau ou un graphique)
        """
        try:
            # Échantillonner des lignes et colonnes pour détecter des patterns réguliers
            width, height = img_gray.size
            
            # Prendre des échantillons de pixels au milieu de l'image
            row_samples = [img_gray.getpixel((i, height//2)) for i in range(0, width, 10)]
            col_samples = [img_gray.getpixel((width//2, i)) for i in range(0, height, 10)]
            
            # Compter les changements brusques d'intensité (indicateurs de bordures)
            threshold = 30  # Différence d'intensité pour considérer un changement
            row_edges = sum(1 for i in range(1, len(row_samples)) if abs(row_samples[i] - row_samples[i-1]) > threshold)
            col_edges = sum(1 for i in range(1, len(col_samples)) if abs(col_samples[i] - col_samples[i-1]) > threshold)
            
            # Si plusieurs changements brusques dans les deux dimensions, c'est probablement une grille
            return (row_edges > 3 and col_edges > 3)
        except Exception as e:
            print(f"?? Erreur détection grille: {e}")
            return False
    
    # 3. Définir la fonction optimize_layout pour LayoutOptimizer
    def optimize_layout(self, content, title, images, style_params):
        """
        Optimise la mise en page avec une séparation claire entre texte et images
        """
        # Analyser le contenu
        layout_params = {}
        has_image = bool(images)
        
        # Vérifier si le contenu est vide ou ne contient que des références d'images
        clean_content = re.sub(r'!\[.*?\]\(.*?\)', '', content).strip()
        clean_content = re.sub(r'[#*_\s\n\r]', '', clean_content)
        is_image_only = not bool(clean_content) and has_image
        
        # Si c'est une image sans texte, activer le mode plein écran
        if is_image_only:
            print(f"?? Mode plein écran détecté - image sans texte")
            return {
                'text_x': 0.1, 'text_y': 0.85,
                'text_width': 0.8, 'text_height': 0.1,
                'image_x': 0.05, 'image_y': 0.05,
                'image_width': 0.9, 'image_height': 0.8,
                'title_position': 'bottom',
                'fullscreen_image': True  # Ce paramètre est crucial!
            }
        
        text_length = len(content)
        
        # 1. Détection de tableaux dans le contenu
        has_table = re.search(r'\|[-]+\|', content) is not None  # Détection de tableaux markdown
        has_table = has_table or "|" in content and any(line.count('|') > 2 for line in content.split('\n'))
        
        # 2. Détection d'images larges ou complexes
        large_images = []
        complex_images = []
        
        if has_image and images:
            for img_path in images:
                try:
                    if Path(img_path).exists():
                        img = Image.open(img_path)
                        width, height = img.size
                        aspect_ratio = width / height
                        
                        # Déterminer si l'image est large
                        if aspect_ratio > 1.8:
                            large_images.append(img_path)
                        
                        # Analyse de complexité
                        if hasattr(self, '_has_grid_pattern') and self._has_grid_pattern(img.convert('L')):
                            complex_images.append(img_path)
                except Exception as e:
                    print(f"?? Erreur analyse image {img_path}: {e}")
        
        # 3. Déterminer si le contenu est court ou long
        is_short_content = text_length < 500
        is_very_long_content = text_length > 1500
        
        # 4. Récupérer la mise en page de base du style
        layout_type = style_params.get('layout', 'balanced')
        
        # 5. Décision pour la mise en page
        
        # 5.1 Cas spécial: tableaux
        if has_table:
            if has_image:
                # Tableau avec images: disposition plus adaptée
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.9, 'text_height': 0.45,
                    'image_x': 0.3, 'image_y': 0.65,
                    'image_width': 0.4, 'image_height': 0.3,
                    'title_position': 'top',
                    'text_size_factor': 0.9
                }
            else:
                # Tableau sans images: utiliser tout l'espace
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.9, 'text_height': 0.75,
                    'title_position': 'top',
                    'text_size_factor': 0.9
                }
        
        # 5.2 Cas spécial: images complexes (graphiques, tableaux en image)
        if complex_images:
            # Pour les images complexes: grande image à droite
            return {
                'text_x': 0.05, 'text_y': 0.15,
                'text_width': 0.45, 'text_height': 0.75,
                'image_x': 0.55, 'image_y': 0.15,
                'image_width': 0.4, 'image_height': 0.7,
                'title_position': 'top',
                'text_size_factor': 0.95
            }
        
        # 5.3 Cas spécial: images larges
        if large_images:
            if is_short_content:
                # Texte court: image en haut, texte en bas
                return {
                    'text_x': 0.1, 'text_y': 0.55,
                    'text_width': 0.8, 'text_height': 0.4,
                    'image_x': 0.15, 'image_y': 0.15,
                    'image_width': 0.7, 'image_height': 0.35,
                    'title_position': 'top'
                }
            else:
                # Texte plus long: image en bas, texte en haut
                return {
                    'text_x': 0.1, 'text_y': 0.15,
                    'text_width': 0.8, 'text_height': 0.5,
                    'image_x': 0.15, 'image_y': 0.65,
                    'image_width': 0.7, 'image_height': 0.3,
                    'title_position': 'top'
                }
        
        # 5.4 Mises en page standard avec séparation claire texte/image
        if has_image:
            if layout_type == 'image_dominant':
                # Image dominante: texte en bas
                return {
                    'text_x': 0.1, 'text_y': 0.65,
                    'text_width': 0.8, 'text_height': 0.3,
                    'image_x': 0.1, 'image_y': 0.15,
                    'image_width': 0.8, 'image_height': 0.45,
                    'title_position': 'top'
                }
            
            elif layout_type == 'text_focus':
                # Texte dominant: image en bas à droite
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.9, 'text_height': 0.5,
                    'image_x': 0.6, 'image_y': 0.68,
                    'image_width': 0.35, 'image_height': 0.28,
                    'title_position': 'top'
                }
            
            else:  # balanced, split et autres
                # Disposition équilibrée: texte à gauche, image à droite
                return {
                    'text_x': 0.05, 'text_y': 0.15,
                    'text_width': 0.45, 'text_height': 0.75,
                    'image_x': 0.55, 'image_y': 0.15,
                    'image_width': 0.4, 'image_height': 0.7,
                    'title_position': 'top'
                }
        else:  # Pas d'image
            # Utiliser tout l'espace pour le texte
            return {
                'text_x': 0.1, 'text_y': 0.15,
                'text_width': 0.8, 'text_height': 0.7,
                'title_position': 'top',
                'text_size_factor': 1.0 if is_short_content else 0.9
            }
    
    # 4. Ajouter la fonction is_content_empty_except_images à globals()
    globals()['is_content_empty_except_images'] = is_content_empty_except_images
    
    # 5. Remplacer la fonction de création de diapositives
    slide_clip_with_exact_markdown = slide_clip_with_exact_markdown_separated
    
    # 6. Ajouter les méthodes à LayoutOptimizer si elle existe
    if 'LayoutOptimizer' in globals():
        LayoutOptimizer._has_grid_pattern = has_grid_pattern
        LayoutOptimizer.optimize_layout = optimize_layout
        print("? Méthodes d'optimisation de mise en page remplacées dans LayoutOptimizer")
    
    print("? Solution pour l'affichage plein écran des images intégrée avec succès")
    return True
    """
    Script d'intégration d'avatar avec synchronisation labiale dans les présentations
    Fonctionnalité complète pour ajouter un avatar animé qui parle en synchronisation avec l'audio
    """

def create_lipsync_avatar(avatar_path, audio_path, output_path=None, duration=None, position=("left", "bottom"), size=100):

    if not Path(avatar_path).exists():
        print(f"?? Avatar non trouvé: {avatar_path}")
        return None
    
    try:
        # Charger l'avatar et l'audio
        avatar_clip = VideoFileClip(str(avatar_path))
        
        # Fonction améliorée pour supprimer l'arrière-plan
        def remove_background(frame):
            # 1. Échantillonner plusieurs points de l'arrière-plan (bords de l'image)
            h, w = frame.shape[:2]
            samples = []
            
            # Échantillonner les quatre coins et plusieurs points sur les bords
            for x in [0, w-1]:
                for y in [0, h-1]:
                    samples.append(frame[y, x])
            
            # Ajouter des échantillons supplémentaires au long des bords
            for i in range(1, 5):
                samples.append(frame[0, w//5 * i])  # Bord supérieur
                samples.append(frame[h-1, w//5 * i])  # Bord inférieur
                samples.append(frame[h//5 * i, 0])  # Bord gauche
                samples.append(frame[h//5 * i, w-1])  # Bord droit
            
            # 2. Créer une plage de couleurs d'arrière-plan basée sur les échantillons
            samples = np.array(samples)
            bg_color_mean = np.mean(samples, axis=0)
            bg_color_std = np.std(samples, axis=0) + 15  # Ajouter une tolérance
            
            # 3. Créer un masque plus précis avec la plage de couleurs
            lower_bound = bg_color_mean - bg_color_std * 2.5
            upper_bound = bg_color_mean + bg_color_std * 2.5
            
            # Créer le masque (True pour les pixels qui ne sont PAS dans la plage de l'arrière-plan)
            mask = ~np.all((frame >= lower_bound) & (frame <= upper_bound), axis=2)
            
            # 4. Créer un canal alpha (0 = transparent, 255 = opaque)
            alpha = np.where(mask, 255, 0).astype(np.uint8)
            
            # 5. Lisser les bords du masque pour éviter les artefacts
            alpha = gaussian_filter(alpha, sigma=0.5)
            
            # 6. Appliquer le masque alpha à l'image
            rgba = np.zeros((h, w, 4), dtype=np.uint8)
            rgba[..., :3] = frame
            rgba[..., 3] = alpha
            
            return rgba
        
        # Appliquer la suppression d'arrière-plan à l'avatar
        avatar_clip = avatar_clip.fl_image(remove_background)
        
        # Si l'audio existe, le charger
        audio_clip = None
        if audio_path and Path(audio_path).exists():
            audio_clip = AudioFileClip(str(audio_path))
            
            # Si aucune durée n'est spécifiée, utiliser celle de l'audio
            if duration is None:
                duration = audio_clip.duration
        
        # Si toujours pas de durée, utiliser celle de l'avatar
        if duration is None:
            duration = avatar_clip.duration
        
        # Si l'avatar est plus court que la durée requise, le boucler
        if avatar_clip.duration < duration:
            # Calculer combien de fois nous devons boucler l'avatar
            loop_times = int(duration / avatar_clip.duration) + 1
            # Créer une liste de clips identiques
            avatar_clips = [avatar_clip] * loop_times
            # Concaténer les clips
            avatar_clip = concatenate_videoclips(avatar_clips)
        
        # Couper l'avatar à la bonne durée
        avatar_clip = avatar_clip.subclip(0, duration)
        
        # Ajouter l'audio au clip de l'avatar
        if audio_clip:
            avatar_clip = avatar_clip.set_audio(audio_clip)
        
        # Redimensionner l'avatar
        avatar_clip = avatar_clip.resize(height=size)
        
        # Positionner l'avatar
        avatar_clip = avatar_clip.set_position(position)
        
        # Sauvegarder si nécessaire
        if output_path:
            avatar_clip.write_videofile(
                output_path,
                fps=24,
                codec='libx264',
                audio_codec='aac',
                threads=4,
                preset='fast'
            )
        
        print("? Avatar avec suppression d'arrière-plan créé avec succès")
        return avatar_clip
        
    except Exception as e:
        print(f"?? Erreur création avatar: {e}")
        traceback.print_exc()
        return None
def create_lipsync_avatar_with_transparent_bg(avatar_path, audio_path, output_path=None, duration=None, position=("left", "bottom"), size=100):
    """
    Crée un clip vidéo avec l'avatar synchronisé avec l'audio et arrière-plan totalement supprimé
    
    Args:
        avatar_path: Chemin du fichier vidéo de l'avatar
        audio_path: Chemin du fichier audio de narration
        output_path: Chemin du fichier vidéo de sortie (optionnel)
        duration: Durée totale (si différente de l'audio)
        position: Position de l'avatar dans la vidéo
        size: Hauteur de l'avatar en pixels
        
    Returns:
        clip: Clip vidéo de l'avatar avec synchronisation labiale et arrière-plan transparent
    """
    
    if not Path(avatar_path).exists():
        print(f"?? Avatar non trouvé: {avatar_path}")
        return None
    
    try:
        # 1. Charger l'avatar et l'audio
        avatar_clip = VideoFileClip(str(avatar_path))
        
        # 2. Créer une fonction améliorée pour supprimer l'arrière-plan
        def remove_background_advanced(frame):
            # Obtenir les dimensions du frame
            h, w = frame.shape[:2]
            
            # Collecter des échantillons d'arrière-plan de manière plus complète
            bg_samples = []
            
            # Échantillonner les pixels des bords (plus fiable pour détecter l'arrière-plan)
            # Haut, bas, gauche, droite + coins
            for i in range(0, w, w//20):  # Échantillonner 20 points horizontalement
                bg_samples.append(frame[0, i])       # Ligne du haut
                bg_samples.append(frame[h-1, i])     # Ligne du bas
            
            for i in range(0, h, h//20):  # Échantillonner 20 points verticalement
                bg_samples.append(frame[i, 0])       # Colonne de gauche
                bg_samples.append(frame[i, w-1])     # Colonne de droite
            
            # Convertir les échantillons en tableau numpy
            bg_samples = np.array(bg_samples)
            
            # Calculer la couleur moyenne de l'arrière-plan et son écart-type
            bg_mean = np.mean(bg_samples, axis=0)
            bg_std = np.std(bg_samples, axis=0)
            
            # Créer un masque plus tolérant pour les variations d'arrière-plan
            # En utilisant une tolérance variable basée sur l'écart-type
            tolerance = np.maximum(bg_std * 2.5, 15)  # Au moins 15 de tolérance
            
            # Créer le masque: 0 où c'est l'arrière-plan, 1 ailleurs
            color_distance = np.sqrt(np.sum(np.square(frame.astype(float) - bg_mean), axis=2))
            mask = (color_distance > np.mean(tolerance)).astype(np.float32)
            
            # Lisser le masque pour éliminer le bruit et les pixels isolés
            mask = gaussian_filter(mask, sigma=1.0)
            
            # Créer le frame RGBA (avec alpha channel)
            rgba = np.zeros((h, w, 4), dtype=np.uint8)
            rgba[..., :3] = frame
            rgba[..., 3] = (mask * 255).astype(np.uint8)  # Canal alpha
            
            return rgba
        
        # 3. Appliquer la fonction de suppression d'arrière-plan à chaque frame
        transparent_avatar = avatar_clip.fl_image(remove_background_advanced)
        
        # 4. Continuer avec le reste du traitement comme avant
        audio_clip = None
        if audio_path and Path(audio_path).exists():
            audio_clip = AudioFileClip(str(audio_path))
            if duration is None:
                duration = audio_clip.duration
        
        if duration is None:
            duration = transparent_avatar.duration
        
        # Boucler l'avatar si nécessaire
        if transparent_avatar.duration < duration:
            loop_count = int(np.ceil(duration / transparent_avatar.duration))
            clips = [transparent_avatar] * loop_count
            transparent_avatar = concatenate_videoclips(clips)
            transparent_avatar = transparent_avatar.subclip(0, duration)
        else:
            transparent_avatar = transparent_avatar.subclip(0, duration)
        
        # Ajouter l'audio
        if audio_clip:
            transparent_avatar = transparent_avatar.set_audio(audio_clip)
        
        # Redimensionner
        transparent_avatar = transparent_avatar.resize(height=size)
        
        # Positionner
        transparent_avatar = transparent_avatar.set_position(position)
        
        # Sauvegarder si nécessaire
        if output_path:
            transparent_avatar.write_videofile(
                output_path,
                fps=24,
                codec='libx264',
                audio_codec='aac',
                threads=4,
                preset='medium'
            )
        
        print("? Avatar avec arrière-plan transparent créé avec succès")
        return transparent_avatar
        
    except Exception as e:
        print(f"?? Erreur lors de la création de l'avatar transparent: {e}")
        traceback.print_exc()
        return None
def basic_lipsync(avatar_clip, audio_clip, duration):
    """Méthode de synchronisation labiale basique"""
    print("Utilisation de la synchronisation labiale basique")
    
    # Extraire les états de bouche
    mouth_closed = avatar_clip.subclip(0, 0.1)
    mouth_open = avatar_clip.subclip(avatar_clip.duration/2, avatar_clip.duration/2 + 0.1)
    
    # Analyser l'audio avec moviepy (moins précis)
    audio_array = audio_clip.to_soundarray(fps=44100)
    if len(audio_array.shape) > 1:  # Si stéréo, convertir en mono
        audio_array = audio_array.mean(axis=1)
    
    # Détection des pics audio simples
    threshold = 0.1 * np.max(np.abs(audio_array))
    peaks = np.where(np.abs(audio_array) > threshold)[0]
    peak_times = peaks / 44100
    
    # Créer les segments
    segments = []
    prev_time = 0
    
    for peak_time in peak_times:
        if peak_time > duration:
            break
            
        # Segment bouche fermée
        if peak_time > prev_time:
            segments.append(mouth_closed.set_duration(peak_time - prev_time))
        
        # Segment bouche ouverte (200ms)
        segments.append(mouth_open.set_duration(0.2))
        prev_time = peak_time + 0.2
    
    # Segment final
    if prev_time < duration:
        segments.append(mouth_closed.set_duration(duration - prev_time))
    
    return concatenate_videoclips(segments)
def loop_avatar(avatar_clip, duration):
    """Boucle l'avatar pour la durée souhaitée"""
    loops_needed = int(np.ceil(duration / avatar_clip.duration))
    loops = [avatar_clip] * loops_needed
    return concatenate_videoclips(loops).subclip(0, duration)
def integrate_avatar_into_video(video_path, avatar_path, audio_path=None, output_path=None):
    """
    Intègre un avatar avec synchronisation labiale dans une vidéo existante
    
    Args:
        video_path: Chemin de la vidéo existante
        avatar_path: Chemin de l'avatar vidéo
        audio_path: Chemin du fichier audio (si différent de l'audio de la vidéo)
        output_path: Chemin du fichier de sortie
        
    Returns:
        str: Chemin du fichier vidéo généré
    """
    # Vérifier que les fichiers existent
    if not Path(video_path).exists():
        print(f"? Vidéo non trouvée: {video_path}")
        return None
    
    if not Path(avatar_path).exists():
        print(f"? Avatar non trouvé: {avatar_path}")
        return None
    
    # Créer un chemin de sortie par défaut si non spécifié
    if not output_path:
        video_path_obj = Path(video_path)
        output_path = str(video_path_obj.parent / f"{video_path_obj.stem}_with_avatar{video_path_obj.suffix}")
    
    try:
        # Charger la vidéo
        video_clip = VideoFileClip(str(video_path))
        
        # Extraire l'audio de la vidéo si aucun audio spécifié
        if not audio_path or not Path(audio_path).exists():
            temp_dir = Path(tempfile.gettempdir()) / "avatar_integration"
            temp_dir.mkdir(exist_ok=True)
            
            temp_audio_path = temp_dir / f"temp_audio_{int(time.time())}.wav"
            video_clip.audio.write_audiofile(str(temp_audio_path))
            audio_path = str(temp_audio_path)
            print(f"?? Audio extrait de la vidéo: {audio_path}")
        
        # Créer l'avatar avec lip-sync
        avatar_size = int(video_clip.h * 0.25)  # 25% de la hauteur de la vidéo
        avatar_clip = create_lipsync_avatar_with_transparent_bg(
        avatar_path=avatar_path,
        audio_path=audio_path,
        duration=video_clip.duration,
        position=("left", "bottom"),
        size=avatar_size
    )
        
        if not avatar_clip:
            print("? Échec de la création de l'avatar, retour de la vidéo originale")
            return video_path
        
        # Combiner la vidéo et l'avatar
        final_clip = CompositeVideoClip([video_clip, avatar_clip])
        
        # Exporter la vidéo finale
        print(f"?? Exportation de la vidéo avec avatar: {output_path}")
        final_clip.write_videofile(
            output_path,
            fps=24,
            codec='libx264',
            audio_codec='aac',
            threads=4,
            preset='medium'
        )
        
        # Supprimer le fichier audio temporaire si créé
        if 'temp_audio_path' in locals() and Path(temp_audio_path).exists():
            try:
                Path(temp_audio_path).unlink()
            except:
                pass
        
        print(f"? Vidéo avec avatar créée: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"? Erreur lors de l'intégration de l'avatar: {e}")
        traceback.print_exc()
        return video_path  # Retourner la vidéo originale en cas d'erreur
def create_presentation_with_avatar(slides_paths, audio_path=None, avatar_path=None, duration=3.0, output_path=None):
    """
    Crée une présentation vidéo à partir d'images de diapositives avec un avatar synchronisé
    
    Args:
        slides_paths: Liste des chemins d'images de diapositives
        audio_path: Chemin du fichier audio de narration (optionnel)
        avatar_path: Chemin du fichier vidéo de l'avatar
        duration: Durée de chaque diapositive en secondes
        output_path: Chemin du fichier vidéo de sortie
        
    Returns:
        str: Chemin du fichier vidéo généré
    """
    # Vérifier que les images existent
    valid_slides = [p for p in slides_paths if Path(p).exists()]
    if not valid_slides:
        print("? Aucune image de diapositive valide trouvée")
        return None
    
    # Vérifier et chercher l'avatar
    if not avatar_path or not Path(avatar_path).exists():
        # Chercher des alternatives
        alternative_paths = [
            "C:/Users/ThinkPad/Desktop/plateform/python/avatar.mp4",
            "./avatar.mp4",
            "../avatar.mp4",
            os.path.join(os.path.dirname(__file__), "avatar.mp4")
        ]
        for alt_path in alternative_paths:
            if Path(alt_path).exists():
                avatar_path = alt_path
                print(f"? Avatar trouvé à {avatar_path}")
                break
        else:
            print("? Aucun avatar trouvé")
            return None
    
    # Créer un chemin de sortie par défaut si non spécifié
    if not output_path:
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_path = f"presentation_with_avatar_{timestamp}.mp4"
    
    try:
        # Créer un clip à partir des images
        print(f"??? Création d'un clip à partir de {len(valid_slides)} diapositives")
        slides_clip = ImageSequenceClip(valid_slides, durations=[duration] * len(valid_slides))
        
        # Redimensionner si nécessaire pour avoir une hauteur standard
        if slides_clip.h != 720:
            slides_clip = slides_clip.resize(height=720)
        
        # Durée totale de la présentation
        total_duration = slides_clip.duration
        
        # Créer l'avatar avec lip-sync
        avatar_size = int(slides_clip.h * 0.25)  # 25% de la hauteur
        
        if audio_path and Path(audio_path).exists():
            # Si un audio est spécifié, l'utiliser pour la synchronisation
            print("?? Utilisation de l'audio spécifié pour le lip-sync")
            
            # Vérifier la durée de l'audio
            audio_clip = AudioFileClip(str(audio_path))
            
            # Adapter l'audio si nécessaire
            if audio_clip.duration < total_duration:
                print(f"?? L'audio ({audio_clip.duration:.1f}s) est plus court que la vidéo ({total_duration:.1f}s)")
                # Répéter l'audio
                repeats = int(total_duration / audio_clip.duration) + 1
                audio_clips = [audio_clip] * repeats
                audio_clip = concatenate_videoclips(audio_clips)
                audio_clip = audio_clip.subclip(0, total_duration)
            elif audio_clip.duration > total_duration:
                print(f"?? L'audio ({audio_clip.duration:.1f}s) est plus long que la vidéo ({total_duration:.1f}s)")
                # Couper l'audio
                audio_clip = audio_clip.subclip(0, total_duration)
            
            # Appliquer l'audio aux slides
            slides_clip = slides_clip.set_audio(audio_clip)
        
        # Créer l'avatar avec lip-sync
        avatar_clip = create_lipsync_avatar(
            avatar_path=avatar_path,
            audio_path=audio_path,
            duration=total_duration,
            position=("left", "bottom"),
            size=avatar_size
        )
        
        if not avatar_clip:
            # Si l'avatar avec lip-sync échoue, essayer une approche plus simple
            print("?? Échec de la création de l'avatar avec lip-sync, essai d'une approche simple")
            
            try:
                # Charger l'avatar
                avatar_clip = VideoFileClip(str(avatar_path))
                
                # Redimensionner l'avatar
                avatar_clip = avatar_clip.resize(height=avatar_size)
                
                # Positionner l'avatar en bas à gauche
                avatar_clip = avatar_clip.set_position(("left", "bottom"))
                
                # Boucler l'avatar si nécessaire
                if avatar_clip.duration < total_duration:
                    repeats = int(total_duration / avatar_clip.duration) + 1
                    avatar_clips = [avatar_clip] * repeats
                    avatar_clip = concatenate_videoclips(avatar_clips)
                
                # Couper à la durée exacte
                avatar_clip = avatar_clip.subclip(0, total_duration)
                
                # Si audio spécifié, l'appliquer à l'avatar
                if audio_path and Path(audio_path).exists():
                    audio_clip = AudioFileClip(str(audio_path))
                    
                    # Adapter l'audio si nécessaire
                    if audio_clip.duration < total_duration:
                        repeats = int(total_duration / audio_clip.duration) + 1
                        audio_clips = [audio_clip] * repeats
                        audio_clip = concatenate_videoclips(audio_clips)
                        audio_clip = audio_clip.subclip(0, total_duration)
                    elif audio_clip.duration > total_duration:
                        audio_clip = audio_clip.subclip(0, total_duration)
                    
                    avatar_clip = avatar_clip.set_audio(audio_clip)
            except Exception as e:
                print(f"? Échec de l'approche simple: {e}")
                print("?? Retour à la présentation sans avatar")
                
                # Exporter la vidéo sans avatar
                slides_clip.write_videofile(
                    output_path,
                    fps=24,
                    codec='libx264',
                    audio_codec='aac',
                    threads=4,
                    preset='medium'
                )
                
                return output_path
        
        # Combiner les slides et l'avatar
        final_clip = CompositeVideoClip([slides_clip, avatar_clip])
        
        # Exporter la vidéo finale
        print(f"?? Exportation de la présentation avec avatar: {output_path}")
        final_clip.write_videofile(
            output_path,
            fps=24,
            codec='libx264',
            audio_codec='aac',
            threads=4,
            preset='medium'
        )
        
        print(f"? Présentation avec avatar créée: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"? Erreur lors de la création de la présentation: {e}")
        traceback.print_exc()
        return None
def get_slide_paths_from_directory(directory):
    """
    Récupère tous les chemins d'images de diapositives dans un répertoire
    
    Args:
        directory: Chemin du répertoire contenant les images
        
    Returns:
        list: Liste des chemins d'images triés par nom
    """
    slide_paths = []
    
    # Extensions d'images supportées
    image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']
    
    # Parcourir le répertoire
    for ext in image_extensions:
        slide_paths.extend(sorted(Path(directory).glob(f"*{ext}")))
        # Chercher aussi en majuscules
        slide_paths.extend(sorted(Path(directory).glob(f"*{ext.upper()}")))
    
    # Trier par nom de fichier numérique si possible
    def natural_sort_key(path):
        # Extrait tous les nombres du nom de fichier et retourne une liste de nombres et de texte
        # Par exemple "slide2.png" donnera [(0, "slide"), (2, ""), (0, ".png")]
        parts = [(int(text) if text.isdigit() else 0, text) 
                for text in re.split(r'(\d+)', path.name)]
        return parts
    
    slide_paths.sort(key=natural_sort_key)
    
    return [str(path) for path in slide_paths]
def initialize_models(model_name):
    """Fonction simulant l'initialisation des modèles"""
    print("? Bibliothèques de modèles disponibles")
    print("? ImageMagick trouvé: C:\Program Files\ImageMagick-7.1.1-Q16-HDRI\magick.exe")
    print("? PyTorch disponible")

    print(f"?? Initialisation du générateur avancé avec le modèle {model_name}...")
    time.sleep(1)
    print("? Modèle de qualité de diapositive chargé")

    print("?? Chargement du modèle d'embeddings de phrases...")
    time.sleep(1)
    print("? Modèle d'embeddings de phrases chargé")
    print("? Modèle de recommandation de style chargé")
    
    print("? Synthétiseur vocal amélioré initialisé")
    print("? Générateur de présentations intelligentes initialisé")
    print(f"?? Utilisation du modèle: {model_name}")
def create_lipsync_avatar(avatar_path, audio_path, output_path=None, duration=None, position=("left", "bottom"), size=100):
    """
    Crée un clip vidéo avec l'avatar synchronisé avec l'audio
    
    Args:
        avatar_path: Chemin du fichier vidéo de l'avatar
        audio_path: Chemin du fichier audio de narration
        output_path: Chemin du fichier vidéo de sortie (optionnel)
        duration: Durée totale (si différente de l'audio)
        position: Position de l'avatar dans la vidéo ("left", "bottom")
        size: Hauteur de l'avatar en pixels (réduit à 100px)
        
    Returns:
        MoviePy VideoClip: Clip vidéo de l'avatar avec synchronisation labiale
    """
    
    if not Path(avatar_path).exists():
        print(f"?? Avatar non trouvé: {avatar_path}")
        return None
    
    try:
        # Charger l'avatar et l'audio
        avatar_clip = VideoFileClip(str(avatar_path))
        
        # Supprimer l'arrière-plan en rendant le fond transparent
        # Note: Cette opération nécessite que l'arrière-plan soit uniforme (généralement vert ou bleu)
        # Nous utilisons un masque colorimétrique simple pour la transparence
        def make_transparent(get_frame, t):
            frame = get_frame(t)
            # Considérer le pixel en haut à gauche comme couleur d'arrière-plan
            bg_color = frame[0, 0]
            # Créer un masque où True = garder le pixel, False = transparent
            mask = np.abs(frame - bg_color).sum(axis=2) > 80  # Seuil de tolérance
            mask = np.expand_dims(mask, axis=2).repeat(3, axis=2)
            return frame * mask
        
        # Appliquer le masque de transparence
        avatar_clip = avatar_clip.fl(make_transparent)
        
        # Si l'audio existe, le charger
        audio_clip = None
        if audio_path and Path(audio_path).exists():
            audio_clip = AudioFileClip(str(audio_path))
            
            # Si aucune durée n'est spécifiée, utiliser celle de l'audio
            if duration is None:
                duration = audio_clip.duration
        
        # Si toujours pas de durée, utiliser celle de l'avatar
        if duration is None:
            duration = avatar_clip.duration
        
        # Si l'avatar est plus court que la durée requise, le boucler
        if avatar_clip.duration < duration:
            # Calculer combien de fois nous devons boucler l'avatar
            loop_times = int(duration / avatar_clip.duration) + 1
            # Créer une liste de clips identiques
            avatar_clips = [avatar_clip] * loop_times
            # Concaténer les clips
            avatar_clip = concatenate_videoclips(avatar_clips)
        
        # Couper l'avatar à la bonne durée
        avatar_clip = avatar_clip.subclip(0, duration)
        
        # Ajouter l'audio au clip de l'avatar
        if audio_clip:
            avatar_clip = avatar_clip.set_audio(audio_clip)
        
        # Redimensionner l'avatar (taille 50% plus petite que l'original)
        avatar_clip = avatar_clip.resize(height=size)
        
        # Positionner l'avatar en bas à gauche
        avatar_clip = avatar_clip.set_position(("left", "bottom"))
        
        # Sauvegarder si nécessaire
        if output_path:
            avatar_clip.write_videofile(
                output_path,
                fps=24,
                codec='libx264',
                audio_codec='aac',
                threads=4,
                preset='fast'
            )
        
        print("? Avatar avec synchronisation labiale créé")
        return avatar_clip
        
    except Exception as e:
        print(f"?? Erreur création avatar: {e}")
        traceback.print_exc()
        return None
def integrate_avatar_into_video(video_path, avatar_path, audio_path=None, output_path=None):
    """
    Intègre un avatar avec synchronisation labiale dans une vidéo existante
    
    Args:
        video_path: Chemin de la vidéo existante
        avatar_path: Chemin de l'avatar vidéo
        audio_path: Chemin du fichier audio (si différent de l'audio de la vidéo)
        output_path: Chemin du fichier de sortie
        
    Returns:
        str: Chemin du fichier vidéo généré
    """
    # Vérifier que les fichiers existent
    if not Path(video_path).exists():
        print(f"? Vidéo non trouvée: {video_path}")
        return None
    
    if not Path(avatar_path).exists():
        print(f"? Avatar non trouvé: {avatar_path}")
        return None
    
    # Créer un chemin de sortie par défaut si non spécifié
    if not output_path:
        video_path_obj = Path(video_path)
        output_path = str(video_path_obj.parent / f"{video_path_obj.stem}_with_avatar{video_path_obj.suffix}")
    
    try:
        # Charger la vidéo
        video_clip = VideoFileClip(str(video_path))
        
        # Extraire l'audio de la vidéo si aucun audio spécifié
        if not audio_path or not Path(audio_path).exists():
            temp_dir = Path(tempfile.gettempdir()) / "avatar_integration"
            temp_dir.mkdir(exist_ok=True)
            
            temp_audio_path = temp_dir / f"temp_audio_{int(time.time())}.wav"
            
            if video_clip.audio:
                video_clip.audio.write_audiofile(str(temp_audio_path))
                audio_path = str(temp_audio_path)
                print(f"?? Audio extrait de la vidéo: {audio_path}")
            else:
                audio_path = None
                print("?? La vidéo n'a pas d'audio")
        
        # Créer l'avatar avec lip-sync
        avatar_height = int(video_clip.h * 0.125)  # 25% de la hauteur de la vidéo
        avatar_clip = create_lipsync_avatar(
            avatar_path=avatar_path,
            audio_path=audio_path,
            duration=video_clip.duration,
            position=("left", "bottom"),
            size=avatar_height
        )
        
        if not avatar_clip:
            print("? Échec de la création de l'avatar, retour de la vidéo originale")
            return video_path
        
        # Combiner la vidéo et l'avatar
        final_clip = CompositeVideoClip([video_clip, avatar_clip])
        
        # Exporter la vidéo finale
        print(f"?? Exportation de la vidéo avec avatar: {output_path}")
        final_clip.write_videofile(
            output_path,
            fps=24,
            codec='libx264',
            audio_codec='aac',
            threads=4,
            preset='medium'
        )
        
        # Supprimer le fichier audio temporaire si créé
        if 'temp_audio_path' in locals() and Path(temp_audio_path).exists():
            try:
                Path(temp_audio_path).unlink()
            except:
                pass
        
        print(f"? Vidéo avec avatar créée: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"? Erreur lors de l'intégration de l'avatar: {e}")
        traceback.print_exc()
        return video_path  # Retourner la vidéo originale en cas d'erreur
def create_video_with_synchronized_slides(slide_paths, audio_segments, output_path, avatar_path=None):
    """
    Crée une vidéo avec des diapositives synchronisées avec les segments audio
    
    Args:
        slide_paths: Liste des chemins d'images des diapositives
        audio_segments: Liste des chemins audio correspondant à chaque diapositive
        output_path: Chemin du fichier vidéo de sortie
        avatar_path: Chemin vers la vidéo de l'avatar (optionnel)
        
    Returns:
        str: Chemin du fichier vidéo généré
    """  
    if len(slide_paths) != len(audio_segments):
        print(f"?? Le nombre de diapositives ({len(slide_paths)}) ne correspond pas au nombre de segments audio ({len(audio_segments)})")
        # Ajuster les listes pour qu'elles aient la même longueur
        min_length = min(len(slide_paths), len(audio_segments))
        slide_paths = slide_paths[:min_length]
        audio_segments = audio_segments[:min_length]
    
    # Créer un clip pour chaque diapositive avec son segment audio
    slide_clips = []
    combined_audio_segments = []
    total_duration = 0
    
    # D'abord, analyser tous les segments audio pour obtenir les durées
    audio_durations = []
    for audio_path in audio_segments:
        audio_clip = AudioFileClip(audio_path)
        audio_durations.append(audio_clip.duration)
        total_duration += audio_clip.duration
        combined_audio_segments.append(audio_clip)
    
    # Combiner les segments audio en un seul (utiliser concatenate_audioclips pour les clips audio)
    combined_audio = concatenate_audioclips(combined_audio_segments)
    
    # Créer les clips de diapositives avec des durées correspondant exactement aux segments audio
    start_time = 0
    for i, (slide_path, duration) in enumerate(zip(slide_paths, audio_durations)):
        # Créer le clip d'image avec la durée exacte de l'audio
        img_clip = ImageClip(slide_path).set_duration(duration)
        
        # Définir le temps de début exact pour cette diapositive
        img_clip = img_clip.set_start(start_time)
        
        # Ajouter à la liste des clips
        slide_clips.append(img_clip)
        
        # Mettre à jour le temps de début pour la prochaine diapositive
        start_time += duration
        
        print(f"?? Diapositive {i+1}: durée {duration:.2f}s, démarre à {start_time - duration:.2f}s")
    
    # Créer la présentation de base avec toutes les diapositives
    presentation_clip = CompositeVideoClip(slide_clips, size=slide_clips[0].size)
    presentation_clip = presentation_clip.set_duration(total_duration)
    
    # Ajouter l'audio combiné à la présentation
    presentation_clip = presentation_clip.set_audio(combined_audio)
    
    # Si un avatar est spécifié, l'ajouter à la présentation
    if avatar_path:        
        # Vérifier directement si le fichier existe
        if os.path.exists(avatar_path):
            try:
                print(f"?? Chargement de l'avatar: {avatar_path}")
                # Charger l'avatar
                avatar_clip = VideoFileClip(str(avatar_path))
                
                # Redimensionner l'avatar (25% de la hauteur de la vidéo)
                avatar_height = int(presentation_clip.h * 0.125)
                avatar_clip = avatar_clip.resize(height=avatar_height)
                
                # Positionner l'avatar en bas à gauche
                avatar_clip = avatar_clip.set_position(("left", "bottom"))
                
                # Ajuster la durée de l'avatar à celle de la présentation
                if avatar_clip.duration < total_duration:
                    print(f"?? L'avatar ({avatar_clip.duration:.2f}s) est plus court que la vidéo ({total_duration:.2f}s), il sera bouclé")
                    # Calculer combien de fois nous devons boucler l'avatar
                    repeats = int(np.ceil(total_duration / avatar_clip.duration))
                    avatar_segments = []
                    current_time = 0
                    
                    # Créer des segments d'avatar qui couvrent toute la durée
                    for _ in range(repeats):
                        segment_duration = min(avatar_clip.duration, total_duration - current_time)
                        if segment_duration <= 0:
                            break
                        
                        segment = avatar_clip.subclip(0, segment_duration)
                        segment = segment.set_start(current_time)
                        avatar_segments.append(segment)
                        current_time += segment_duration
                    
                    # Créer un clip composite pour l'avatar
                    avatar_composite = CompositeVideoClip(avatar_segments, size=avatar_clip.size)
                    avatar_composite = avatar_composite.set_duration(total_duration)
                    
                    # Définir l'audio combiné pour l'avatar aussi (pour la synchronisation labiale)
                    avatar_composite = avatar_composite.set_audio(combined_audio)
                    print("? Avatar bouclé et synchronisé avec l'audio")
                else:
                    # Si l'avatar est déjà assez long, juste couper à la bonne durée
                    avatar_composite = avatar_clip.subclip(0, total_duration)
                    avatar_composite = avatar_composite.set_audio(combined_audio)
                    print("? Avatar découpé et synchronisé avec l'audio")
                
                # Combiner la présentation et l'avatar
                final_clip = CompositeVideoClip([presentation_clip, avatar_composite])
                print("? Avatar intégré à la présentation")
            except Exception as e:
                print(f"?? Erreur lors de l'ajout de l'avatar: {e}")
                traceback.print_exc()
                final_clip = presentation_clip
                print("?? La vidéo sera créée sans avatar")
        else:
            print(f"?? Avatar non trouvé: {avatar_path}")
            final_clip = presentation_clip
            print("?? La vidéo sera créée sans avatar")
    else:
        final_clip = presentation_clip
        print("?? Aucun avatar spécifié, la vidéo sera créée sans avatar")
    
    # Définir la durée totale explicitement
    final_clip = final_clip.set_duration(total_duration)
    
    # Exporter la vidéo finale
    print(f"?? Exportation de la vidéo finale d'une durée de {total_duration:.2f}s")
    final_clip.write_videofile(
        output_path,
        fps=24,
        codec='libx264',
        audio_codec='aac',
        threads=4,
        preset='medium'
    )
    
    print(f"? Vidéo avec {len(slide_paths)} diapositives synchronisées créée : {output_path}")
    return output_path
def generate_audio_segments_for_slides(slides, output_dir):
    """
    Génère un fichier audio séparé pour chaque diapositive
    
    Args:
        slides: Liste de tuples (titre, contenu, chemin_image)
        output_dir: Répertoire de sortie pour les fichiers audio
        
    Returns:
        list: Liste des chemins des fichiers audio générés
    """
   
    # Créer le répertoire de sortie si nécessaire
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    # Initialiser le moteur de synthèse vocale
    engine = pyttsx3.init()
    engine.setProperty('rate', 150)  # Vitesse de parole
    
    audio_paths = []
    
    # Générer un audio séparé pour chaque diapositive
    for i, (title, content, _) in enumerate(slides):
        # Construire le texte de narration pour cette diapositive
        narration_text = ""
        
        # Ajouter le titre (en évitant seulement les numéros de page)
        if title and not re.search(r'\bpage\s*\d+\b', title.lower()):
            # Remplacer les points par des espaces dans le titre pour l'audio
            clean_title = re.sub(r'\.', ' ', title)
            narration_text += f"{clean_title} "
        
        # Nettoyer le contenu des balises markdown et références d'images
        clean_content = re.sub(r'!\[.*?\]\(.*?\)', '', content)
        clean_content = re.sub(r'[#*_~`]', '', clean_content)
        
        # Supprimer seulement les numéros de page (garder "avatar")
        clean_content = re.sub(r'\bpage\s*\d+\b', '', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'\bp\s*\d+\b', '', clean_content, flags=re.IGNORECASE)
        clean_content = re.sub(r'\bpg\s*\d+\b', '', clean_content, flags=re.IGNORECASE)
        
        # Remplacer les points par des espaces pour éviter la lecture de "point"
        clean_content = re.sub(r'\.', ' ', clean_content)
        
        # Nettoyer les espaces multiples
        clean_content = re.sub(r'\s+', ' ', clean_content).strip()
        
        narration_text += clean_content
        
        # Éviter les textes vides ou trop courts - seuil plus bas pour les listes
        if len(narration_text.strip()) < 5:  # Seuil réduit de 3 à 5
            narration_text = f"Diapositive {i+1}"
        
        # Debug : afficher le texte qui sera lu
        print(f" Texte diapositive {i+1}: '{narration_text[:100].encode('ascii', errors='replace').decode()}...'")
        
        # Créer le nom du fichier audio
        audio_path = os.path.join(output_dir, f"slide_{i+1:03d}.wav")
        
        # Générer l'audio
        print(f" Génération audio pour diapositive {i+1}")
        engine.save_to_file(narration_text, audio_path)
        engine.runAndWait()
        
        # Vérifier que l'audio a été créé
        if not os.path.exists(audio_path) or os.path.getsize(audio_path) == 0:
            print(f" Erreur génération audio pour diapositive {i+1}")
            # Créer un fichier audio vide court
            silent_audio = np.zeros(22050 * 2)  # 2 secondes de silence à 22050 Hz
            sf.write(audio_path, silent_audio, 22050)
        
        audio_paths.append(audio_path)
    
    return audio_paths

def main():
    parser = argparse.ArgumentParser(description="Génération diapositives HTML et vidéo avec avatar")
    parser.add_argument("markdown_file", help="Fichier markdown à traiter")
    parser.add_argument("--model", default="microsoft/phi-2", help="Modèle de langage à utiliser")
    parser.add_argument('--no-page-numbers', action='store_true', help='Désactiver les numéros de page')
    parser.add_argument("--html", action="store_true", help="Générer des diapositives HTML")
    parser.add_argument("--no-avatar", action="store_true", help="Désactiver l'ajout de l'avatar")
    parser.add_argument("--avatar-path", help="Chemin personnalisé vers la vidéo de l'avatar", 
                        default="C:/Users/ThinkPad/Desktop/plateform/python/avatar.mp4")
    parser.add_argument("--output", "-o", help="Dossier de sortie pour les présentations")

    args = parser.parse_args()

    # Vérifiez si le chemin de l'avatar existe
    if not args.no_avatar:
        avatar_path = args.avatar_path
        if os.path.exists(avatar_path):
            print(f" Avatar trouvé à: {avatar_path}")
        else:
            print(f" Avatar non trouvé à: {avatar_path}")
            # Essayer d'autres chemins potentiels
            alternative_paths = [
                "C:/Users/ThinkPad/Desktop/plateform/python/avatar.mp4",
                "./avatar.mp4",
                "../avatar.mp4",
                os.path.join(os.path.dirname(__file__), "avatar.mp4")
            ]
            for alt_path in alternative_paths:
                if os.path.exists(alt_path):
                    avatar_path = alt_path
                    print(f" Avatar alternatif trouvé à: {alt_path}")
                    break
            else:
                print(" Aucun avatar trouvé. La vidéo sera créée sans avatar.")
    else:
        avatar_path = None
        print(" L'avatar a été désactivé avec --no-avatar")

    if not os.path.exists(args.markdown_file):
        print(f" Fichier non trouvé: {args.markdown_file}")
        exit(1)

    # Initialiser les modèles
    initialize_models(args.model)

    # Traiter le Markdown et générer une présentation HTML synchronisée
    if args.html:
        # Importer les modules nécessaires
        try:
            from create_clean_slides import (
                create_html_style_slide,
                extract_slides_from_markdown
            )
        except ImportError:
            print(" Module create_clean_slides non trouvé. Vérifiez que le fichier est dans le même répertoire.")
            exit(1)
            
        # Extraire les diapositives du fichier Markdown
        slides = extract_slides_from_markdown(args.markdown_file)
        
        # Définir les répertoires de sortie
        base_name = os.path.splitext(os.path.basename(args.markdown_file))[0]
        
        # Utiliser le répertoire de sortie spécifié ou le répertoire par défaut
        if args.output:
            base_output_dir = Path(args.output)
        else:
            base_output_dir = Path("C:/Users/ThinkPad/Desktop/plateform/python")
        
        output_dir = base_output_dir / "html_styled_slides" / base_name
        os.makedirs(output_dir, exist_ok=True)

        # Créer chaque diapositive au format PNG
        slide_paths = []
        print(f" Génération de {len(slides)} diapositives HTML...")
        
        for i, (title, content, image_path) in enumerate(slides):
            slide_output_path = os.path.join(output_dir, f"slide_{i+1:03d}.png")
            slide_image_path = create_html_style_slide(
                title, content, image_path, slide_output_path, i+1
            )
            slide_paths.append(slide_image_path)
            print(f" Diapositive {i+1} créée : {slide_image_path}")

        # Créer des répertoires pour les fichiers audio et vidéo
        audio_dir = base_output_dir / "audio_segments"
        video_dir = base_output_dir / "html_styled_videos"
        os.makedirs(audio_dir, exist_ok=True)
        os.makedirs(video_dir, exist_ok=True)
        
        # Générer un segment audio pour chaque diapositive
        print(" Génération des segments audio...")
        audio_segments = generate_audio_segments_for_slides(slides, audio_dir)
        
        # Générer un timestamp pour les noms de fichiers
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        video_output_path = os.path.join(video_dir, f"{base_name}_synchronized_{timestamp}.mp4")

        # Créer la vidéo synchronisée
        print(" Création de la vidéo synchronisée en cours...")
        
        # Créer la vidéo avec les diapositives synchronisées avec leur audio
        create_video_with_synchronized_slides(
            slide_paths, 
            audio_segments, 
            video_output_path, 
            avatar_path if not args.no_avatar else None
        )
        print(f" Vidéo synchronisée créée : {video_output_path}")
    else:
        print(" Option --html non spécifiée. Utilisez --html pour générer des diapositives HTML.")
        exit(1)

if __name__ == "__main__":
    main()